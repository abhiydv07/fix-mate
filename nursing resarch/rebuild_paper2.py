import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.oxml.ns import qn

print("Loading original document...")
original = Document('Document3.docx')

print(f"Original: {len(original.paragraphs)} paragraphs, {len(original.tables)} tables")

# Work with XML to preserve everything
body = original.element.body
children = list(body)

nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

def get_text(element):
    """Get all text from an XML element"""
    text = ''
    for t in element.iter(qn('w:t')):
        if t.text:
            text += t.text
    return text.strip()

# Remove consecutive empty paragraphs (keep max 1)
elements_to_remove = []
empty_streak = 0

for child in children:
    if child.tag.endswith('}p'):
        text = get_text(child)
        if not text:
            empty_streak += 1
            if empty_streak > 1:
                elements_to_remove.append(child)
        else:
            empty_streak = 0
    else:
        empty_streak = 0

print(f"Removing {len(elements_to_remove)} extra empty paragraphs...")

for elem in elements_to_remove:
    body.remove(elem)

# Count after cleanup
children_after = list(body)
table_count = sum(1 for c in children_after if c.tag.endswith('}tbl'))

print(f"After cleanup: {len(children_after)} children, {table_count} tables")

# Save with different filename
output_path = 'Document3_WithTables.docx'
original.save(output_path)

# Verify
print("\nVerifying...")
verify_doc = Document(output_path)
print(f"Final: {len(verify_doc.paragraphs)} paragraphs, {len(verify_doc.tables)} tables")

# Count images
image_count = 0
for rel in verify_doc.part.rels.values():
    if 'image' in rel.reltype:
        image_count += 1
print(f"Images: {image_count}")

# Word count
total_words = sum(len(p.text.split()) for p in verify_doc.paragraphs)
for table in verify_doc.tables:
    for row in table.rows:
        for cell in row.cells:
            total_words += len(cell.text.split())
print(f"Total words: {total_words}")

print(f"\nSaved to: {output_path}")
print("Done!")
