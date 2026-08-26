import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from lxml import etree
import copy

doc = Document('Document3_Final.docx')

print(f"Starting: {len(doc.paragraphs)} paragraphs")

# Access the document body XML
body = doc.element.body

# Get all paragraph elements
nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
all_paras = body.findall('.//w:p', nsmap)

print(f"Found {len(all_paras)} paragraph elements in XML")

# Identify empty paragraphs (no text content)
empty_paras = []
for p in all_paras:
    # Get all text within the paragraph
    text_nodes = p.findall('.//w:t', nsmap)
    text = ''.join([t.text or '' for t in text_nodes]).strip()
    if not text:
        empty_paras.append(p)

print(f"Found {len(empty_paras)} empty paragraphs")

# Keep only 1 empty paragraph between content sections
# First, identify all non-empty paragraphs and their positions
non_empty_positions = []
for i, p in enumerate(all_paras):
    text_nodes = p.findall('.//w:t', nsmap)
    text = ''.join([t.text or '' for t in text_nodes]).strip()
    if text:
        non_empty_positions.append(i)

print(f"Found {len(non_empty_positions)} non-empty paragraphs")

# Determine which empty paragraphs to keep (max 1 between non-empty)
to_remove = set()
empty_streak = 0

for i, p in enumerate(all_paras):
    text_nodes = p.findall('.//w:t', nsmap)
    text = ''.join([t.text or '' for t in text_nodes]).strip()
    
    if not text:
        empty_streak += 1
        if empty_streak > 1:
            to_remove.add(i)
    else:
        empty_streak = 0

print(f"Will remove {len(to_remove)} empty paragraphs")

# Remove identified empty paragraphs
removed_count = 0
for i in sorted(to_remove, reverse=True):
    p = all_paras[i]
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)
        removed_count += 1

print(f"Removed {removed_count} paragraphs from XML")

# Save and verify
output_path = 'C:\\Users\\abhiy\\OneDrive\\Desktop\\Document3_Final.docx'
doc.save(output_path)

# Verify
doc2 = Document(output_path)
total_paras = len(doc2.paragraphs)
empty_count = sum(1 for p in doc2.paragraphs if not p.text.strip())
total_words = sum(len(p.text.split()) for p in doc2.paragraphs)

print(f"\n=== FINAL VERIFICATION ===")
print(f"Total paragraphs: {total_paras}")
print(f"Empty paragraphs: {empty_count}")
print(f"Total words: {total_words}")
print(f"\nSaved to: {output_path}")
print("Done!")
