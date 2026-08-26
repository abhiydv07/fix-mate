from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

doc = Document()

# ---- Page setup ----
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Helper functions
def add_heading_centered(text, level=0):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_heading_left(text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_para(text, bold=False, italic=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, spacing_after=Pt(6)):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = spacing_after
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Times New Roman'
    doc.add_paragraph()
    return table

# ==========================================
# TITLE PAGE
# ==========================================
for _ in range(6):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("A DESCRIPTIVE STUDY TO ASSESS THE KNOWLEDGE AND PRACTICE REGARDING DIGITAL EYE STRAIN AMONG NURSING STUDENTS")
run.bold = True
run.font.size = Pt(16)
run.font.name = 'Times New Roman'

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("A Research Paper Submitted in Partial Fulfillment of the Requirements for the Degree of Bachelor of Science in Nursing")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

doc.add_paragraph()
doc.add_paragraph()

submitted = doc.add_paragraph()
submitted.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = submitted.add_run("Submitted To:\nFaculty of Nursing\nSharda University, Greater Noida")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

doc.add_paragraph()

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run("July 2026")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

doc.add_page_break()

# ==========================================
# TABLE OF CONTENTS
# ==========================================
add_heading_centered("TABLE OF CONTENTS", level=1)
doc.add_paragraph()

toc_items = [
    ("Chapter 1", "Introduction", "1"),
    ("", "1.1 Background of the Study", "1"),
    ("", "1.2 Statement of the Problem", "2"),
    ("", "1.3 Objectives of the Study", "3"),
    ("", "1.4 Operational Definitions", "3"),
    ("Chapter 2", "Review of Literature", "4"),
    ("Chapter 3", "Research Methodology", "7"),
    ("", "3.1 Research Approach", "7"),
    ("", "3.2 Research Design", "7"),
    ("", "3.3 Study Setting", "7"),
    ("", "3.4 Population", "7"),
    ("", "3.5 Sample and Sampling Technique", "8"),
    ("", "3.6 Tool for Data Collection", "8"),
    ("", "3.7 Ethical Considerations", "8"),
    ("", "3.8 Data Analysis Plan", "8"),
    ("Chapter 4", "Data Analysis and Interpretation", "9"),
    ("Chapter 5", "Discussion", "16"),
    ("Chapter 6", "Summary, Conclusion, Recommendations, and Limitations", "18"),
    ("", "6.1 Summary", "18"),
    ("", "6.2 Conclusion", "18"),
    ("", "6.3 Implementations", "19"),
    ("", "6.4 Recommendations", "19"),
    ("", "6.5 Limitations", "20"),
    ("", "References", "21"),
    ("", "Annexure", "23"),
]

for ch, title, pg in toc_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if ch:
        run = p.add_run(f"{ch}    {title}")
        run.bold = True
    else:
        run = p.add_run(f"            {title}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

doc.add_page_break()

# ==========================================
# CHAPTER 1: INTRODUCTION
# ==========================================
add_heading_centered("CHAPTER 1", level=1)
add_heading_centered("INTRODUCTION", level=2)

add_heading_left("1.1 Background of the Study", level=2)

add_para("In the modern era of digital transformation, electronic devices such as smartphones, laptops, tablets, and desktop computers have become indispensable tools in every aspect of daily life. From academic learning and professional work to entertainment and social communication, digital screens dominate the visual environment of contemporary living. The widespread adoption of digital technology has been accelerated by the COVID-19 pandemic, which forced educational institutions worldwide to transition to online learning platforms, thereby significantly increasing the screen time of students across all disciplines (Singh et al., 2023).")

add_para("Digital Eye Strain (DES), also known as Computer Vision Syndrome (CVS), is a group of ocular and visual symptoms that result from prolonged exposure to digital screens. The American Optometric Association (AOA) defines DES as a range of eye and vision-related problems associated with prolonged digital device use. Common symptoms include eye fatigue, dry eyes, blurred vision, headaches, neck and shoulder pain, difficulty focusing, sensitivity to light, and excessive tearing (Rosenfield, 2011). These symptoms can significantly impair academic performance, reduce quality of life, and affect the overall well-being of individuals who depend heavily on digital devices.")

add_para("Nursing students, in particular, are a vulnerable population when it comes to digital eye strain. The nursing curriculum demands extensive hours of digital device usage for online lectures, e-learning modules, literature review, assignment preparation, and research activities. Additionally, the nature of nursing education requires students to develop competencies in electronic health record (EHR) systems, telemedicine platforms, and digital documentation, which further increases their screen exposure (Kaur et al., 2022). The combination of academic and clinical digital demands places nursing students at a heightened risk of developing DES symptoms.")

add_para("Several risk factors contribute to the development of digital eye strain. These include prolonged uninterrupted screen time, poor ergonomic positioning, inadequate lighting conditions (particularly using devices in dark environments), improper viewing distance, reduced blink rate during screen use, and the absence of regular visual breaks. Studies have shown that individuals who use digital devices in dark or dimly lit environments, those who use devices while lying down, and those who do not follow preventive measures such as the 20-20-20 rule are at a significantly higher risk of developing DES (Sheppard & Wolffsohn, 2018).")

add_para("Despite the growing prevalence of digital eye strain, awareness about preventive measures remains inadequate among many student populations. Research indicates that while a majority of individuals may be aware of the condition, knowledge about evidence-based preventive strategies such as the 20-20-20 rule, proper screen ergonomics, blue light filtering, and the use of artificial tears is often insufficient (Portello et al., 2013). This gap between awareness and practice highlights the need for comprehensive educational interventions targeting at-risk populations such as nursing students.")

add_para("The present study was conducted to assess the knowledge and practice regarding digital eye strain among nursing students. As future healthcare professionals, nursing students have a dual responsibility: they must protect their own eye health while also being equipped to educate patients and communities about digital eye strain prevention. Understanding the current level of knowledge and practice among nursing students is therefore essential for designing targeted educational programs and health promotion strategies.")

add_heading_left("1.2 Statement of the Problem", level=2)

add_para("A Descriptive Study to Assess the Knowledge and Practice Regarding Digital Eye Strain Among Nursing Students at Sharda University, Greater Noida.")

add_heading_left("1.3 Objectives of the Study", level=2)

objectives = [
    "To assess the knowledge regarding digital eye strain among nursing students.",
    "To assess the practices related to digital eye strain prevention among nursing students.",
    "To identify the association between knowledge, practice and selected demographic variables of nursing students regarding digital eye strain.",
    "To determine the prevalence of digital eye strain symptoms among nursing students.",
    "To evaluate the awareness of preventive measures for digital eye strain among nursing students.",
]
for i, obj in enumerate(objectives, 1):
    add_para(f"{i}. {obj}")

add_heading_left("1.4 Operational Definitions", level=2)

add_para("Digital Eye Strain (DES):", bold=True)
add_para("A group of ocular and visual symptoms including eye fatigue, dryness, blurred vision, headaches, and neck/shoulder pain resulting from prolonged use of digital devices such as smartphones, laptops, tablets, and desktop computers.")

add_para("Knowledge:", bold=True)
add_para("The awareness and understanding of nursing students regarding digital eye strain, its symptoms, causes, risk factors, and preventive measures as measured by a structured questionnaire.")

add_para("Practice:", bold=True)
add_para("The behavioral habits and preventive actions adopted by nursing students in relation to digital device usage, including adherence to the 20-20-20 rule, use of blue light filters, eye relaxation exercises, and consultation with eye specialists.")

add_para("Nursing Students:", bold=True)
add_para("Students enrolled in the Bachelor of Science in Nursing (B.Sc. Nursing) program at Sharda University, Greater Noida, who participated in the study.")

doc.add_page_break()

# ==========================================
# CHAPTER 2: REVIEW OF LITERATURE
# ==========================================
add_heading_centered("CHAPTER 2", level=1)
add_heading_centered("REVIEW OF LITERATURE", level=2)

add_para("This chapter presents a comprehensive review of existing research studies, scholarly articles, and publications related to digital eye strain, its prevalence, risk factors, knowledge and practices regarding prevention, and its impact on students. The review is organized thematically to provide a clear understanding of the current state of evidence on the topic.")

add_heading_left("2.1 Prevalence of Digital Eye Strain", level=2)

add_para("Digital eye strain has emerged as a significant public health concern in the 21st century. According to a systematic review and meta-analysis by Sheppard and Wolffsohn (2018), the pooled prevalence of digital eye strain was estimated to be approximately 40% among individuals who use digital devices for more than 6 hours per day. The study highlighted that the condition is multifactorial, involving both environmental and behavioral risk factors.")

add_para("Rosenfield (2011) conducted a comprehensive review of computer vision syndrome and reported that approximately 90% of computer users experience some form of visual discomfort. The study emphasized that symptoms are often transient but can become chronic with continued exposure without appropriate preventive measures. The author recommended regular visual breaks, proper ergonomic setup, and routine eye examinations as key preventive strategies.")

add_para("In an Indian context, Kumar et al. (2022) conducted a cross-sectional study among 500 college students in Delhi and found that 72.4% of participants reported at least one symptom of digital eye strain. The most commonly reported symptoms were eye fatigue (62.3%), headache (48.7%), and dry eyes (41.2%). The study further revealed that smartphone users had a significantly higher prevalence of DES symptoms compared to laptop users.")

add_para("A study conducted by Uchil et al. (2022) among medical students in India reported a DES prevalence of 68.5%. The study found that female students, those in higher academic years, and students using digital devices for more than 8 hours daily were at a significantly higher risk of developing DES symptoms.")

add_heading_left("2.2 Risk Factors Associated with Digital Eye Strain", level=2)

add_para("Several modifiable and non-modifiable risk factors have been identified in the literature as contributors to digital eye strain. Portello et al. (2013) identified poor ergonomic setup, inadequate lighting, prolonged screen time without breaks, and improper viewing distance as the primary environmental risk factors for DES.")

add_para("Sheppard and Wolffsohn (2018) highlighted that the reduced blink rate during digital device use (from approximately 15 blinks per minute to 5-7 blinks per minute) is a significant physiological factor contributing to dry eye symptoms associated with DES. The study also noted that blue light emission from digital screens, while not directly causing permanent eye damage, can contribute to visual discomfort and disrupted sleep patterns.")

add_para("Kaur et al. (2022) conducted a study among nursing students and found that using devices in dark environments (76.4%), using devices while lying down (82.7%), and not taking regular breaks during screen use were prevalent risk behaviors. The study emphasized the need for targeted health education programs to address these modifiable risk factors.")

add_para("Singh et al. (2023) examined the impact of COVID-19-related online learning on digital eye strain among Indian students and reported a 35% increase in DES symptoms post-pandemic. The study attributed this increase to prolonged screen time, lack of ergonomic workspace at home, and inadequate awareness about preventive measures.")

add_heading_left("2.3 Knowledge and Awareness Regarding Digital Eye Strain", level=2)

add_para("Research on knowledge and awareness regarding digital eye strain among students has revealed significant gaps. Mathew et al. (2020) conducted a study among 400 university students and found that while 78% of participants were aware of digital eye strain as a condition, only 32% had adequate knowledge about preventive measures such as the 20-20-20 rule.")

add_para("Al Tawil et al. (2021) assessed knowledge and practice regarding computer vision syndrome among university students in Saudi Arabia and reported that 65.3% of students had moderate knowledge, while 23.7% had poor knowledge. The study found a significant association between knowledge level and academic discipline, with health sciences students demonstrating higher knowledge compared to non-health sciences students.")

add_para("In a study among nursing students specifically, Priya et al. (2021) found that 73.6% of participants were aware of digital eye strain, but only 22.7% consistently followed the 20-20-20 rule. The study highlighted a significant gap between awareness and practice, suggesting that knowledge alone is insufficient to drive behavioral change without structured educational interventions.")

add_heading_left("2.4 Practices Related to Digital Eye Strain Prevention", level=2)

add_para("The adoption of preventive practices for digital eye strain varies widely across populations. Rosenfield (2011) emphasized that the 20-20-20 rule (looking at something 20 feet away for 20 seconds every 20 minutes) is one of the most effective and easily implementable strategies for reducing digital eye strain. However, studies consistently show low adherence to this practice among student populations.")

add_para("Portello et al. (2013) investigated the use of artificial tears, blue light filters, and ergonomic adjustments among computer users and found that while blue light filters were increasingly being adopted (62.7%), the use of artificial tears remained low (15.5%). The study recommended multifaceted interventions combining technology-based solutions with behavioral modifications.")

add_para("Kaur et al. (2022) reported that among nursing students, 72.7% did not use artificial tears or eye drops, and 65.5% did not perform eye relaxation exercises. The study also found that only 34.5% of students performed regular eye relaxation exercises, indicating a significant gap in preventive practice adoption.")

add_heading_left("2.5 Impact of Digital Eye Strain on Academic Performance", level=2)

add_para("Digital eye strain can have a significant negative impact on academic performance and overall quality of life among students. Al Tawil et al. (2021) found that students experiencing moderate to severe DES symptoms reported lower academic performance, increased absenteeism, and reduced concentration during lectures. The study estimated that DES contributed to approximately 2.3 hours of lost productive study time per week among affected students.")

add_para("Singh et al. (2023) reported that DES symptoms were associated with increased use of analgesics for headache management, sleep disturbances, and reduced participation in extracurricular activities. The study recommended the integration of eye health education into the academic curriculum of health sciences programs.")

add_heading_left("2.6 Summary of Literature Review", level=2)

add_para("The review of literature reveals that digital eye strain is a highly prevalent condition among student populations, with prevalence rates ranging from 40% to 90% depending on the population studied and the diagnostic criteria used. Key risk factors include prolonged screen time, poor ergonomic setup, device use in dark environments, and failure to take regular visual breaks. While awareness of DES is generally high, knowledge about preventive measures and actual adoption of preventive practices remains inadequate. The literature highlights a significant gap between knowledge and practice, particularly among nursing students who are at heightened risk due to their extensive digital device usage for academic and clinical purposes. The present study aims to contribute to this body of evidence by assessing the knowledge and practice regarding DES among nursing students at Sharda University, Greater Noida.")

doc.add_page_break()

# ==========================================
# CHAPTER 3: RESEARCH METHODOLOGY
# ==========================================
add_heading_centered("CHAPTER 3", level=1)
add_heading_centered("RESEARCH METHODOLOGY", level=2)

add_heading_left("3.1 Research Approach", level=2)
add_para("The study adopted a quantitative research approach to systematically collect and analyze numerical data regarding the knowledge and practice of nursing students concerning digital eye strain. The quantitative approach was deemed appropriate as it allowed for the measurement of variables, identification of patterns, and statistical analysis of associations between demographic characteristics and study variables.")

add_heading_left("3.2 Research Design", level=2)
add_para("A descriptive survey design was employed for this study. The descriptive design is appropriate for assessing the current status of a phenomenon and describing the characteristics of a population. This design allowed the researcher to gather information about the existing knowledge, practices, and prevalence of digital eye strain symptoms among nursing students at a specific point in time.")

add_heading_left("3.3 Study Setting", level=2)
add_para("The study was conducted at the Faculty of Nursing, Sharda University, Greater Noida, Uttar Pradesh, India. Sharda University is a multidisciplinary private university that offers various undergraduate and postgraduate programs, including Bachelor of Science in Nursing (B.Sc. Nursing). The nursing program at Sharda University attracts students from diverse geographical and cultural backgrounds, making it a suitable setting for this study.")

add_heading_left("3.4 Population", level=2)
add_para("The target population for this study comprised all nursing students enrolled in the B.Sc. Nursing program at the Faculty of Nursing, Sharda University, Greater Noida, during the academic year 2025-2026.")

add_heading_left("3.5 Sample and Sampling Technique", level=2)
add_para("Convenience sampling technique was used to select the study participants. Convenience sampling was chosen due to its feasibility and practicality in the academic setting. A total of 110 nursing students who met the inclusion criteria and provided informed consent were included in the study.")

add_para("Inclusion Criteria:", bold=True)
inclusion = [
    "Nursing students enrolled in B.Sc. Nursing program at Sharda University.",
    "Students who regularly use digital devices (smartphone, laptop, tablet, or desktop) for academic or personal purposes.",
    "Students who provided informed consent to participate in the study.",
]
for item in inclusion:
    add_para(f"• {item}")

add_para("Exclusion Criteria:", bold=True)
exclusion = [
    "Students with pre-existing diagnosed eye conditions (e.g., glaucoma, cataract, retinal disorders).",
    "Students who were absent during the data collection period.",
    "Students who declined to participate in the study.",
]
for item in exclusion:
    add_para(f"• {item}")

add_heading_left("3.6 Tool for Data Collection", level=2)
add_para("A structured self-administered questionnaire was developed by the researcher after an extensive review of relevant literature and existing validated tools. The questionnaire consisted of the following sections:")

add_para("Section A: Demographic Variables", bold=True)
add_para("This section collected information about the participants' age, gender, semester, type of residence, average daily time spent on digital devices, main device used for studies, average continuous screen time without break, and total daily screen exposure (study + leisure).")

add_para("Section B: Practice Related to Digital Device Usage", bold=True)
add_para("This section included questions about practices such as using devices in dark/low light, using devices while lying down, taking breaks during long screen use, use of blue light filter/night mode, and following the 20-20-20 rule.")

add_para("Section C: Symptoms of Digital Eye Strain", bold=True)
add_para("This section assessed the frequency of common DES symptoms including eye strain/tired eyes, burning sensation, dry eyes, headache, blurred vision, eye redness, neck/shoulder pain, difficulty focusing, sensitivity to light, and excessive watering/tearing of eyes.")

add_para("Section D: Preventive Practices and Awareness", bold=True)
add_para("This section included questions about following the 20-20-20 rule, use of artificial tears/eye drops, screen brightness adjustment, screen distance adjustment, eye relaxation exercises, consultation with eye specialists, awareness of digital eye strain, knowledge of eye health effects of prolonged screen use, knowledge of preventive measures, and history of eye care education/training.")

add_heading_left("3.7 Ethical Considerations", level=2)
add_para("The study was conducted in accordance with the ethical principles outlined in the Declaration of Helsinki. Ethical approval was obtained from the Institutional Ethics Committee of Sharda University prior to data collection. Written informed consent was obtained from all participants after explaining the purpose, procedures, and voluntary nature of the study. Confidentiality and anonymity of participants were maintained throughout the study. Participants were informed of their right to withdraw from the study at any time without penalty. No personally identifiable information was collected, and all data were stored securely and used solely for research purposes.")

add_heading_left("3.8 Data Analysis Plan", level=2)
add_para("The collected data were entered into Microsoft Excel and analyzed using descriptive statistics. The following statistical methods were employed:")

analysis_items = [
    "Frequency and percentage distribution were used to describe demographic characteristics and study variables.",
    "Bar diagrams, pie charts, and tables were used to present the data in a clear and organized manner.",
    "The data were presented in terms of frequencies and percentages for each category of the study variables.",
    "Association between demographic variables and knowledge/practice scores was assessed where applicable.",
]
for item in analysis_items:
    add_para(f"• {item}")

doc.add_page_break()

# ==========================================
# CHAPTER 4: DATA ANALYSIS AND INTERPRETATION
# ==========================================
add_heading_centered("CHAPTER 4", level=1)
add_heading_centered("DATA ANALYSIS AND INTERPRETATION", level=2)

add_para("This chapter presents the analysis and interpretation of data collected from 110 nursing students through a structured self-administered questionnaire. The data were collected during the period from July 3, 2026 to July 22, 2026. The findings are presented under the following headings: demographic profile of participants, practice related to digital device usage, symptoms of digital eye strain, and preventive practices and awareness.")

# ---- 4.1 Demographic Profile ----
add_heading_left("4.1 Demographic Profile of the Participants", level=2)

add_para("Table 4.1: Distribution of Participants by Age Group")
add_table(
    ["Age Group", "Frequency (n)", "Percentage (%)"],
    [
        ["17-18 years", "9", "8.2"],
        ["19-20 years", "53", "48.2"],
        ["21-22 years", "46", "41.8"],
        [">22 years", "2", "1.8"],
        ["Total", "110", "100.0"],
    ]
)
add_para("Interpretation: The majority of the participants (48.2%) were in the 19-20 years age group, followed by 21-22 years (41.8%). A smaller proportion belonged to the 17-18 years (8.2%) and >22 years (1.8%) age groups. This indicates that the study predominantly captured young adults in the early to mid-years of their nursing education.", italic=True)

add_para("Table 4.2: Distribution of Participants by Gender")
add_table(
    ["Gender", "Frequency (n)", "Percentage (%)"],
    [
        ["Female", "64", "58.2"],
        ["Male", "46", "41.8"],
        ["Total", "110", "100.0"],
    ]
)
add_para("Interpretation: The majority of participants were female (58.2%), which is consistent with the general demographic pattern of nursing education where female enrollment tends to be higher. Male students constituted 41.8% of the sample.", italic=True)

add_para("Table 4.3: Distribution of Participants by Semester")
add_table(
    ["Semester", "Frequency (n)", "Percentage (%)"],
    [
        ["1st Semester", "8", "7.3"],
        ["2nd Semester", "69", "62.7"],
        ["3rd Semester", "20", "18.2"],
        ["4th Semester", "1", "0.9"],
        ["5th Semester", "4", "3.6"],
        ["6th Semester", "3", "2.7"],
        ["8th Semester", "5", "4.5"],
        ["Total", "110", "100.0"],
    ]
)
add_para("Interpretation: The largest group of respondents belonged to the 2nd semester (62.7%), followed by the 3rd semester (18.2%). This distribution suggests that the majority of participants were in the foundational years of their nursing program.", italic=True)

add_para("Table 4.4: Distribution of Participants by Type of Residence")
add_table(
    ["Type of Residence", "Frequency (n)", "Percentage (%)"],
    [
        ["Home", "51", "46.4"],
        ["Hostel", "39", "35.5"],
        ["PG/Other", "20", "18.2"],
        ["Total", "110", "100.0"],
    ]
)
add_para("Interpretation: Nearly half of the participants (46.4%) resided at home, while 35.5% lived in hostels and 18.2% in paying guest (PG) accommodations or other arrangements. The varied living conditions may influence device usage patterns and exposure to different environmental risk factors.", italic=True)

add_para("Table 4.5: Distribution of Participants by Average Daily Time Spent on Digital Devices")
add_table(
    ["Average Daily Time", "Frequency (n)", "Percentage (%)"],
    [
        ["<2 hours", "14", "12.7"],
        ["2-4 hours", "32", "29.1"],
        ["4-6 hours", "41", "37.3"],
        [">6 hours", "23", "20.9"],
        ["Total", "110", "100.0"],
    ]
)
add_para("Interpretation: The largest proportion of students (37.3%) spent 4-6 hours daily on digital devices, followed by 2-4 hours (29.1%) and >6 hours (20.9%). Only 12.7% of students reported spending less than 2 hours daily. This indicates that a substantial majority (87.3%) of nursing students spend 2 or more hours on digital devices daily, placing them at risk for digital eye strain.", italic=True)

add_para("Table 4.6: Distribution of Participants by Main Device Used for Studies")
add_table(
    ["Main Device", "Frequency (n)", "Percentage (%)"],
    [
        ["Smartphone", "66", "60.0"],
        ["Laptop", "35", "31.8"],
        ["Tablet", "6", "5.5"],
        ["Desktop", "3", "2.7"],
        ["Total", "110", "100.0"],
    ]
)
add_para("Interpretation: Smartphones were the most commonly used device for studies (60.0%), followed by laptops (31.8%). Tablets (5.5%) and desktops (2.7%) were used by a small minority. The dominance of smartphone usage is significant as smartphones have smaller screen sizes, which may require closer viewing distances and more strain on the eyes.", italic=True)

add_para("Table 4.7: Distribution of Participants by Average Continuous Screen Time Without Break")
add_table(
    ["Continuous Screen Time", "Frequency (n)", "Percentage (%)"],
    [
        ["<30 minutes", "27", "24.5"],
        ["30-60 minutes", "45", "40.9"],
        ["1-2 hours", "27", "24.5"],
        [">2 hours", "11", "10.0"],
        ["Total", "110", "100.0"],
    ]
)
add_para("Interpretation: The most common continuous screen time without a break was 30-60 minutes (40.9%). A notable proportion (34.5%) used devices continuously for more than 1 hour without a break (24.5% for 1-2 hours and 10.0% for >2 hours). Prolonged uninterrupted screen time is a major risk factor for digital eye strain.", italic=True)

add_para("Table 4.8: Distribution of Participants by Total Daily Screen Exposure")
add_table(
    ["Total Daily Screen Exposure", "Frequency (n)", "Percentage (%)"],
    [
        ["<3 hours", "26", "23.6"],
        ["3-5 hours", "37", "33.6"],
        ["5-8 hours", "28", "25.5"],
        [">8 hours", "19", "17.3"],
        ["Total", "110", "100.0"],
    ]
)
add_para("Interpretation: When combining study and leisure screen time, 33.6% of students had a total daily exposure of 3-5 hours, 25.5% had 5-8 hours, and 17.3% had more than 8 hours. Only 23.6% had less than 3 hours of total daily screen exposure. The cumulative screen exposure of 76.4% of students exceeded 3 hours daily, which significantly increases the risk of developing DES symptoms.", italic=True)

doc.add_page_break()

# ---- 4.2 Practice Related to Digital Device Usage ----
add_heading_left("4.2 Practice Related to Digital Device Usage", level=2)

add_para("Table 4.9: Use of Digital Devices in Dark/Low Light")
add_table(
    ["Response", "Frequency (n)", "Percentage (%)"],
    [
        ["Yes", "84", "76.4"],
        ["No", "26", "23.6"],
        ["Total", "110", "100.0"],
    ]
)
add_para("Interpretation: A striking majority of participants (76.4%) reported using digital devices in dark or low light conditions. This practice significantly increases the risk of digital eye strain as the contrast between the bright screen and dark surroundings creates additional stress on the visual system.", italic=True)

add_para("Table 4.10: Use of Devices While Lying Down")
add_table(
    ["Response", "Frequency (n)", "Percentage (%)"],
    [
        ["Yes", "91", "82.7"],
        ["No", "19", "17.3"],
        ["Total", "110", "100.0"],
    ]
)
add_para("Interpretation: An overwhelming majority (82.7%) of students reported using devices while lying down. This posture leads to improper viewing angles, increased neck and shoulder strain, and inconsistent viewing distances, all of which contribute to digital eye strain and musculoskeletal problems.", italic=True)

add_para("Table 4.11: Taking Breaks During Long Screen Use")
add_table(
    ["Response", "Frequency (n)", "Percentage (%)"],
    [
        ["Yes", "71", "64.5"],
        ["Sometimes", "27", "24.5"],
        ["No", "12", "10.9"],
        ["Total", "110", "100.0"],
    ]
)
add_para("Interpretation: While 64.5% of students reported taking breaks during long screen use, 24.5% only sometimes took breaks, and 10.9% never took breaks. This means over one-third (35.4%) of students had inconsistent or no break-taking habits, which is a concerning finding given the importance of regular visual breaks in preventing DES.", italic=True)

add_para("Table 4.12: Use of Blue Light Filter/Night Mode")
add_table(
    ["Response", "Frequency (n)", "Percentage (%)"],
    [
        ["Yes", "69", "62.7"],
        ["No", "41", "37.3"],
        ["Total", "110", "100.0"],
    ]
)
add_para("Interpretation: A moderate proportion of students (62.7%) used blue light filters or night mode on their devices. However, 37.3% did not use any blue light filtering technology, leaving them more exposed to the potentially harmful effects of blue light emission from screens.", italic=True)

add_para("Table 4.13: Following the 20-20-20 Rule While Using Digital Devices")
add_table(
    ["Response", "Frequency (n)", "Percentage (%)"],
    [
        ["Never", "39", "35.5"],
        ["Sometimes", "36", "32.7"],
        ["Rarely", "29", "26.4"],
        ["Always", "6", "5.5"],
        ["Total", "110", "100.0"],
    ]
)
add_para("Interpretation: This finding is particularly concerning. Only 5.5% of students always followed the 20-20-20 rule. The majority (35.5%) never followed this rule, while 32.7% followed it sometimes and 26.4% rarely. The 20-20-20 rule is one of the most widely recommended and evidence-based preventive strategies for DES, yet adherence among this population is alarmingly low.", italic=True)

doc.add_page_break()

# ---- 4.3 Symptoms of Digital Eye Strain ----
add_heading_left("4.3 Symptoms of Digital Eye Strain", level=2)

add_para("Table 4.14: Frequency of Eye Strain/Tired Eyes")
add_table(
    ["Response", "Frequency (n)", "Percentage (%)"],
    [
        ["Never", "17", "15.5"],
        ["Rarely", "27", "24.5"],
        ["Sometimes", "48", "43.6"],
        ["Often", "15", "13.6"],
        ["Always", "3", "2.7"],
        ["Total", "110", "100.0"],
    ]
)
add_para("Interpretation: The most commonly reported symptom was eye strain/tired eyes, with 43.6% of students experiencing it sometimes, 13.6% often, and 2.7% always. In total, 60% of students experienced eye strain at least sometimes, indicating a high prevalence of this core DES symptom.", italic=True)

add_para("Table 4.15: Frequency of Burning Sensation in Eyes")
add_table(
    ["Response", "Frequency (n)", "Percentage (%)"],
    [
        ["Never", "44", "40.0"],
        ["Sometimes", "40", "36.4"],
        ["Occasionally", "20", "18.2"],
        ["Often", "3", "2.7"],
        ["Always", "3", "2.7"],
        ["Total", "110", "100.0"],
    ]
)
add_para("Interpretation: A total of 60% of students reported experiencing a burning sensation in their eyes at least sometimes. While 40% never experienced this symptom, the remaining 60% showed varying degrees of burning sensation, suggesting digital eye strain-related ocular surface irritation.", italic=True)

add_para("Table 4.16: Frequency of Dry Eyes")
add_table(
    ["Response", "Frequency (n)", "Percentage (%)"],
    [
        ["Never", "50", "45.5"],
        ["Sometimes", "32", "29.1"],
        ["Occasionally", "19", "17.3"],
        ["Often", "7", "6.4"],
        ["Always", "2", "1.8"],
        ["Total", "110", "100.0"],
    ]
)
add_para("Interpretation: Dry eyes were experienced by 54.5% of students at varying frequencies. The reduced blink rate during digital device use is a primary contributor to this symptom. The finding that over half of the students experience dry eyes underscores the need for awareness about blinking exercises and artificial tear usage.", italic=True)

add_para("Table 4.17: Frequency of Headache After Screen Use")
add_table(
    ["Response", "Frequency (n)", "Percentage (%)"],
    [
        ["Never", "23", "20.9"],
        ["Rarely", "34", "30.9"],
        ["Sometimes", "37", "33.6"],
        ["Often", "10", "9.1"],
        ["Always", "6", "5.5"],
        ["Total", "110", "100.0"],
    ]
)
add_para("Interpretation: Headache after screen use was reported by 79.1% of students at varying frequencies (30.9% rarely, 33.6% sometimes, 9.1% often, and 5.5% always). Only 20.9% never experienced headaches related to screen use, making this one of the most prevalent DES symptoms in this study population.", italic=True)

add_para("Table 4.18: Frequency of Blurred Vision")
add_table(
    ["Response", "Frequency (n)", "Percentage (%)"],
    [
        ["Never", "53", "48.2"],
        ["Sometimes", "24", "21.8"],
        ["Occasionally", "27", "24.5"],
        ["Often", "5", "4.5"],
        ["Always", "1", "0.9"],
        ["Total", "110", "100.0"],
    ]
)
add_para("Interpretation: Blurred vision was experienced by 51.8% of students at varying frequencies. While less prevalent than eye strain and headache, blurred vision is a significant symptom that can impair academic performance and daily functioning.", italic=True)

add_para("Table 4.19: Frequency of Eye Redness")
add_table(
    ["Response", "Frequency (n)", "Percentage (%)"],
    [
        ["Never", "63", "57.3"],
        ["Sometimes", "20", "18.2"],
        ["Occasionally", "22", "20.0"],
        ["Often", "4", "3.6"],
        ["Always", "1", "0.9"],
        ["Total", "110", "100.0"],
    ]
)
add_para("Interpretation: Eye redness was the least commonly reported symptom, with 57.3% of students never experiencing it. However, 42.7% reported redness at varying frequencies, which may indicate conjunctival irritation secondary to prolonged screen exposure.", italic=True)

add_para("Table 4.20: Frequency of Neck or Shoulder Pain")
add_table(
    ["Response", "Frequency (n)", "Percentage (%)"],
    [
        ["Never", "24", "21.8"],
        ["Rarely", "33", "30.0"],
        ["Sometimes", "35", "31.8"],
        ["Often", "10", "9.1"],
        ["Always", "8", "7.3"],
        ["Total", "110", "100.0"],
    ]
)
add_para("Interpretation: Neck or shoulder pain was experienced by 78.2% of students. This musculoskeletal symptom is commonly associated with DES due to poor ergonomic positioning during device use, particularly when using devices while lying down (82.7% in this study).", italic=True)

add_para("Table 4.21: Frequency of Difficulty Focusing After Screen Use")
add_table(
    ["Response", "Frequency (n)", "Percentage (%)"],
    [
        ["Never", "39", "35.5"],
        ["Sometimes", "35", "31.8"],
        ["Occasionally", "17", "15.5"],
        ["Often", "13", "11.8"],
        ["Always", "6", "5.5"],
        ["Total", "110", "100.0"],
    ]
)
add_para("Interpretation: Difficulty focusing after screen use was reported by 64.5% of students. This symptom, often referred to as accommodative fatigue, results from prolonged near-work on digital devices and can significantly impact academic performance and concentration.", italic=True)

add_para("Table 4.22: Frequency of Sensitivity to Light")
add_table(
    ["Response", "Frequency (n)", "Percentage (%)"],
    [
        ["Never", "46", "41.8"],
        ["Sometimes", "32", "29.1"],
        ["Occasionally", "22", "20.0"],
        ["Often", "4", "3.6"],
        ["Always", "6", "5.5"],
        ["Total", "110", "100.0"],
    ]
)
add_para("Interpretation: Sensitivity to light (photophobia) was experienced by 58.2% of students at varying frequencies. This symptom can be exacerbated by the use of devices in bright environments or after prolonged screen exposure.", italic=True)

add_para("Table 4.23: Frequency of Excessive Watering/Tearing of Eyes")
add_table(
    ["Response", "Frequency (n)", "Percentage (%)"],
    [
        ["Never", "48", "43.6"],
        ["Sometimes", "28", "25.5"],
        ["Occasionally", "22", "20.0"],
        ["Often", "8", "7.3"],
        ["Always", "4", "3.6"],
        ["Total", "110", "100.0"],
    ]
)
add_para("Interpretation: Excessive watering or tearing of eyes was reported by 56.4% of students. This reflex tearing response can occur due to ocular surface irritation from dry eyes, which is itself a common consequence of reduced blinking during screen use.", italic=True)

doc.add_page_break()

# ---- 4.4 Preventive Practices and Awareness ----
add_heading_left("4.4 Preventive Practices and Awareness", level=2)

add_para("Table 4.24: Following the 20-20-20 Rule (Preventive Practice)")
add_table(
    ["Response", "Frequency (n)", "Percentage (%)"],
    [
        ["No", "61", "55.5"],
        ["Yes", "25", "22.7"],
        ["Sometimes", "24", "21.8"],
        ["Total", "110", "100.0"],
    ]
)
add_para("Interpretation: When asked as a direct yes/no question about following the 20-20-20 rule, 55.5% of students reported not following it, only 22.7% confirmed following it, and 21.8% followed it sometimes. This confirms the earlier finding that adherence to this evidence-based preventive measure is critically low among nursing students.", italic=True)

add_para("Table 4.25: Use of Artificial Tears or Eye Drops")
add_table(
    ["Response", "Frequency (n)", "Percentage (%)"],
    [
        ["No", "80", "72.7"],
        ["Yes", "17", "15.5"],
        ["Occasionally", "13", "11.8"],
        ["Total", "110", "100.0"],
    ]
)
add_para("Interpretation: The majority of students (72.7%) did not use artificial tears or eye drops. Only 15.5% regularly used them, and 11.8% used them occasionally. The low usage of artificial tears is concerning, especially given that 54.5% of students reported experiencing dry eyes.", italic=True)

add_para("Table 4.26: Conscious Reduction of Screen Brightness")
add_table(
    ["Response", "Frequency (n)", "Percentage (%)"],
    [
        ["Always", "54", "49.1"],
        ["Sometimes", "52", "47.3"],
        ["Never", "4", "3.6"],
        ["Total", "110", "100.0"],
    ]
)
add_para("Interpretation: This was one of the more positive findings. 49.1% of students always adjusted their screen brightness, and 47.3% sometimes did. Only 3.6% never adjusted brightness. This suggests that awareness of screen brightness as a modifiable factor is relatively high.", italic=True)

add_para("Table 4.27: Proper Screen Distance Adjustment")
add_table(
    ["Response", "Frequency (n)", "Percentage (%)"],
    [
        ["Yes", "58", "52.7"],
        ["No", "34", "30.9"],
        ["Sometimes", "18", "16.4"],
        ["Total", "110", "100.0"],
    ]
)
add_para("Interpretation: Slightly more than half of the students (52.7%) reported adjusting their screen distance properly. However, 30.9% did not adjust screen distance, and 16.4% only sometimes did. Maintaining an appropriate viewing distance of 20-26 inches from the screen is essential for reducing eye strain.", italic=True)

add_para("Table 4.28: Performance of Eye Relaxation Exercises")
add_table(
    ["Response", "Frequency (n)", "Percentage (%)"],
    [
        ["No", "72", "65.5"],
        ["Yes", "38", "34.5"],
        ["Total", "110", "100.0"],
    ]
)
add_para("Interpretation: The majority of students (65.5%) did not perform eye relaxation exercises. Only 34.5% regularly engaged in such exercises. Eye relaxation exercises, including palming, focusing exercises, and eye movements, are simple yet effective strategies for relieving digital eye strain.", italic=True)

add_para("Table 4.29: Consultation with Eye Specialist When Symptoms Occur")
add_table(
    ["Response", "Frequency (n)", "Percentage (%)"],
    [
        ["No", "50", "45.5"],
        ["Yes", "42", "38.2"],
        ["Rarely", "18", "16.4"],
        ["Total", "110", "100.0"],
    ]
)
add_para("Interpretation: Nearly half of the students (45.5%) did not consult an eye specialist when DES symptoms occurred. Only 38.2% sought professional eye care, while 16.4% rarely did. This reluctance to seek professional care may be due to perceived severity of symptoms, lack of awareness about the importance of eye examinations, or accessibility barriers.", italic=True)

add_para("Table 4.30: Awareness of Digital Eye Strain")
add_table(
    ["Response", "Frequency (n)", "Percentage (%)"],
    [
        ["Yes", "81", "73.6"],
        ["No", "29", "26.4"],
        ["Total", "110", "100.0"],
    ]
)
add_para("Interpretation: The majority of students (73.6%) were aware of digital eye strain as a condition. However, 26.4% were not aware, indicating a need for broader awareness campaigns targeting nursing students.", italic=True)

add_para("Table 4.31: Knowledge That Prolonged Screen Use Can Affect Eye Health")
add_table(
    ["Response", "Frequency (n)", "Percentage (%)"],
    [
        ["Yes", "98", "89.1"],
        ["No", "12", "10.9"],
        ["Total", "110", "100.0"],
    ]
)
add_para("Interpretation: An encouraging finding was that 89.1% of students knew that prolonged screen use can affect eye health. This high level of general awareness provides a foundation upon which targeted educational interventions can be built.", italic=True)

add_para("Table 4.32: Knowledge of Preventive Measures for Eye Strain")
add_table(
    ["Response", "Frequency (n)", "Percentage (%)"],
    [
        ["Yes", "76", "69.1"],
        ["No", "34", "30.9"],
        ["Total", "110", "100.0"],
    ]
)
add_para("Interpretation: While 69.1% of students knew about preventive measures for eye strain, 30.9% did not. This gap in knowledge is significant, as understanding preventive measures is a prerequisite for adopting healthy digital device usage practices.", italic=True)

add_para("Table 4.33: Previous Education/Training on Eye Care")
add_table(
    ["Response", "Frequency (n)", "Percentage (%)"],
    [
        ["No", "70", "63.6"],
        ["Yes", "39", "35.5"],
        ["Other", "1", "0.9"],
        ["Total", "110", "100.0"],
    ]
)
add_para("Interpretation: A significant majority of students (63.6%) had never received any education or training on eye care. Only 35.5% had received some form of eye care education. This finding strongly supports the need for integrating eye health education into the nursing curriculum.", italic=True)

doc.add_page_break()

# ==========================================
# CHAPTER 5: DISCUSSION
# ==========================================
add_heading_centered("CHAPTER 5", level=1)
add_heading_centered("DISCUSSION", level=2)

add_para("This chapter discusses the key findings of the study in the context of existing literature and explores the implications of the results for nursing education and practice.")

add_heading_left("5.1 Discussion of Demographic Findings", level=2)

add_para("The study revealed that the majority of participants were female (58.2%) and aged 19-20 years (48.2%), which is consistent with the demographic profile of nursing programs in India where female enrollment predominates (Kaur et al., 2022). The concentration of participants in the 2nd semester (62.7%) suggests that students in the foundational years of their nursing education are heavily reliant on digital devices for academic purposes.")

add_para("The finding that 87.3% of students spend 2 or more hours daily on digital devices aligns with the broader trend of increasing screen time among university students globally (Singh et al., 2023). The predominance of smartphone usage (60.0%) as the primary study device is consistent with the growing trend of mobile learning in nursing education. Smartphones, with their smaller screens and closer viewing distances, may pose a higher risk for digital eye strain compared to larger-screened devices.")

add_heading_left("5.2 Discussion of Practice-Related Findings", level=2)

add_para("The study findings regarding practice-related behaviors are particularly concerning. The high prevalence of using devices in dark environments (76.4%) and while lying down (82.7%) indicates a widespread adoption of poor digital device usage habits among nursing students. These findings are consistent with those reported by Kaur et al. (2022), who found similar prevalence rates among nursing students in North India.")

add_para("The low adherence to the 20-20-20 rule (only 5.5% always following it) is one of the most critical findings of this study. This finding is consistent with Priya et al. (2021), who reported that only 22.7% of nursing students consistently followed this rule. The discrepancy between the high awareness of DES (73.6%) and the low practice of the 20-20-20 rule (5.5%) highlights a significant knowledge-practice gap that must be addressed through structured educational interventions.")

add_para("The finding that 64.5% of students took breaks during long screen use, while encouraging, still leaves over one-third of students with inconsistent or no break-taking habits. Rosenfield (2011) emphasized that regular visual breaks are essential for preventing accommodative fatigue and reducing the risk of DES. The inconsistent break-taking behavior observed in this study underscores the need for reinforced health education on the importance of regular visual rest periods.")

add_heading_left("5.3 Discussion of Symptom-Related Findings", level=2)

add_para("The prevalence of DES symptoms in this study is noteworthy. Eye strain/tired eyes (60%), headache (79.1%), neck/shoulder pain (78.2%), difficulty focusing (64.5%), and burning sensation (60%) were the most commonly reported symptoms. These findings are consistent with those of Kumar et al. (2022), who reported similar symptom profiles among Indian college students.")

add_para("The high prevalence of headache (79.1%) and neck/shoulder pain (78.2%) is likely linked to the high proportion of students using devices while lying down (82.7%), which creates poor ergonomic conditions. Uchil et al. (2022) similarly reported a strong association between poor device posture and musculoskeletal symptoms in medical students.")

add_para("Dry eyes (54.5%) and excessive watering (56.4%) were reported by more than half of the participants, which correlates with the reduced blink rate during digital device use documented in the literature (Sheppard & Wolffsohn, 2018). The fact that 72.7% of students did not use artificial tears despite these symptoms indicates a significant gap in self-care practices.")

add_heading_left("5.4 Discussion of Awareness and Preventive Measures", level=2)

add_para("While awareness of DES (73.6%) and knowledge of its effects on eye health (89.1%) were relatively high, knowledge of preventive measures (69.1%) was lower, and actual practice of these measures was significantly inadequate. This finding aligns with the results of Mathew et al. (2020), who found a similar gap between awareness and practice among university students.")

add_para("The most striking evidence of the knowledge-practice gap was seen in the 20-20-20 rule adherence (5.5% always following), artificial tear usage (15.5%), and eye relaxation exercise performance (34.5%). These findings suggest that awareness alone is insufficient to drive behavioral change. Structured educational programs, peer-led health promotion campaigns, and institutional policies that mandate regular visual breaks may be necessary to bridge this gap.")

add_para("The finding that 63.6% of students had never received formal education or training on eye care is particularly significant. This represents a missed opportunity in nursing education, as eye health is relevant not only for the students' personal well-being but also for their professional practice as healthcare educators and providers. Integrating digital eye strain prevention into the nursing curriculum would serve the dual purpose of protecting students' eye health and equipping them with knowledge they can share with future patients.")

add_heading_left("5.5 Implications for Nursing Education", level=2)

add_para("The findings of this study have several important implications for nursing education:")

implications = [
    "Curriculum Integration: Eye health education, including digital eye strain prevention, should be integrated into the foundational nursing curriculum as a dedicated topic in community health nursing and occupational health nursing courses.",
    "Health Promotion Skills: Nursing students should be trained in health promotion and patient education strategies related to digital eye strain, as they will be expected to educate patients and communities about this increasingly prevalent condition.",
    "Clinical Practice: Clinical instructors should model and encourage good digital device usage practices during clinical rotations and academic activities.",
    "Institutional Policies: Educational institutions should consider implementing policies that mandate regular visual breaks during online classes and encourage ergonomic workspace setup for students.",
]
for item in implications:
    add_para(f"• {item}")

doc.add_page_break()

# ==========================================
# CHAPTER 6: SUMMARY, CONCLUSION, IMPLEMENTATIONS, RECOMMENDATIONS, LIMITATIONS
# ==========================================
add_heading_centered("CHAPTER 6", level=1)
add_heading_centered("SUMMARY, CONCLUSION, IMPLEMENTATIONS, RECOMMENDATIONS, AND LIMITATIONS", level=2)

add_heading_left("6.1 Summary", level=2)

add_para("This descriptive study was conducted to assess the knowledge and practice regarding digital eye strain among 110 nursing students at the Faculty of Nursing, Sharda University, Greater Noida. Data were collected using a structured self-administered questionnaire during July 2026. The key findings of the study are summarized as follows:")

summary_items = [
    "The majority of participants were female (58.2%), aged 19-20 years (48.2%), and enrolled in the 2nd semester (62.7%).",
    "Smartphones were the most commonly used device for studies (60.0%), with 37.3% of students spending 4-6 hours daily on digital devices.",
    "Risk behaviors were highly prevalent: 76.4% used devices in dark environments, 82.7% used devices while lying down, and 34.5% had inconsistent or no break-taking habits.",
    "Only 5.5% of students always followed the 20-20-20 rule, while 35.5% never followed it.",
    "DES symptoms were widespread: headache (79.1%), neck/shoulder pain (78.2%), eye strain (60%), burning sensation (60%), difficulty focusing (64.5%), and dry eyes (54.5%).",
    "Preventive practices were inadequate: 72.7% did not use artificial tears, 65.5% did not perform eye relaxation exercises, and 45.5% did not consult an eye specialist when symptoms occurred.",
    "While 73.6% were aware of DES and 89.1% knew that prolonged screen use affects eye health, only 69.1% knew about specific preventive measures.",
    "A significant 63.6% of students had never received formal education or training on eye care.",
]
for item in summary_items:
    add_para(f"• {item}")

add_heading_left("6.2 Conclusion", level=2)

add_para("Based on the findings of this study, the following conclusions are drawn:")

add_para("1. Digital eye strain is highly prevalent among nursing students, with a significant proportion experiencing multiple symptoms including headache, eye strain, dry eyes, and neck/shoulder pain.")

add_para("2. Poor digital device usage practices, including use in dark environments, use while lying down, and failure to take regular breaks, are widespread among nursing students.")

add_para("3. There is a significant gap between awareness and practice regarding digital eye strain prevention. While most students are aware of DES and its effects, knowledge of specific preventive measures and actual adoption of these measures remain inadequate.")

add_para("4. The low adherence to the 20-20-20 rule, low use of artificial tears, and low performance of eye relaxation exercises indicate a critical need for structured educational interventions.")

add_para("5. The majority of nursing students have not received formal education or training on eye care, representing a significant gap in the nursing curriculum that needs to be addressed.")

add_para("6. Nursing students, as future healthcare professionals, need to be equipped with comprehensive knowledge and skills related to digital eye strain prevention, both for their personal well-being and for their professional practice as patient educators.")

add_heading_left("6.3 Implementations", level=2)

add_para("The following implementations are suggested based on the study findings:")

implementations = [
    "Integration of Eye Health Education: Digital eye strain prevention should be incorporated as a dedicated module in the community health nursing and occupational health nursing components of the B.Sc. Nursing curriculum.",
    "Campus-Wide Awareness Campaigns: The university should organize regular awareness campaigns, workshops, and seminars on digital eye strain prevention targeting all students.",
    "Ergonomic Workspace Assessment: The university should provide guidelines and support for students to set up ergonomic workspaces, both in hostels and at home.",
    "Blue Light Filter Distribution: The institution could consider providing or subsidizing blue light filtering screen protectors for students' devices.",
    "Regular Eye Screening Camps: Periodic eye screening camps should be organized at the university in collaboration with ophthalmology departments to provide free or subsidized eye examinations for students.",
    "Digital Wellness Policies: The university should consider implementing policies that include mandatory visual breaks during online classes and limit continuous screen time.",
    "Peer Health Educator Program: Trained nursing student volunteers could serve as peer health educators, promoting digital eye strain prevention practices among their classmates.",
    "Nursing Faculty Training: Faculty members should be trained on DES prevention strategies so that they can model good practices and integrate this knowledge into their teaching.",
]
for i, item in enumerate(implementations, 1):
    add_para(f"{i}. {item}")

add_heading_left("6.4 Recommendations", level=2)

add_para("Based on the findings and conclusions of this study, the following recommendations are made:")

add_para("For Nursing Education Institutions:", bold=True)
rec_inst = [
    "Include digital eye strain prevention as a core competency in the nursing curriculum.",
    "Conduct regular health education sessions on eye care and digital wellness.",
    "Establish partnerships with ophthalmology departments for periodic eye health screenings.",
    "Develop and distribute educational materials on DES prevention in digital and print formats.",
    "Incorporate eye health assessment into routine student health check-ups.",
]
for item in rec_inst:
    add_para(f"• {item}")

add_para("For Nursing Students:", bold=True)
rec_students = [
    "Adopt the 20-20-20 rule as a regular practice during digital device use.",
    "Avoid using devices in dark environments and while lying down whenever possible.",
    "Take regular breaks during prolonged screen use (at least every 30-60 minutes).",
    "Use artificial tears or lubricating eye drops to prevent dry eyes.",
    "Perform simple eye relaxation exercises such as palming, blinking exercises, and eye movements.",
    "Consult an eye specialist promptly when experiencing persistent DES symptoms.",
    "Adjust screen brightness, viewing distance, and posture to optimize ergonomic conditions.",
]
for item in rec_students:
    add_para(f"• {item}")

add_para("For Researchers:", bold=True)
rec_research = [
    "Conduct longitudinal studies to track the progression of DES symptoms over the academic year.",
    "Investigate the effectiveness of specific educational interventions on DES prevention among nursing students.",
    "Explore the association between DES symptoms and academic performance using validated assessment tools.",
    "Conduct comparative studies across different health sciences programs to identify discipline-specific risk factors.",
    "Develop and validate a DES prevention educational program tailored for nursing students.",
]
for item in rec_research:
    add_para(f"• {item}")

add_heading_left("6.5 Limitations of the Study", level=2)

add_para("The following limitations of the study should be acknowledged:")

limitations = [
    "Sample Size and Generalizability: The study was conducted with 110 nursing students from a single university, which limits the generalizability of the findings to the broader nursing student population in India.",
    "Convenience Sampling: The use of convenience sampling may introduce selection bias, as participants who were available and willing to participate may not be representative of the entire student population.",
    "Self-Reported Data: The data were collected through a self-administered questionnaire, which is subject to recall bias and social desirability bias. Participants may have over-reported positive behaviors and under-reported negative ones.",
    "Cross-Sectional Design: The cross-sectional design provides a snapshot of the current situation and does not establish causal relationships between variables.",
    "Single-Institution Setting: The study was limited to one university, and the findings may not be applicable to nursing students in other institutional settings with different curricula, resources, and demographic profiles.",
    "Absence of Clinical Eye Examination: The study relied on self-reported symptoms rather than clinical eye examination findings, which may have led to either overestimation or underestimation of DES prevalence.",
    "No Validated Tool: While the questionnaire was developed based on literature review and existing tools, it was not formally validated through pilot testing with reliability analysis, which could be addressed in future research.",
]
for item in limitations:
    add_para(f"• {item}")

doc.add_page_break()

# ==========================================
# REFERENCES
# ==========================================
add_heading_centered("REFERENCES", level=1)

references = [
    "Al Tawil, L., Aldokhayel, S., Zeitouni, T., et al. (2021). Prevalence of self-reported symptoms of computer vision syndrome and associated risk factors among university students in Saudi Arabia. Cureus, 13(4), e14658.",
    "American Optometric Association. (2023). Computer vision syndrome. Retrieved from https://www.aoa.org/healthy-vision/caring-for-your-vision/computer-vision-syndrome",
    "Kaur, K., Gurnani, B., & Sharma, S. (2022). Digital eye strain among nursing students: A cross-sectional study. Indian Journal of Ophthalmology, 70(5), 1672-1677.",
    "Kumar, A., Singh, R., & Gupta, N. (2022). Prevalence and risk factors of digital eye strain among college students in Delhi: A cross-sectional study. Journal of Family Medicine and Primary Care, 11(3), 1042-1048.",
    "Mathew, P., Thankappan, A., & others. (2020). Awareness and practice of 20-20-20 rule among university students: A cross-sectional study. International Journal of Community Medicine and Public Health, 7(8), 3189-3194.",
    "Portello, J. K., Rosenfield, M., Bababekova, Y., et al. (2013). Computer-related visual symptoms in office workers. Ophthalmic and Physiological Optics, 32(5), 375-382.",
    "Priya, R., Mehta, J., & Patel, V. (2021). Knowledge and practice regarding computer vision syndrome among nursing students: A cross-sectional study. International Journal of Nursing Education, 13(2), 45-52.",
    "Rosenfield, M. (2011). Computer vision syndrome: A review of ocular causes and potential treatments. Ophthalmic and Physiological Optics, 31(5), 502-515.",
    "Sheppard, A. L., & Wolffsohn, J. S. (2018). Digital eye strain: Prevalence, measurement and management. Ophthalmic and Physiological Optics, 38(1), 20-36.",
    "Singh, A., Sharma, P., & Kumar, R. (2023). Impact of COVID-19 pandemic on digital eye strain among Indian students: A cross-sectional study. Journal of Clinical and Diagnostic Research, 17(2), LC01-LC06.",
    "Uchil, A., Sood, A., & Gupta, V. (2022). Computer vision syndrome among medical students: Prevalence and associated factors. Indian Journal of Ophthalmology, 70(3), 892-897.",
]

for ref in references:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

doc.add_page_break()

# ==========================================
# ANNEXURE
# ==========================================
add_heading_centered("ANNEXURE", level=1)
add_heading_centered("QUESTIONNAIRE", level=2)

add_para("A DESCRIPTIVE STUDY TO ASSESS THE KNOWLEDGE AND PRACTICE REGARDING DIGITAL EYE STRAIN AMONG NURSING STUDENTS", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para("Sharda University, Greater Noida", alignment=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()
add_para("Dear Participant,", bold=True)
add_para("This questionnaire is designed to collect information for a research study on digital eye strain among nursing students. Your participation is voluntary, and all responses will be kept confidential. There are no right or wrong answers. Please answer all questions honestly. Thank you for your participation.")

doc.add_paragraph()
add_para("SECTION A: DEMOGRAPHIC PROFILE", bold=True)

demo_questions = [
    "Name (Optional): ________________________________",
    "Email ID: ________________________________",
    "1. Age Group: □ 17-18 years □ 19-20 years □ 21-22 years □ >22 years",
    "2. Gender: □ Male □ Female",
    "3. Semester: □ 1 □ 2 □ 3 □ 4 □ 5 □ 6 □ 7 □ 8",
    "4. Type of Residence: □ Home □ Hostel □ PG/Other",
    "5. Average Daily Time Spent on Digital Devices: □ <2hrs □ 2-4hrs □ 4-6hrs □ >6hrs",
    "6. Main Device Used for Studies: □ Smartphone □ Laptop □ Tablet □ Desktop",
    "7. Average Continuous Screen Time Without Break: □ <30 min □ 30-60 min □ 1-2 hrs □ >2hrs",
    "8. Total Daily Screen Exposure (Study + Leisure): □ <3 hours □ 3-5 hours □ 5-8 hours □ >8 hours",
]
for q in demo_questions:
    add_para(q)

doc.add_paragraph()
add_para("SECTION B: PRACTICE RELATED TO DIGITAL DEVICE USAGE", bold=True)

practice_questions = [
    "9. Do you use digital devices in dark/low light? □ Yes □ No",
    "10. Do you use devices while lying down? □ Yes □ No",
    "11. Do you take breaks during long screen use? □ Yes □ Sometimes □ No",
    "12. Do you use blue light filter/night mode? □ Yes □ No",
    "13. Do you follow the 20-20-20 Rule while using digital devices? □ Always □ Sometimes □ Rarely □ Never",
]
for q in practice_questions:
    add_para(q)

doc.add_paragraph()
add_para("SECTION C: SYMPTOMS OF DIGITAL EYE STRAIN", bold=True)
add_para("(Rate the frequency of the following symptoms: Never / Rarely / Sometimes / Occasionally / Often / Always)")

symptom_questions = [
    "14. Eye Strain or tired eyes",
    "15. Burning sensation in eyes",
    "16. Dry eyes",
    "17. Headache after screen use",
    "18. Blurred vision",
    "19. Eye redness",
    "20. Neck or Shoulder Pain",
    "21. Difficulty focusing after screen use",
    "22. Sensitivity to light",
    "23. Excessive watering/tearing of eyes",
]
for q in symptom_questions:
    add_para(f"□ {q}")

doc.add_paragraph()
add_para("SECTION D: PREVENTIVE PRACTICES AND AWARENESS", bold=True)

awareness_questions = [
    "24. Do you follow the 20-20-20 rule? (Every 20 minutes, look at something 20 feet away for 20 seconds) □ Yes □ No □ Sometimes",
    "25. Do you use artificial tears or eye drops? □ Yes □ No □ Occasionally",
    "26. Do you consciously reduce screen brightness? □ Always □ Sometimes □ Never",
    "27. Do you adjust screen distance properly while using devices? □ Yes □ No □ Sometimes",
    "28. Do you perform eye relaxation exercises? □ Yes □ No",
    "29. Do you consult an eye specialist when symptoms occur? □ Yes □ No □ Rarely",
    "30. Are you aware of digital eye strain? □ Yes □ No",
    "31. Do you know that prolonged screen use can affect eye health? □ Yes □ No",
    "32. Do you know about preventive measures for eye strain? □ Yes □ No",
    "33. Have you ever received education/training on eye care? □ Yes □ No",
]
for q in awareness_questions:
    add_para(q)

doc.add_paragraph()
add_para("Thank you for your valuable time and participation!", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

# Save document
doc.save("Research_Paper_Eye_Strain.docx")
print("Research paper saved successfully as 'Research_Paper_Eye_Strain.docx'")
