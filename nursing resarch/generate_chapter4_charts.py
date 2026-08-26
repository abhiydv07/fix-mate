import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

# ---- Chart styling ----
plt.rcParams['font.family'] = 'DejaVu Sans'
plt.rcParams['font.size'] = 11
plt.rcParams['axes.titlesize'] = 13
plt.rcParams['axes.labelsize'] = 11

# Color palettes
bar_colors_3 = ['#2196F3', '#FF9800', '#4CAF50']
bar_colors_4 = ['#2196F3', '#FF9800', '#4CAF50', '#F44336']
bar_colors_5 = ['#2196F3', '#FF9800', '#4CAF50', '#F44336', '#9C27B0']
pie_colors = ['#2196F3', '#FF9800', '#4CAF50', '#F44336', '#9C27B0', '#00BCD4', '#795548']

def make_bar(title, labels, values, filename, colors=None, ylabel='Percentage (%)'):
    if colors is None:
        colors = bar_colors_5[:len(labels)]
    fig, ax = plt.subplots(figsize=(8, 5))
    bars = ax.bar(labels, values, color=colors, edgecolor='white', linewidth=0.8, width=0.6)
    ax.set_title(title, fontsize=14, fontweight='bold', pad=15)
    ax.set_ylabel(ylabel, fontsize=11)
    ax.set_ylim(0, max(values) + 15)
    for bar, val in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 1.5, f'{val}%', ha='center', va='bottom', fontweight='bold', fontsize=10)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.tick_params(axis='x', rotation=0)
    plt.tight_layout()
    plt.savefig(filename, dpi=150, bbox_inches='tight')
    plt.close()

def make_pie(title, labels, values, filename, colors=None):
    if colors is None:
        colors = pie_colors[:len(labels)]
    fig, ax = plt.subplots(figsize=(7, 5))
    wedges, texts, autotexts = ax.pie(values, labels=None, autopct='%1.1f%%', colors=colors,
                                       startangle=90, pctdistance=0.75, textprops={'fontsize': 10})
    for t in autotexts:
        t.set_fontweight('bold')
    ax.legend(wedges, [f'{l} ({v}%)' for l, v in zip(labels, values)], loc='center left',
              bbox_to_anchor=(1, 0.5), fontsize=9)
    ax.set_title(title, fontsize=14, fontweight='bold', pad=15)
    plt.tight_layout()
    plt.savefig(filename, dpi=150, bbox_inches='tight')
    plt.close()

def make_grouped_bar(title, labels, data_dict, filename, colors=None):
    """data_dict: {'Category A': [vals], 'Category B': [vals]}"""
    x = np.arange(len(labels))
    n = len(data_dict)
    width = 0.7 / n
    if colors is None:
        colors = bar_colors_5[:n]
    fig, ax = plt.subplots(figsize=(10, 6))
    for i, (cat, vals) in enumerate(data_dict.items()):
        offset = (i - n/2 + 0.5) * width
        bars = ax.bar(x + offset, vals, width, label=cat, color=colors[i], edgecolor='white')
        for bar, val in zip(bars, vals):
            if val > 0:
                ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.5, f'{val}%', ha='center', va='bottom', fontsize=8)
    ax.set_title(title, fontsize=14, fontweight='bold', pad=15)
    ax.set_ylabel('Percentage (%)', fontsize=11)
    ax.set_xticks(x)
    ax.set_xticklabels(labels, fontsize=9)
    ax.legend(fontsize=9)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.set_ylim(0, max([max(v) for v in data_dict.values()]) + 10)
    plt.tight_layout()
    plt.savefig(filename, dpi=150, bbox_inches='tight')
    plt.close()

# ============ Generate all charts ============

# Demographics
make_bar('Age Group Distribution', ['17-18 yrs', '19-20 yrs', '21-22 yrs', '>22 yrs'], [8.2, 48.2, 41.8, 1.8], 'chart_01_age.png', ['#00BCD4', '#2196F3', '#FF9800', '#9C27B0'])
make_pie('Gender Distribution', ['Female', 'Male'], [58.2, 41.8], 'chart_02_gender.png', ['#E91E63', '#2196F3'])
make_bar('Semester Distribution', ['1st', '2nd', '3rd', '4th', '5th', '6th', '8th'], [7.3, 62.7, 18.2, 0.9, 3.6, 2.7, 4.5], 'chart_03_semester.png', pie_colors)
make_pie('Type of Residence', ['Home', 'Hostel', 'PG/Other'], [46.4, 35.5, 18.2], 'chart_04_residence.png')
make_bar('Average Daily Time on Digital Devices', ['<2 hrs', '2-4 hrs', '4-6 hrs', '>6 hrs'], [12.7, 29.1, 37.3, 20.9], 'chart_05_dailytime.png', ['#4CAF50', '#2196F3', '#FF9800', '#F44336'])
make_pie('Main Device Used for Studies', ['Smartphone', 'Laptop', 'Tablet', 'Desktop'], [60.0, 31.8, 5.5, 2.7], 'chart_06_device.png')
make_bar('Continuous Screen Time Without Break', ['<30 min', '30-60 min', '1-2 hrs', '>2 hrs'], [24.5, 40.9, 24.5, 10.0], 'chart_07_contscreen.png', ['#4CAF50', '#2196F3', '#FF9800', '#F44336'])
make_bar('Total Daily Screen Exposure', ['<3 hrs', '3-5 hrs', '5-8 hrs', '>8 hrs'], [23.6, 33.6, 25.5, 17.3], 'chart_08_totalexposure.png', ['#4CAF50', '#2196F3', '#FF9800', '#F44336'])

# Practice
make_pie('Use of Devices in Dark/Low Light', ['Yes', 'No'], [76.4, 23.6], 'chart_09_darklight.png', ['#F44336', '#4CAF50'])
make_pie('Use of Devices While Lying Down', ['Yes', 'No'], [82.7, 17.3], 'chart_10_lyingdown.png', ['#F44336', '#4CAF50'])
make_bar('Taking Breaks During Long Screen Use', ['Yes', 'Sometimes', 'No'], [64.5, 24.5, 10.9], 'chart_11_breaks.png', ['#4CAF50', '#FF9800', '#F44336'])
make_pie('Use of Blue Light Filter', ['Yes', 'No'], [62.7, 37.3], 'chart_12_bluelight.png', ['#4CAF50', '#F44336'])
make_bar('Following 20-20-20 Rule (Usage)', ['Never', 'Sometimes', 'Rarely', 'Always'], [35.5, 32.7, 26.4, 5.5], 'chart_13_2020_1.png', ['#F44336', '#FF9800', '#9C27B0', '#4CAF50'])

# Symptoms - grouped bar chart
symptom_labels = ['Eye\nStrain', 'Burning\nSensation', 'Dry\nEyes', 'Headache', 'Blurred\nVision', 'Eye\nRedness', 'Neck/\nShoulder\nPain', 'Difficulty\nFocusing', 'Sensitivity\nto Light', 'Excessive\nWatering']
symptom_yes = [60.0, 60.0, 54.5, 79.1, 51.8, 42.7, 78.2, 64.5, 58.2, 56.4]
symptom_never = [15.5, 40.0, 45.5, 20.9, 48.2, 57.3, 21.8, 35.5, 41.8, 43.6]

fig, ax = plt.subplots(figsize=(12, 6))
x = np.arange(len(symptom_labels))
width = 0.35
bars1 = ax.bar(x - width/2, symptom_yes, width, label='Experiencing (Sometimes+)', color='#F44336', edgecolor='white')
bars2 = ax.bar(x + width/2, symptom_never, width, label='Never Experiencing', color='#4CAF50', edgecolor='white')
ax.set_title('Prevalence of Digital Eye Strain Symptoms', fontsize=14, fontweight='bold', pad=15)
ax.set_ylabel('Percentage (%)', fontsize=11)
ax.set_xticks(x)
ax.set_xticklabels(symptom_labels, fontsize=8, ha='center')
ax.legend(fontsize=10)
ax.spines['top'].set_visible(False)
ax.spines['right'].set_visible(False)
ax.set_ylim(0, 100)
for bar in bars1:
    ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 1, f'{bar.get_height():.0f}%', ha='center', va='bottom', fontsize=8, fontweight='bold')
for bar in bars2:
    ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 1, f'{bar.get_height():.0f}%', ha='center', va='bottom', fontsize=8, fontweight='bold')
plt.tight_layout()
plt.savefig('chart_14_symptoms_grouped.png', dpi=150, bbox_inches='tight')
plt.close()

# Individual symptom bars
symptoms_data = [
    ('Eye Strain/Tired Eyes', ['Never','Rarely','Sometimes','Often','Always'], [15.5,24.5,43.6,13.6,2.7], ['#4CAF50','#00BCD4','#FF9800','#F44336','#9C27B0']),
    ('Burning Sensation', ['Never','Sometimes','Occasionally','Often','Always'], [40.0,36.4,18.2,2.7,2.7], ['#4CAF50','#FF9800','#00BCD4','#F44336','#9C27B0']),
    ('Dry Eyes', ['Never','Sometimes','Occasionally','Often','Always'], [45.5,29.1,17.3,6.4,1.8], ['#4CAF50','#FF9800','#00BCD4','#F44336','#9C27B0']),
    ('Headache After Screen Use', ['Never','Rarely','Sometimes','Often','Always'], [20.9,30.9,33.6,9.1,5.5], ['#4CAF50','#00BCD4','#FF9800','#F44336','#9C27B0']),
    ('Blurred Vision', ['Never','Sometimes','Occasionally','Often','Always'], [48.2,21.8,24.5,4.5,0.9], ['#4CAF50','#FF9800','#00BCD4','#F44336','#9C27B0']),
    ('Eye Redness', ['Never','Sometimes','Occasionally','Often','Always'], [57.3,18.2,20.0,3.6,0.9], ['#4CAF50','#FF9800','#00BCD4','#F44336','#9C27B0']),
    ('Neck/Shoulder Pain', ['Never','Rarely','Sometimes','Often','Always'], [21.8,30.0,31.8,9.1,7.3], ['#4CAF50','#00BCD4','#FF9800','#F44336','#9C27B0']),
    ('Difficulty Focusing', ['Never','Sometimes','Occasionally','Often','Always'], [35.5,31.8,15.5,11.8,5.5], ['#4CAF50','#FF9800','#00BCD4','#F44336','#9C27B0']),
    ('Sensitivity to Light', ['Never','Sometimes','Occasionally','Often','Always'], [41.8,29.1,20.0,3.6,5.5], ['#4CAF50','#FF9800','#00BCD4','#F44336','#9C27B0']),
    ('Excessive Watering', ['Never','Sometimes','Occasionally','Often','Always'], [43.6,25.5,20.0,7.3,3.6], ['#4CAF50','#FF9800','#00BCD4','#F44336','#9C27B0']),
]
for i, (name, labs, vals, cols) in enumerate(symptoms_data):
    make_bar(f'Frequency of {name}', labs, vals, f'chart_sym_{i+1}.png', cols)

# Preventive practices
make_bar('Following 20-20-20 Rule (Preventive)', ['No', 'Yes', 'Sometimes'], [55.5, 22.7, 21.8], 'chart_15_2020_2.png', ['#F44336', '#4CAF50', '#FF9800'])
make_pie('Use of Artificial Tears/Eye Drops', ['No', 'Yes', 'Occasionally'], [72.7, 15.5, 11.8], 'chart_16_artificial.png', ['#F44336', '#4CAF50', '#FF9800'])
make_bar('Conscious Reduction of Screen Brightness', ['Always', 'Sometimes', 'Never'], [49.1, 47.3, 3.6], 'chart_17_brightness.png', ['#4CAF50', '#FF9800', '#F44336'])
make_bar('Proper Screen Distance Adjustment', ['Yes', 'No', 'Sometimes'], [52.7, 30.9, 16.4], 'chart_18_distance.png', ['#4CAF50', '#F44336', '#FF9800'])
make_pie('Eye Relaxation Exercises', ['No', 'Yes'], [65.5, 34.5], 'chart_19_exercise.png', ['#F44336', '#4CAF50'])
make_bar('Consultation with Eye Specialist', ['No', 'Yes', 'Rarely'], [45.5, 38.2, 16.4], 'chart_20_specialist.png', ['#F44336', '#4CAF50', '#9C27B0'])

# Awareness
make_pie('Awareness of Digital Eye Strain', ['Yes', 'No'], [73.6, 26.4], 'chart_21_awareness.png', ['#4CAF50', '#F44336'])
make_pie('Knowledge: Screen Use Affects Eye Health', ['Yes', 'No'], [89.1, 10.9], 'chart_22_knowledge_health.png', ['#4CAF50', '#F44336'])
make_pie('Knowledge of Preventive Measures', ['Yes', 'No'], [69.1, 30.9], 'chart_23_knowledge_prev.png', ['#4CAF50', '#F44336'])
make_pie('Previous Eye Care Education', ['No', 'Yes'], [63.6, 35.5], 'chart_24_education.png', ['#F44336', '#4CAF50'])

# Summary comparison chart
summary_labels = ['Eye\nStrain', 'Burning\nSensation', 'Dry\nEyes', 'Headache', 'Blurred\nVision', 'Eye\nRedness', 'Neck/\nShoulder\nPain', 'Difficulty\nFocusing', 'Sensitivity\nto Light', 'Excessive\nWatering']
summary_vals = [60.0, 60.0, 54.5, 79.1, 51.8, 42.7, 78.2, 64.5, 58.2, 56.4]
fig, ax = plt.subplots(figsize=(10, 5))
sorted_pairs = sorted(zip(summary_vals, summary_labels), reverse=True)
vals_sorted = [v for v, l in sorted_pairs]
labels_sorted = [l for v, l in sorted_pairs]
gradient = plt.cm.RdYlGn_r(np.linspace(0.2, 0.9, len(vals_sorted)))
bars = ax.barh(labels_sorted, vals_sorted, color=gradient, edgecolor='white', height=0.6)
ax.set_title('DES Symptoms Ranked by Prevalence', fontsize=14, fontweight='bold', pad=15)
ax.set_xlabel('Percentage (%)', fontsize=11)
ax.spines['top'].set_visible(False)
ax.spines['right'].set_visible(False)
ax.set_xlim(0, 100)
for bar, val in zip(bars, vals_sorted):
    ax.text(bar.get_width() + 1, bar.get_y() + bar.get_height()/2, f'{val}%', va='center', fontweight='bold', fontsize=10)
plt.tight_layout()
plt.savefig('chart_25_summary_rank.png', dpi=150, bbox_inches='tight')
plt.close()

print("All charts generated successfully!")

# ============ Now build the Word document ============
doc = Document()
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

def add_hl(text, level=2):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in h.runs: r.font.color.rgb = RGBColor(0,0,0)
    return h

def add_p(text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, sp=Pt(6)):
    p = doc.add_paragraph()
    p.alignment = align; p.paragraph_format.space_after = sp
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(12)
    r.bold = bold; r.italic = italic
    return p

def add_tbl(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.size = Pt(10); r.font.name = 'Times New Roman'
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = str(val)
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs: r.font.size = Pt(10); r.font.name = 'Times New Roman'
    doc.add_paragraph()
    return t

def add_chart(img_path):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=Inches(5.8))
    doc.add_paragraph()

# ============ CHAPTER 4 TITLE ============
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("CHAPTER 4"); r.bold = True; r.font.size = Pt(16); r.font.name = 'Times New Roman'
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("DATA ANALYSIS AND INTERPRETATION"); r.bold = True; r.font.size = Pt(14); r.font.name = 'Times New Roman'
doc.add_paragraph()
add_p("This chapter presents the analysis and interpretation of data collected from 110 nursing students through a structured self-administered questionnaire. Data were collected from July 3 to July 22, 2026. Findings are organized under four headings: demographic profile, practice related to digital device usage, symptoms of digital eye strain, and preventive practices and awareness.")

# ============ 4.1 DEMOGRAPHICS ============
add_hl("4.1 Demographic Profile of the Participants")

# 4.1 Age
add_p("Table 4.1: Distribution of Participants by Age Group")
add_tbl(["Age Group","Frequency (n)","Percentage (%)"], [["17-18 years","9","8.2"],["19-20 years","53","48.2"],["21-22 years","46","41.8"],[">22 years","2","1.8"],["Total","110","100.0"]])
add_p("Figure 4.1: Age Group Distribution", bold=True)
add_chart('chart_01_age.png')
add_p("Interpretation: The majority of participants (48.2%) were in the 19-20 years age group, followed by 21-22 years (41.8%).", italic=True)

# 4.2 Gender
add_p("Table 4.2: Distribution of Participants by Gender")
add_tbl(["Gender","Frequency (n)","Percentage (%)"], [["Female","64","58.2"],["Male","46","41.8"],["Total","110","100.0"]])
add_p("Figure 4.2: Gender Distribution", bold=True)
add_chart('chart_02_gender.png')
add_p("Interpretation: 58.2% were female, consistent with nursing education demographics.", italic=True)

# 4.3 Semester
add_p("Table 4.3: Distribution of Participants by Semester")
add_tbl(["Semester","Frequency (n)","Percentage (%)"], [["1st","8","7.3"],["2nd","69","62.7"],["3rd","20","18.2"],["4th","1","0.9"],["5th","4","3.6"],["6th","3","2.7"],["8th","5","4.5"],["Total","110","100.0"]])
add_p("Figure 4.3: Semester Distribution", bold=True)
add_chart('chart_03_semester.png')
add_p("Interpretation: 62.7% were in the 2nd semester, the foundational year.", italic=True)

# 4.4 Residence
add_p("Table 4.4: Distribution of Participants by Type of Residence")
add_tbl(["Type of Residence","Frequency (n)","Percentage (%)"], [["Home","51","46.4"],["Hostel","39","35.5"],["PG/Other","20","18.2"],["Total","110","100.0"]])
add_p("Figure 4.4: Type of Residence", bold=True)
add_chart('chart_04_residence.png')
add_p("Interpretation: Varied living conditions may influence device usage patterns.", italic=True)

# 4.5 Daily Time
add_p("Table 4.5: Distribution of Participants by Average Daily Time Spent on Digital Devices")
add_tbl(["Average Daily Time","Frequency (n)","Percentage (%)"], [["<2 hours","14","12.7"],["2-4 hours","32","29.1"],["4-6 hours","41","37.3"],[">6 hours","23","20.9"],["Total","110","100.0"]])
add_p("Figure 4.5: Average Daily Time on Digital Devices", bold=True)
add_chart('chart_05_dailytime.png')
add_p("Interpretation: 87.3% spend 2+ hours daily on digital devices, at DES risk.", italic=True)

# 4.6 Device
add_p("Table 4.6: Distribution of Participants by Main Device Used for Studies")
add_tbl(["Main Device","Frequency (n)","Percentage (%)"], [["Smartphone","66","60.0"],["Laptop","35","31.8"],["Tablet","6","5.5"],["Desktop","3","2.7"],["Total","110","100.0"]])
add_p("Figure 4.6: Main Device Used for Studies", bold=True)
add_chart('chart_06_device.png')
add_p("Interpretation: Smartphones dominate (60.0%), posing higher DES risk due to smaller screens.", italic=True)

# 4.7 Continuous
add_p("Table 4.7: Distribution of Participants by Average Continuous Screen Time Without Break")
add_tbl(["Continuous Screen Time","Frequency (n)","Percentage (%)"], [["<30 min","27","24.5"],["30-60 min","45","40.9"],["1-2 hrs","27","24.5"],[">2 hrs","11","10.0"],["Total","110","100.0"]])
add_p("Figure 4.7: Continuous Screen Time Without Break", bold=True)
add_chart('chart_07_contscreen.png')
add_p("Interpretation: 34.5% use devices continuously for more than 1 hour without a break.", italic=True)

# 4.8 Total exposure
add_p("Table 4.8: Distribution of Participants by Total Daily Screen Exposure")
add_tbl(["Total Daily Screen Exposure","Frequency (n)","Percentage (%)"], [["<3 hours","26","23.6"],["3-5 hours","37","33.6"],["5-8 hours","28","25.5"],[">8 hours","19","17.3"],["Total","110","100.0"]])
add_p("Figure 4.8: Total Daily Screen Exposure", bold=True)
add_chart('chart_08_totalexposure.png')
add_p("Interpretation: 76.4% have total daily exposure exceeding 3 hours.", italic=True)

doc.add_page_break()

# ============ 4.2 PRACTICE ============
add_hl("4.2 Practice Related to Digital Device Usage")

add_p("Table 4.9: Use of Digital Devices in Dark/Low Light")
add_tbl(["Response","Frequency (n)","Percentage (%)"], [["Yes","84","76.4"],["No","26","23.6"],["Total","110","100.0"]])
add_p("Figure 4.9: Use of Devices in Dark/Low Light", bold=True)
add_chart('chart_09_darklight.png')
add_p("Interpretation: 76.4% used devices in dark environments, significantly increasing DES risk.", italic=True)

add_p("Table 4.10: Use of Devices While Lying Down")
add_tbl(["Response","Frequency (n)","Percentage (%)"], [["Yes","91","82.7"],["No","19","17.3"],["Total","110","100.0"]])
add_p("Figure 4.10: Use of Devices While Lying Down", bold=True)
add_chart('chart_10_lyingdown.png')
add_p("Interpretation: 82.7% used devices while lying down, leading to poor ergonomics.", italic=True)

add_p("Table 4.11: Taking Breaks During Long Screen Use")
add_tbl(["Response","Frequency (n)","Percentage (%)"], [["Yes","71","64.5"],["Sometimes","27","24.5"],["No","12","10.9"],["Total","110","100.0"]])
add_p("Figure 4.11: Taking Breaks During Screen Use", bold=True)
add_chart('chart_11_breaks.png')
add_p("Interpretation: 35.4% had inconsistent or no break-taking habits.", italic=True)

add_p("Table 4.12: Use of Blue Light Filter/Night Mode")
add_tbl(["Response","Frequency (n)","Percentage (%)"], [["Yes","69","62.7"],["No","41","37.3"],["Total","110","100.0"]])
add_p("Figure 4.12: Use of Blue Light Filter", bold=True)
add_chart('chart_12_bluelight.png')
add_p("Interpretation: 37.3% did not use blue light filtering technology.", italic=True)

add_p("Table 4.13: Following the 20-20-20 Rule While Using Digital Devices")
add_tbl(["Response","Frequency (n)","Percentage (%)"], [["Never","39","35.5"],["Sometimes","36","32.7"],["Rarely","29","26.4"],["Always","6","5.5"],["Total","110","100.0"]])
add_p("Figure 4.13: Following 20-20-20 Rule", bold=True)
add_chart('chart_13_2020_1.png')
add_p("Interpretation: Only 5.5% always followed the 20-20-20 rule — a critical finding.", italic=True)

doc.add_page_break()

# ============ 4.3 SYMPTOMS ============
add_hl("4.3 Symptoms of Digital Eye Strain")

add_p("Figure 4.14: Overview of DES Symptoms Prevalence (Experiencing vs Never)", bold=True)
add_chart('chart_14_symptoms_grouped.png')
add_p("Interpretation: The grouped bar chart clearly shows that headache (79.1%) and neck/shoulder pain (78.2%) are the most prevalent DES symptoms, while eye redness (42.7% never) is the least commonly experienced symptom.", italic=True)

add_p("Figure 4.15: DES Symptoms Ranked by Prevalence", bold=True)
add_chart('chart_25_summary_rank.png')
add_p("Interpretation: This ranked visualization confirms headache (79.1%) and neck/shoulder pain (78.2%) as the top two symptoms, followed by difficulty focusing (64.5%) and eye strain (60%).", italic=True)

doc.add_page_break()

# Individual symptom tables and charts
symptoms = [
    ("Table 4.14", "Frequency of Eye Strain/Tired Eyes", [["Never","17","15.5"],["Rarely","27","24.5"],["Sometimes","48","43.6"],["Often","15","13.6"],["Always","3","2.7"],["Total","110","100.0"]], "60% experienced eye strain at least sometimes.", 'chart_sym_1.png'),
    ("Table 4.15", "Frequency of Burning Sensation in Eyes", [["Never","44","40.0"],["Sometimes","40","36.4"],["Occasionally","20","18.2"],["Often","3","2.7"],["Always","3","2.7"],["Total","110","100.0"]], "60% reported burning sensation at least sometimes.", 'chart_sym_2.png'),
    ("Table 4.16", "Frequency of Dry Eyes", [["Never","50","45.5"],["Sometimes","32","29.1"],["Occasionally","19","17.3"],["Often","7","6.4"],["Always","2","1.8"],["Total","110","100.0"]], "54.5% experienced dry eyes due to reduced blinking.", 'chart_sym_3.png'),
    ("Table 4.17", "Frequency of Headache After Screen Use", [["Never","23","20.9"],["Rarely","34","30.9"],["Sometimes","37","33.6"],["Often","10","9.1"],["Always","6","5.5"],["Total","110","100.0"]], "79.1% reported headaches — the most prevalent symptom.", 'chart_sym_4.png'),
    ("Table 4.18", "Frequency of Blurred Vision", [["Never","53","48.2"],["Sometimes","24","21.8"],["Occasionally","27","24.5"],["Often","5","4.5"],["Always","1","0.9"],["Total","110","100.0"]], "51.8% experienced blurred vision.", 'chart_sym_5.png'),
    ("Table 4.19", "Frequency of Eye Redness", [["Never","63","57.3"],["Sometimes","20","18.2"],["Occasionally","22","20.0"],["Often","4","3.6"],["Always","1","0.9"],["Total","110","100.0"]], "42.7% reported eye redness.", 'chart_sym_6.png'),
    ("Table 4.20", "Frequency of Neck or Shoulder Pain", [["Never","24","21.8"],["Rarely","33","30.0"],["Sometimes","35","31.8"],["Often","10","9.1"],["Always","8","7.3"],["Total","110","100.0"]], "78.2% experienced neck/shoulder pain, linked to poor posture.", 'chart_sym_7.png'),
    ("Table 4.21", "Frequency of Difficulty Focusing", [["Never","39","35.5"],["Sometimes","35","31.8"],["Occasionally","17","15.5"],["Often","13","11.8"],["Always","6","5.5"],["Total","110","100.0"]], "64.5% reported difficulty focusing, impacting academics.", 'chart_sym_8.png'),
    ("Table 4.22", "Frequency of Sensitivity to Light", [["Never","46","41.8"],["Sometimes","32","29.1"],["Occasionally","22","20.0"],["Often","4","3.6"],["Always","6","5.5"],["Total","110","100.0"]], "58.2% experienced light sensitivity.", 'chart_sym_9.png'),
    ("Table 4.23", "Frequency of Excessive Watering/Tearing", [["Never","48","43.6"],["Sometimes","28","25.5"],["Occasionally","22","20.0"],["Often","8","7.3"],["Always","4","3.6"],["Total","110","100.0"]], "56.4% experienced excessive tearing.", 'chart_sym_10.png'),
]

for tbl_no, tbl_name, data, interp, chart in symptoms:
    add_p(f"{tbl_no}: {tbl_name}")
    add_tbl(["Response","Frequency (n)","Percentage (%)"], data)
    add_p(f"Figure: {tbl_name} - Visual Representation", bold=True)
    add_chart(chart)
    add_p(f"Interpretation: {interp}", italic=True)

doc.add_page_break()

# ============ 4.4 PREVENTIVE PRACTICES ============
add_hl("4.4 Preventive Practices and Awareness")

preventive_data = [
    ("Table 4.24", "Following the 20-20-20 Rule (Preventive Practice)", [["No","61","55.5"],["Yes","25","22.7"],["Sometimes","24","21.8"],["Total","110","100.0"]], "55.5% did not follow the 20-20-20 rule.", 'chart_15_2020_2.png'),
    ("Table 4.25", "Use of Artificial Tears or Eye Drops", [["No","80","72.7"],["Yes","17","15.5"],["Occasionally","13","11.8"],["Total","110","100.0"]], "72.7% did not use artificial tears despite 54.5% having dry eyes.", 'chart_16_artificial.png'),
    ("Table 4.26", "Conscious Reduction of Screen Brightness", [["Always","54","49.1"],["Sometimes","52","47.3"],["Never","4","3.6"],["Total","110","100.0"]], "Positive finding: 96.4% adjusted brightness at least sometimes.", 'chart_17_brightness.png'),
    ("Table 4.27", "Proper Screen Distance Adjustment", [["Yes","58","52.7"],["No","34","30.9"],["Sometimes","18","16.4"],["Total","110","100.0"]], "47.3% did not consistently adjust screen distance.", 'chart_18_distance.png'),
    ("Table 4.28", "Performance of Eye Relaxation Exercises", [["No","72","65.5"],["Yes","38","34.5"],["Total","110","100.0"]], "65.5% did not perform eye relaxation exercises.", 'chart_19_exercise.png'),
    ("Table 4.29", "Consultation with Eye Specialist", [["No","50","45.5"],["Yes","42","38.2"],["Rarely","18","16.4"],["Total","110","100.0"]], "45.5% did not consult an eye specialist.", 'chart_20_specialist.png'),
    ("Table 4.30", "Awareness of Digital Eye Strain", [["Yes","81","73.6"],["No","29","26.4"],["Total","110","100.0"]], "73.6% were aware of DES.", 'chart_21_awareness.png'),
    ("Table 4.31", "Knowledge: Prolonged Screen Use Affects Eye Health", [["Yes","98","89.1"],["No","12","10.9"],["Total","110","100.0"]], "89.1% knew about effects.", 'chart_22_knowledge_health.png'),
    ("Table 4.32", "Knowledge of Preventive Measures", [["Yes","76","69.1"],["No","34","30.9"],["Total","110","100.0"]], "30.9% did not know about preventive measures.", 'chart_23_knowledge_prev.png'),
    ("Table 4.33", "Previous Education/Training on Eye Care", [["No","70","63.6"],["Yes","39","35.5"],["Other","1","0.9"],["Total","110","100.0"]], "63.6% never received eye care education.", 'chart_24_education.png'),
]

for tbl_no, tbl_name, data, interp, chart in preventive_data:
    add_p(f"{tbl_no}: {tbl_name}")
    add_tbl(["Response","Frequency (n)","Percentage (%)"], data)
    add_p(f"Figure: {tbl_name} - Visual Representation", bold=True)
    add_chart(chart)
    add_p(f"Interpretation: {interp}", italic=True)

doc.save("Chapter_4_Data_Analysis_with_Charts.docx")
print("Chapter 4 with charts saved successfully!")
