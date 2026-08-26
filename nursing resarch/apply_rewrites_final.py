import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.oxml.ns import qn

doc = Document('Document3_WithTables.docx')

print(f"Starting: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")

# Define ALL rewrites as a dict mapping paragraph text start to new text
rewrites = {}

# CHAPTER 1: INTRODUCTION
rewrites['Digital technology has become a central part'] = """Today's nursing students spend an extraordinary amount of time staring at screens. Between online lectures, digital assignments, clinical documentation, and the occasional late-night scroll through social media, the hours add up fast. And while smartphones and laptops have made studying more convenient than ever, they have also introduced a problem that many students overlook until it becomes impossible to ignore: digital eye strain."""

rewrites['Digital Eye Strain, also referred to as Computer Vision Syndrome'] = """Think of digital eye strain as your eyes' way of telling you they have had enough. Medically, it falls under the umbrella of computer vision syndrome, and it shows up as a bundle of complaints: eyes that feel heavy or tired, a burning or gritty sensation, redness, watering, headaches that seem to start right behind the eyes, and sometimes even blurred vision that comes and goes. None of these are dangerous in the short term, but they are incredibly disruptive — especially when you are trying to focus on a lecture or study for an exam. The root cause is fairly straightforward. When we stare at screens, we blink far less often than normal. Our blink rate can drop from roughly 15 times a minute to as low as 5 or 6. That reduced blinking allows the tear film on the surface of the eye to thin out, which causes dryness and irritation."""

rewrites['Research suggests that roughly 50% to 90%'] = """What makes digital eye strain tricky is how normalised it has become. Most students treat it as an unavoidable side effect of modern education rather than something worth addressing. A student who gets a headache after three hours of screen time does not typically think, "I need to change my habits" — they reach for a painkiller and keep scrolling. For nursing students specifically, the stakes are a bit higher than for the general population. When you are sitting in a classroom, a headache is annoying. But when you are in a clinical setting — monitoring vitals, reading medication charts, documenting patient observations — even mild visual fatigue can compromise your performance and attention."""

rewrites['Self-care plays an important role in both preventing'] = """Self-care, in this context, is not about grand gestures or expensive eye treatments. It is about small, consistent habits that protect the eyes during everyday screen use. The 20-20-20 rule, which we mentioned earlier, is probably the most widely cited recommendation. Beyond that, keeping your screen at roughly arm's length, ensuring the room is adequately lit (avoiding screen use in total darkness is a big one), and consciously making an effort to blink can all reduce symptom severity. The catch is that awareness does not equal action. Several studies have found that students can correctly identify preventive strategies when asked about them on a questionnaire, yet their actual behaviour tells a very different story."""

rewrites['Since nursing students are training to become healthcare professionals'] = """There is a broader point here that goes beyond individual comfort. Nursing education is built on the idea that healthcare professionals should model the behaviours they recommend to patients. If a nurse advises a patient to take regular breaks from screens, limit night-time phone use, and maintain proper posture while working — but does none of those things personally — something does not add up. Teaching students to look after their own visual health is therefore not just a wellness issue; it is a professional development issue."""

rewrites['Technology has reshaped nearly every aspect of life'] = """It is hard to overstate how much digital technology has reshaped education. Twenty years ago, nursing students relied on printed textbooks, handwritten notes, and physical library visits. Today, entire courses are delivered online, textbooks exist as PDFs and e-books, assignments are submitted through learning management systems, and clinical guidelines are accessible through smartphone apps. The transformation happened gradually at first, then all at once when the pandemic forced institutions worldwide to shift to remote learning almost overnight."""

rewrites['Among its many benefits are faster communication'] = """The American Optometric Association defines digital eye strain as the collection of visual symptoms that arise from prolonged use of digital devices — desktops, laptops, smartphones, tablets, and similar screens. It is worth noting that the term does not refer to a single clinical diagnosis. Instead, it captures a range of overlapping discomforts that share a common trigger: extended, uninterrupted screen time. What makes digital eye strain particularly common is that the conditions under which most people use screens are far from ideal."""

rewrites['The Digital Eye Strain Report of 2016'] = """A large-scale survey conducted in the United States in 2016 gathered responses from over 10,000 adults and found that roughly 65% reported symptoms consistent with digital eye strain. Women were somewhat more likely to report symptoms than men — 69% compared to 60% — though researchers were cautious about attributing this entirely to biological differences. Several factors contribute to the development of these symptoms. Reduced contrast between text and background, screen glare, viewing screens at awkward angles, poor posture, and insufficient ambient light all increase the visual workload placed on the eyes."""

rewrites['While digital devices have undeniably improved work efficiency'] = """There is little doubt that the COVID-19 pandemic turned digital eye strain from a niche complaint into a widespread issue. When offices, universities, and schools shut down, millions of people suddenly found themselves spending the entire day in front of screens. The shift was abrupt, and for many, the habits formed during lockdowns persisted long after restrictions were lifted. Even now, screen dependence has not meaningfully decreased. Digital tools have simply become too embedded in academic workflows to abandon."""

rewrites['The widespread adoption of smartphones and portable devices'] = """For nursing students, the digital workload is particularly intense. Their academic responsibilities span a wide range of screen-dependent activities: attending online lectures, reading electronic textbooks, preparing presentations, documenting clinical observations, searching research databases, and communicating with faculty. What compounds the issue is that these activities are not evenly distributed across the day. Students often have clusters of screen-intensive tasks, particularly during assignment submission weeks or exam preparation periods."""

rewrites['The severity of digital eye strain depends not only'] = """It is worth emphasising that digital eye strain is not just about how long you use a screen. How you use it matters enormously. A student who uses a laptop for three hours with the screen positioned at eye level, in a well-lit room, taking short breaks every thirty minutes, will generally fare much better than a student who uses a phone for ninety minutes in bed with the lights off. The "how" also includes factors that students rarely think about: whether they are blinking regularly, whether their prescription glasses are up to date, whether they are tilting their head at an awkward angle to avoid screen glare."""

rewrites['Awareness of preventive measures is another key factor'] = """Here is what makes this topic genuinely interesting from a research perspective: most students are not ignorant about digital eye strain. They have heard of it. They can name at least a couple of preventive measures. But there is a persistent, well-documented gap between what students know and what they actually do. A student might correctly identify the 20-20-20 rule on a questionnaire and then spend four consecutive hours on a laptop without looking up once."""

rewrites['Digital eye strain should also be viewed as a multidisciplinary concern'] = """From a broader perspective, digital eye strain sits at the intersection of several disciplines: ophthalmology, ergonomics, behavioural science, and health education. Nurses and nursing educators are uniquely positioned to address it, not only because they understand the health implications, but because they have direct access to the affected population."""

rewrites['In nursing education, visual comfort is especially important'] = """Nursing education, perhaps more than most other disciplines, demands sustained visual attention. Students spend hours reading dense academic texts, studying anatomical diagrams, observing clinical procedures, interpreting patient records, and writing care plans. When visual fatigue sets in, concentration suffers, errors increase, and the overall learning experience deteriorates."""

rewrites['The assessment of digital eye strain among nursing students'] = """This study was conducted at Sharda University, Greater Noida, where nursing students rely heavily on digital devices as part of their academic routine. By surveying students about their screen habits, symptoms, and self-care practices, the research aimed to generate a clear picture of how prevalent digital eye strain is in this specific population."""

rewrites['The present study is intended to generate evidence'] = """Ultimately, the goal was straightforward: find out how common digital eye strain is among nursing students at this institution, understand what self-care habits they follow, and use that information to support better awareness and healthier practices."""

rewrites['In conclusion, digital technology is now an essential component'] = """To sum up the introduction: digital technology is here to stay in nursing education, and screen exposure is not going to decrease any time soon. What can change is how students manage that exposure. Digital eye strain is largely preventable, and the preventive strategies are neither complicated nor expensive. What is needed is consistent awareness, practical education, and a shift from passive knowledge to active self-care behaviour."""

# Background of the Study
rewrites['The rapid advancement of digital technology has transformed'] = """The classrooms and libraries at Sharda University look quite different today than they did a decade ago. Nursing students carry laptops instead of heavy textbooks, submit assignments through online portals rather than handing in paper copies, and attend lectures that are simultaneously livestreamed for students who cannot be physically present. Digital devices have become so woven into the fabric of nursing education that it is almost impossible to imagine the programme functioning without them. But this reliance on technology, while undeniably efficient, comes with consequences that rarely get discussed in academic planning meetings."""

rewrites['Digital Eye Strain (DES), also referred to as Computer Vision Syndrome, is a group'] = """Digital eye strain is not a single symptom but rather a collection of related complaints that tend to appear together after prolonged screen use. Students describe it in different ways: "my eyes feel tired," "they burn after a while," "I get headaches," "things look blurry when I stop looking at the screen." The severity of these symptoms varies from person to person and depends on several factors: how long the screen is used, the type of device, ambient lighting, screen settings, posture, and whether the individual takes adequate breaks."""

rewrites['Nursing students represent an important population for studying'] = """Nursing students are an especially relevant group to study for several reasons. Their academic workload inherently demands heavy screen use — online modules, evidence-based practice research, clinical simulation exercises, digital documentation, and communication with faculty. As future healthcare professionals, nursing students are expected to understand and promote health literacy among their patients. Being knowledgeable about digital eye strain and its prevention aligns with their professional development objectives."""

rewrites['Self-care practices play an important role in preventing or reducing'] = """Self-care in the context of digital eye strain revolves around a handful of well-established practices: taking regular breaks, following the 20-20-20 rule, maintaining appropriate screen distance, ensuring adequate ambient lighting, blinking deliberately, and adjusting screen brightness to match the environment. The problem, as many studies have pointed out, is not a lack of awareness but a lack of consistent implementation."""

rewrites['The prevalence of digital eye strain and the self-care practices adopted'] = """Several demographic and behavioural factors might influence how severely a student experiences digital eye strain. Age, gender, academic year, daily screen duration, type of device used, and living arrangements could all play a role. By collecting data on these variables alongside symptom reports, this study aimed to identify which groups of students are most affected and which habits are most closely linked to symptom severity."""

rewrites['Therefore, the present study, entitled'] = """The study, titled "A Descriptive Study to Assess the Prevalence of Digital Eye Strain and Associated Self-Care Practices Among Nursing Students," was designed to provide a snapshot of the current situation. It does not claim to establish causal relationships — that would require a different research design. Instead, it maps the territory."""

# Need of the Study
rewrites['Digital devices have become an integral component of nursing education. Nursing students spend considerable'] = """The simple truth is that nursing students are using screens more than ever, and no one is systematically monitoring the consequences. Universities have invested heavily in digital infrastructure — learning management systems, online examination platforms, digital libraries — but have given comparatively little attention to the physical side effects of this digital transformation. Digital eye strain is not life-threatening, but it is life-affecting."""

rewrites['Digital eye strain is particularly relevant among students because its symptoms'] = """Consider the daily routine of a typical nursing student at this university. Morning lectures might involve projected slides and digital note-taking. Afternoon sessions could include online modules or simulation exercises. Evening study hours are spent on research, assignments, and exam preparation — almost entirely screen-based activities. By the end of such a day, many students report eye discomfort, but few connect it to their screen habits or take meaningful steps to address it."""

# Apply rewrites
count = 0
for para in doc.paragraphs:
    text = para.text.strip()
    for key, new_text in rewrites.items():
        if text.startswith(key):
            # Preserve formatting from first run
            if para.runs:
                first_run = para.runs[0]
                # Clear all runs
                for run in para.runs:
                    run.text = ""
                first_run.text = new_text
            else:
                para.text = new_text
            count += 1
            break

print(f"Applied {count} text rewrites")

# Save
output_path = 'Document3_Final_Reviewed.docx'
doc.save(output_path)

# Verify
print("\nVerifying...")
verify_doc = Document(output_path)
print(f"Paragraphs: {len(verify_doc.paragraphs)}")
print(f"Tables: {len(verify_doc.tables)}")

image_count = sum(1 for rel in verify_doc.part.rels.values() if 'image' in rel.reltype)
print(f"Images: {image_count}")

total_words = sum(len(p.text.split()) for p in verify_doc.paragraphs)
for table in verify_doc.tables:
    for row in table.rows:
        for cell in row.cells:
            total_words += len(cell.text.split())
print(f"Total words: {total_words}")

print(f"\nSaved to: {output_path}")
print("Done!")
