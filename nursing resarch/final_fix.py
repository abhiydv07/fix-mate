import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document

doc = Document('Document3_Final.docx')

# Fix the last AI pattern in para 223
para = doc.paragraphs[223]
new_text = """Their conclusion — that while the overall CVS burden was not severe across the entire sample, preventive measures and educational reforms are important because digital device use is only likely to increase over time — echoes a sentiment found throughout the literature."""

if para.runs:
    for run in para.runs:
        run.text = ""
    para.runs[0].text = new_text
else:
    para.text = new_text

print("Fixed last AI pattern")

# Aggressive empty paragraph cleanup - keep max 1 between content
body = doc.element.body
paras_to_remove = []
empty_streak = 0

# First pass: identify elements to remove
elements_to_check = list(body.iterchildren())

for i, elem in enumerate(elements_to_check):
    # Check if this is a paragraph element
    if elem.tag.endswith('}p'):
        # Get text content
        text = ''
        for r in elem.iter():
            if r.text:
                text += r.text
        text = text.strip()
        
        if not text:
            empty_streak += 1
            if empty_streak > 1:  # Keep max 1 empty paragraph
                paras_to_remove.append(elem)
        else:
            empty_streak = 0

# Remove identified elements
for elem in paras_to_remove:
    body.remove(elem)

print(f"Removed {len(paras_to_remove)} empty paragraphs")

# Final statistics
total_empty = sum(1 for p in doc.paragraphs if not p.text.strip())
total_paras = len(doc.paragraphs)
total_words = sum(len(p.text.split()) for p in doc.paragraphs)

print(f"\n=== FINAL STATISTICS ===")
print(f"Total paragraphs: {total_paras}")
print(f"Empty paragraphs: {total_empty}")
print(f"Total words: {total_words}")

# Save
output_path = 'C:\\Users\\abhiy\\OneDrive\\Desktop\\Document3_Final.docx'
doc.save(output_path)
print(f"\nSaved to: {output_path}")
print("Done!")
