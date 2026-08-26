import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.oxml.ns import qn
from lxml import etree
import copy

print("Loading original document...")
original = Document('Document3.docx')

print(f"Original: {len(original.paragraphs)} paragraphs, {len(original.tables)} tables")

# We'll work with the XML directly to preserve everything
# First, let's identify all body children (paragraphs, tables, etc.)
body = original.element.body
children = list(body)

print(f"Total body children: {len(children)}")

# Count by type
para_count = sum(1 for c in children if c.tag.endswith('}p'))
table_count = sum(1 for c in children if c.tag.endswith('}tbl'))
other_count = len(children) - para_count - table_count

print(f"Paragraphs: {para_count}")
print(f"Tables: {table_count}")
print(f"Other: {other_count}")

# Now remove consecutive empty paragraphs (keep max 1)
nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

def get_text(element):
    """Get all text from an XML element"""
    text = ''
    for t in element.iter(qn('w:t')):
        if t.text:
            text += t.text
    return text.strip()

# Find empty paragraphs to remove (consecutive ones)
elements_to_remove = []
empty_streak = 0

for child in children:
    if child.tag.endswith('}p'):  # Is a paragraph
        text = get_text(child)
        if not text:
            empty_streak += 1
            if empty_streak > 1:  # Keep max 1 empty paragraph
                elements_to_remove.append(child)
        else:
            empty_streak = 0
    else:
        # Not a paragraph (e.g., table) - reset streak
        empty_streak = 0

print(f"\nRemoving {len(elements_to_remove)} extra empty paragraphs...")

# Remove identified elements
for elem in elements_to_remove:
    body.remove(elem)

# Count after cleanup
children_after = list(body)
para_count_after = sum(1 for c in children_after if c.tag.endswith('}p'))
table_count_after = sum(1 for c in children_after if c.tag.endswith('}tbl'))

print(f"\nAfter cleanup:")
print(f"  Total children: {len(children_after)}")
print(f"  Paragraphs: {para_count_after}")
print(f"  Tables: {table_count_after}")

# Now apply text rewrites to specific paragraphs
print("\nApplying text rewrites...")

# Rebuild paragraph index (only non-empty paragraphs)
para_index = 0
rewrite_map = {}

# Define rewrites for key paragraphs
rewrites = {
    3: """Today's nursing students spend an extraordinary amount of time staring at screens. Between online lectures, digital assignments, clinical documentation, and the occasional late-night scroll through social media, the hours add up fast. And while smartphones and laptops have made studying more convenient than ever, they have also introduced a problem that many students overlook until it becomes impossible to ignore: digital eye strain.""",
    
    5: """Think of digital eye strain as your eyes' way of telling you they have had enough. Medically, it falls under the umbrella of computer vision syndrome, and it shows up as a bundle of complaints: eyes that feel heavy or tired, a burning or gritty sensation, redness, watering, headaches that seem to start right behind the eyes, and sometimes even blurred vision that comes and go""",
}

# Apply rewrites to the document
applied = 0
for child in list(body):
    if child.tag.endswith('}p'):
        text = get_text(child)
        if text:
            # Find matching rewrite
            for idx, new_text in rewrites.items():
                # Check if first few words match
                if text[:50] in new_text[:50] or new_text[:50] in text[:50]:
                    # Find the w:t elements and update
                    for t in child.iter(qn('w:t')):
                        if t.text:
                            t.text = new_text
                            applied += 1
                            break
                    break

print(f"Applied {applied} text rewrites")

# Save the document
output_path = 'Document3_Clean.docx'
original.save(output_path)

# Verify
print("\nVerifying...")
verify_doc = Document(output_path)
print(f"Final document: {len(verify_doc.paragraphs)} paragraphs, {len(verify_doc.tables)} tables")

# Count images
image_count = 0
for rel in verify_doc.part.rels.values():
    if 'image' in rel.reltype:
        image_count += 1
print(f"Images: {image_count}")

# Word count
total_words = sum(len(p.text.split()) for p in verify_doc.paragraphs)
for table in verify_doc.tables:
    for row in table.rows:
        for cell in row.cells:
            total_words += len(cell.text.split())
print(f"Total words (including tables): {total_words}")

print(f"\nSaved to: {output_path}")
print("Done!")
