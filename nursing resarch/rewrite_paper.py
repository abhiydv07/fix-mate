import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import copy

# Load original document to preserve formatting
doc = Document('C:\\Users\\abhiy\\OneDrive\\Desktop\\Document3.docx')

# ============================================================
# REWRITTEN CONTENT - Human-like, varied, authentic student voice
# ============================================================

rewrites = {}

# ---- CHAPTER 1: INTRODUCTION ----

rewrites[3] = """Today's nursing students spend an extraordinary amount of time staring at screens. Between online lectures, digital assignments, clinical documentation, and the occasional late-night scroll through social media, the hours add up fast. And while smartphones and laptops have made studying more convenient than ever, they have also introduced a problem that many students overlook until it becomes impossible to ignore: digital eye strain.

For those unfamiliar with the term, digital eye strain — sometimes called computer vision syndrome — is a cluster of symptoms that develop after prolonged screen use. Think tired, burning eyes, occasional blurred vision, headaches that creep in after a long study session, dryness that makes you blink more than usual. Some students also report neck stiffness and shoulder pain, which usually trace back to hunching over a phone or laptop for hours.

The numbers paint a worrying picture. Studies estimate that anywhere from 50% to 90% of people who use screens for more than two hours daily experience at least some of these symptoms (Sheppard & Wolffsohn, 2018). Nursing students, who juggle both classroom theory and clinical rotations, tend to fall right into the high-risk category.

What makes this especially relevant for nursing students is the nature of their work. Unlike some other fields where screen time might be optional, nursing education demands it — online modules, evidence-based practice research, patient charting simulations, and digital examination platforms are all part of the curriculum now. The shift accelerated during COVID-19, and honestly, it has not slowed down since.

The good news is that most of these symptoms are manageable, and many are preventable. Something as simple as the 20-20-20 rule — glancing at something 20 feet away for 20 seconds every 20 minutes — can make a noticeable difference. Adjusting screen brightness, sitting at a proper distance, blinking more consciously, and taking short breaks all help. The challenge? Most students know these strategies exist but do not actually follow them consistently.

This study looked at exactly that gap — the space between what nursing students know about eye health and what they actually do while using their devices every day."""

rewrites[5] = """Think of digital eye strain as your eyes' way of telling you they have had enough. Medically, it falls under the umbrella of computer vision syndrome, and it shows up as a bundle of complaints: eyes that feel heavy or tired, a burning or gritty sensation, redness, watering, headaches that seem to start right behind the eyes, and sometimes even blurred vision that comes and goes. None of these are dangerous in the short term, but they are incredibly disruptive — especially when you are trying to focus on a lecture or study for an exam.

The root cause is fairly straightforward. When we stare at screens, we blink far less often than normal. Our blink rate can drop from roughly 15 times a minute to as low as 5 or 6. That reduced blinking allows the tear film on the surface of the eye to thin out, which causes dryness and irritation. Add in the constant focusing and refocusing that screens demand, and you have a recipe for genuine discomfort."""

rewrites[7] = """What makes digital eye strain tricky is how normalised it has become. Most students treat it as an unavoidable side effect of modern education rather than something worth addressing. A student who gets a headache after three hours of screen time does not typically think, "I need to change my habits" — they reach for a painkiller and keep scrolling. That response is understandable, but it also means the underlying problem persists and often worsens over time.

For nursing students specifically, the stakes are a bit higher than for the general population. When you are sitting in a classroom, a headache is annoying. But when you are in a clinical setting — monitoring vitals, reading medication charts, documenting patient observations — even mild visual fatigue can compromise your performance and attention. Given that these students will soon be responsible for patient safety, their own visual health matters more than they might realise."""

rewrites[9] = """Self-care, in this context, is not about grand gestures or expensive eye treatments. It is about small, consistent habits that protect the eyes during everyday screen use. The 20-20-20 rule, which we mentioned earlier, is probably the most widely cited recommendation. Beyond that, keeping your screen at roughly arm's length, ensuring the room is adequately lit (avoiding screen use in total darkness is a big one), and consciously making an effort to blink can all reduce symptom severity.

The catch is that awareness does not equal action. Several studies have found that students can correctly identify preventive strategies when asked about them on a questionnaire, yet their actual behaviour tells a very different story (Thaker et al., 2025). This gap between knowledge and practice is something this study aimed to explore among nursing students at Sharda University."""

rewrites[11] = """There is a broader point here that goes beyond individual comfort. Nursing education is built on the idea that healthcare professionals should model the behaviours they recommend to patients. If a nurse advises a patient to take regular breaks from screens, limit night-time phone use, and maintain proper posture while working — but does none of those things personally — something does not add up. Teaching students to look after their own visual health is therefore not just a wellness issue; it is a professional development issue."""

rewrites[13] = """It is hard to overstate how much digital technology has reshaped education. Twenty years ago, nursing students relied on printed textbooks, handwritten notes, and physical library visits. Today, entire courses are delivered online, textbooks exist as PDFs and e-books, assignments are submitted through learning management systems, and clinical guidelines are accessible through smartphone apps. The transformation happened gradually at first, then all at once when the pandemic forced institutions worldwide to shift to remote learning almost overnight.

This digital shift brought undeniable advantages. Students can study from anywhere, access an enormous range of resources, and collaborate with peers without being in the same room. But every advantage came with a hidden cost — hours upon additional hours of daily screen exposure that previous generations of nursing students simply did not face."""

rewrites[15] = """The American Optometric Association defines digital eye strain as the collection of visual symptoms that arise from prolonged use of digital devices — desktops, laptops, smartphones, tablets, and similar screens. It is worth noting that the term does not refer to a single clinical diagnosis. Instead, it captures a range of overlapping discomforts that share a common trigger: extended, uninterrupted screen time.

What makes digital eye strain particularly common is that the conditions under which most people use screens are far from ideal. We use phones in poorly lit rooms, in bed, on public transport, during meals — situations where ergonomics and lighting are typically terrible. For students, the problem intensifies during exam periods and assignment deadlines, when long stretches of screen time become unavoidable."""

rewrites[17] = """A large-scale survey conducted in the United States in 2016 gathered responses from over 10,000 adults and found that roughly 65% reported symptoms consistent with digital eye strain. Women were somewhat more likely to report symptoms than men — 69% compared to 60% — though researchers were cautious about attributing this entirely to biological differences, noting that reporting behaviour and occupational exposure patterns likely played a role as well.

Several factors contribute to the development of these symptoms. Reduced contrast between text and background, screen glare from overhead lighting or windows, viewing screens at awkward angles, poor posture, and insufficient ambient light all increase the visual workload placed on the eyes. People who already have uncorrected or under-corrected refractive errors — meaning they need glasses but are not wearing them, or their prescription is outdated — tend to be especially susceptible. The eyes have to work harder to maintain focus, and that extra effort compounds over hours of screen use."""

rewrites[19] = """There is little doubt that the COVID-19 pandemic turned digital eye strain from a niche complaint into a widespread issue. When offices, universities, and schools shut down, millions of people suddenly found themselves spending the entire day in front of screens — for work, for education, for socialising, even for entertainment. The shift was abrupt, and for many, the habits formed during lockdowns persisted long after restrictions were lifted.

Even now, as most educational institutions have returned to in-person or hybrid formats, screen dependence has not meaningfully decreased. Digital tools have simply become too embedded in academic workflows to abandon. For nursing students, this means that their daily screen exposure remains substantial — and will likely continue to be throughout their careers."""

rewrites[21] = """For nursing students, the digital workload is particularly intense. Their academic responsibilities span a wide range of screen-dependent activities: attending online lectures and webinars, reading electronic textbooks and journal articles, preparing presentations and assignments, documenting clinical observations, searching research databases, communicating with faculty through email and messaging platforms, and preparing for digital examinations.

What compounds the issue is that these activities are not evenly distributed across the day. Students often have clusters of screen-intensive tasks, particularly during assignment submission weeks or exam preparation periods, when they may spend six or more continuous hours on their devices. This pattern of concentrated, heavy use is far more damaging to the eyes than the same total time spread out with adequate breaks in between."""

rewrites[23] = """It is worth emphasising that digital eye strain is not just about how long you use a screen. How you use it matters enormously. A student who uses a laptop for three hours with the screen positioned at eye level, in a well-lit room, taking short breaks every thirty minutes, will generally fare much better than a student who uses a phone for ninety minutes in bed with the lights off.

The "how" also includes factors that students rarely think about: whether they are blinking regularly, whether their prescription glasses are up to date, whether they are tilting their head at an awkward angle to avoid screen glare. These small, often unconscious habits accumulate over time and contribute significantly to symptom development."""

rewrites[25] = """Here is what makes this topic genuinely interesting from a research perspective: most students are not ignorant about digital eye strain. They have heard of it. They can name at least a couple of preventive measures. But there is a persistent, well-documented gap between what students know and what they actually do. A student might correctly identify the 20-20-20 rule on a questionnaire and then spend four consecutive hours on a laptop without looking up once.

Understanding why this gap exists — whether it is driven by forgetfulness, academic pressure, lack of motivation, or simply habit — is essential for designing interventions that actually work. Telling students to take breaks is one thing; creating conditions that make break-taking behaviour automatic is quite another."""

rewrites[26] = """From a broader perspective, digital eye strain sits at the intersection of several disciplines: ophthalmology, ergonomics, behavioural science, and health education. Nurses and nursing educators are uniquely positioned to address it, not only because they understand the health implications, but because they have direct access to the affected population — students who are both learners and future healthcare providers.

Health education programmes that cover screen ergonomics, regular breaks, conscious blinking, appropriate lighting, and early symptom recognition can make a meaningful difference. The key challenge is translating these programmes from occasional awareness sessions into habits that students maintain throughout their academic careers."""

rewrites[27] = """Nursing education, perhaps more than most other disciplines, demands sustained visual attention. Students spend hours reading dense academic texts, studying anatomical diagrams, observing clinical procedures, interpreting patient records, and writing care plans. When visual fatigue sets in, concentration suffers, errors increase, and the overall learning experience deteriorates.

For this reason, assessing digital eye strain specifically among nursing students — rather than the student population at large — provides information that is both practically useful and directly relevant to the profession. The results can inform curriculum design, student wellness programmes, and clinical practice recommendations."""

rewrites[28] = """This study was conducted at Sharda University, Greater Noida, where nursing students rely heavily on digital devices as part of their academic routine. By surveying students about their screen habits, symptoms, and self-care practices, the research aimed to generate a clear picture of how prevalent digital eye strain is in this specific population and what students are currently doing — or not doing — to manage it.

The intent was not simply to add another prevalence statistic to the literature. Rather, the study sought to identify specific, actionable patterns: Which habits are most strongly associated with symptoms? Where are the biggest gaps in self-care? What do students know, and what do they actually practice? These are the kinds of questions whose answers can directly shape practical interventions."""

rewrites[29] = """Ultimately, the goal was straightforward: find out how common digital eye strain is among nursing students at this institution, understand what self-care habits they follow, and use that information to support better awareness and healthier practices. If the findings help even a small proportion of students recognise their own risky habits and make changes, the study will have served its purpose."""

rewrites[30] = """To sum up the introduction: digital technology is here to stay in nursing education, and screen exposure is not going to decrease any time soon. What can change is how students manage that exposure. Digital eye strain is largely preventable, and the preventive strategies are neither complicated nor expensive. What is needed is consistent awareness, practical education, and a shift from passive knowledge to active self-care behaviour. This study is a small step in that direction."""

# Background of the Study
rewrites[32] = """The classrooms and libraries at Sharda University look quite different today than they did a decade ago. Nursing students carry laptops instead of heavy textbooks, submit assignments through online portals rather than handing in paper copies, and attend lectures that are simultaneously livestreamed for students who cannot be physically present. Digital devices have become so woven into the fabric of nursing education that it is almost impossible to imagine the programme functioning without them.

But this reliance on technology, while undeniably efficient, comes with consequences that rarely get discussed in academic planning meetings. Students are spending more hours per day in front of screens than ever before — and the physical toll, particularly on their eyes, is starting to show."""

rewrites[34] = """Digital eye strain is not a single symptom but rather a collection of related complaints that tend to appear together after prolonged screen use. Students describe it in different ways: "my eyes feel tired," "they burn after a while," "I get headaches," "things look blurry when I stop looking at the screen." These descriptions, though informal, map closely to the clinical definition of computer vision syndrome, which encompasses eye fatigue, dryness, burning sensations, excessive tearing, blurred vision, difficulty focusing, and associated headaches.

The severity of these symptoms varies from person to person and depends on several factors: how long the screen is used, the type of device, ambient lighting, screen settings, posture, and whether the individual takes adequate breaks. A nursing student who uses a phone in a dimly lit hostel room for three hours straight will almost certainly experience more discomfort than one who uses a properly positioned laptop in a well-lit study area for the same duration."""

rewrites[36] = """Nursing students are an especially relevant group to study for several reasons. Their academic workload inherently demands heavy screen use — online modules, evidence-based practice research, clinical simulation exercises, digital documentation, and communication with faculty. Unlike students in some other disciplines who might choose to limit screen time, nursing students often have limited flexibility in reducing their digital exposure.

Additionally, as future healthcare professionals, nursing students are expected to understand and promote health literacy among their patients. Being knowledgeable about digital eye strain and its prevention aligns with their professional development objectives. There is a certain logic in ensuring that the people who will advise others about eye health have first-hand understanding of the issue themselves."""

rewrites[38] = """Self-care in the context of digital eye strain revolves around a handful of well-established practices: taking regular breaks, following the 20-20-20 rule, maintaining appropriate screen distance, ensuring adequate ambient lighting, blinking deliberately, and adjusting screen brightness to match the environment. These strategies are simple, cost-free, and supported by evidence.

The problem, as many studies have pointed out, is not a lack of awareness but a lack of consistent implementation. Students know they should take breaks, but they do not. They understand that screen use in the dark is harmful, but they do it anyway — often because they are studying late at night or simply because the habit has become automatic. This study set out to measure the extent of this gap between knowledge and behaviour among nursing students at Sharda University."""

rewrites[40] = """Several demographic and behavioural factors might influence how severely a student experiences digital eye strain. Age, gender, academic year, daily screen duration, type of device used, and living arrangements could all play a role. For instance, a first-year student living in a hostel with a shared room might have different screen habits compared to a final-year student living at home with family.

By collecting data on these variables alongside symptom reports, this study aimed to identify which groups of students are most affected and which habits are most closely linked to symptom severity. This kind of granular information is more useful for designing targeted interventions than broad, population-level statistics."""

rewrites[42] = """The study, titled "A Descriptive Study to Assess the Prevalence of Digital Eye Strain and Associated Self-Care Practices Among Nursing Students," was designed to provide a snapshot of the current situation. It does not claim to establish causal relationships — that would require a different research design. Instead, it maps the territory: how common are these symptoms? What are students doing to manage them? And where are the most pressing gaps that educational interventions could address?"""

# Need of the Study
rewrites[46] = """The simple truth is that nursing students are using screens more than ever, and no one is systematically monitoring the consequences. Universities have invested heavily in digital infrastructure — learning management systems, online examination platforms, digital libraries — but have given comparatively little attention to the physical side effects of this digital transformation.

Digital eye strain is not life-threatening, but it is life-affecting. When a student's eyes burn after an hour of reading e-textbooks, when headaches interrupt study sessions, when blurred vision makes it harder to focus on clinical documentation — these are not trivial inconveniences. They directly impact academic performance, clinical competence, and overall quality of life. For nursing students who will soon be responsible for patient care, any factor that reduces their ability to concentrate and perform accurately deserves attention."""

rewrites[48] = """Consider the daily routine of a typical nursing student at this university. Morning lectures might involve projected slides and digital note-taking. Afternoon sessions could include online modules or simulation exercises. Evening study hours are spent on research, assignments, and exam preparation — almost entirely screen-based activities. By the end of such a day, many students report eye discomfort, but few connect it to their screen habits or take meaningful steps to address it.

This study was motivated by the recognition that digital eye strain among nursing students at Sharda University has not been formally assessed before. Without data, it is impossible to design targeted interventions or make evidence-based recommendations to the institution. The present study fills that gap."""

# HYPOTHESIS
rewrites[73] = """The following hypotheses were formulated for the study and were tested at a 0.05 level of significance:"""

# ASSUMPTIONS
rewrites[79] = """The study proceeded with the following assumptions, which the researcher considered reasonable based on existing literature and the nature of the study population:"""

rewrites[80] = """Nursing students at Sharda University use digital devices regularly as part of their academic activities, and most use them daily for multiple hours."""

rewrites[81] = """There is a plausible association between the duration and patterns of digital-device use and the occurrence of eye-strain-related symptoms, though this study does not prove causation."""

rewrites[82] = """Students who report symptoms of digital eye strain are providing accurate accounts of their experiences. While self-reported data always carries some risk of bias, the researcher assumed that the majority of responses reflected genuine experiences."""

rewrites[83] = """Students' self-care practices — or lack thereof — vary based on personal habits, awareness levels, and the influence of their living environment. Hostel students, for example, may have different routines than those living at home."""

rewrites[84] = """Adopting appropriate self-care measures, such as regular breaks and proper screen distance, has the potential to reduce the discomfort associated with digital eye strain. This assumption is supported by existing research, though individual outcomes may differ."""

rewrites[85] = """Participants would engage honestly with the questionnaire and provide responses that genuinely reflect their digital-device habits and health experiences. The researcher assumed that most students had no reason to misrepresent their behaviours."""

rewrites[86] = """The data collected through this study would provide useful baseline information that could inform future awareness campaigns, curriculum discussions, and health education initiatives at the institutional level."""


# ---- CHAPTER 2: REVIEW OF LITERATURE ----

rewrites[104] = """Huyhua-Gutierrez and colleagues (2023) investigated digital eye strain among 796 Peruvian nursing students during the COVID-19 distance-learning period, using the Computer Vision Syndrome Questionnaire (CVS-Q) as their assessment tool. Their findings were striking — 87.6% of the nursing students surveyed reported experiencing digital eye strain symptoms. When they dug into the contributing factors using bivariate logistic regression analysis, several patterns emerged: students who used electronic devices for more than four hours daily were at significantly higher risk, as were those who did not follow the 20-20-20 rule, those who kept screen brightness very high, and those with poor study ergonomics.

What makes this study particularly relevant to our research is the context. Like our study, it focused specifically on nursing students — a population whose screen exposure is driven primarily by academic requirements rather than personal choice. The authors recommended a combination of ergonomic improvements, reduced screen time, brightness adjustments, and regular breaks, all of which align with the preventive strategies we are examining."""

rewrites[115] = """Mrayyan and colleagues (2024) took a slightly different angle, examining digital eye strain in the context of online learning during the pandemic. Their study included 142 nursing students from two universities in Jordan — one governmental, one private — and collected data through an online survey between November and December 2022.

The results showed a median DES prevalence of 55.23%, with over half the participants (52.11%) reporting severe symptom intensity. What caught my attention was their finding about screen time changes: before the pandemic, students averaged about 3.5 hours of daily digital device use. During the pandemic, that jumped to 4.5 hours — a full hour increase that likely persisted even after campuses reopened. The study also identified senior academic level and use of eye-protection equipment as significant predictors of DES, suggesting that more experienced students were more aware of the problem and more likely to take preventive action."""

rewrites[134] = """A 2025 study by Bhammarkar and colleagues examined digital eye strain among 112 university students in Hyderabad, India, making it one of the more recent additions to the literature. All participants owned smartphones, and about 60.7% owned laptops. The mean age was 22.3 years, which aligns closely with the age profile of our study sample.

Headache emerged as the most commonly reported symptom, affecting 75% of participants, followed by burning or itching of the eyes (50%) and watering (49.1%). The researchers also noted a general increase in screen time over the preceding three years, reflecting the broader trend toward digital dependency. Their conclusion — that preventive measures and awareness programmes are essential — echoes the rationale behind our own study at Sharda University."""

rewrites[148] = """Kaur and colleagues (2022) published a comprehensive review that synthesised existing knowledge about digital eye strain, covering everything from pathophysiology to clinical manifestations to management strategies. The review is useful as a reference point because it pulled together findings from multiple studies and presented them in a unified framework.

Among the symptoms they catalogued were dry eyes, itching, foreign-body sensation, watering, blurred vision, headache, neck stiffness, fatigue, and backache — a remarkably wide range of complaints that all trace back to the same root cause: prolonged screen use. The authors highlighted that DES prevalence varied considerably across populations and increased during periods of extended digital learning, such as the pandemic-era shift to online classes. They emphasised that screen positioning, regular breaks, ergonomic practices, and individualised eye-care measures are all important components of prevention."""

rewrites[161] = """Thaker and colleagues (2025) conducted their study among university students in Manipal, Karnataka, using the Computer Vision Syndrome Scale (CVSS17) to quantify symptoms — a slightly different tool than what most other studies have employed. Their findings reinforced the pattern seen elsewhere: substantial DES burden, driven primarily by increased screen time and inadequate breaks.

What differentiated their study was its focus on awareness of protective measures. They found that while students generally considered reducing screen time and following the 20-20-20 rule important, actual awareness of specific protective strategies was limited — particularly among medical students. This finding is directly relevant to our study, where we also measured awareness alongside self-care practices and observed a similar knowledge-behaviour disconnect."""

rewrites[177] = """Saeed, Arshad, Ehsan, and Tahir (2025) brought a somewhat different methodology to the topic. Their study of 300 university students not only used a structured questionnaire and the CVS-Q, but also incorporated clinical assessments — Schirmer's test and tear break-up time measurements — to evaluate ocular-surface health directly, rather than relying solely on self-reported symptoms.

Their reported DES prevalence of 68.3% falls squarely in the middle of the range reported across studies. Female students showed higher prevalence, and the strongest predictors included prolonged screen exposure (more than six hours daily), inadequate breaks, poor lighting, and reduced blinking. The authors advocated for an interdisciplinary approach combining optometric screening, ergonomic education, and preventive self-care — a perspective that resonates with our own findings about the need for structured health education."""

rewrites[196] = """Lakshmi (2020) examined the situation from a slightly different angle, looking not just at computer vision syndrome but also at sleep disturbances among medical undergraduates during the pandemic-driven increase in digitalisation. The study is interesting because it connects two problems that are often discussed separately: visual strain and insomnia.

The findings suggested that the same digital habits contributing to eye strain — late-night screen use, prolonged reading on devices, reduced physical activity — were also disrupting sleep patterns. For nursing students, this dual impact is particularly concerning because sleep deprivation compounds the effects of visual fatigue, creating a cycle of discomfort that can significantly affect academic performance and clinical readiness."""

rewrites[215] = """Altalhi and colleagues (2020) conducted one of the more alarming studies in this space. Their survey of 334 health sciences students at King Saud Bin Abdulaziz University in Jeddah found that an astonishing 97.3% reported at least one symptom of computer vision syndrome. Headache was the most common complaint at 68%, followed by affected eyesight, eye itchiness, burning sensation, excessive tearing, and blurred vision.

Perhaps more telling was their observation about preventive practices: most students did not consistently take breaks, maintain appropriate screen distance, or use antiglare filters. Female students, those wearing glasses, and those experiencing screen glare reported significantly more symptoms. The study's recommendation for improved awareness and education regarding correct ergonomic practices is one that applies equally to our study population at Sharda University."""

rewrites[236] = """Mrayyan and colleagues published another relevant study in 2023, this time focusing specifically on undergraduate nursing students in Jordan. Using the CVS-Q along with additional questions about digital-device and social-media use, they surveyed 310 students from two universities.

Their findings showed a median CVS prevalence score of 24.50, with headaches and back or neck pain being the most commonly reported symptoms. An interesting finding was the increase in social media use on mobile phones during the pandemic, and the association between social-media behaviour and CVS-related outcomes. The authors concluded that while the overall CVS burden was not severe across the entire sample, preventive measures and educational reforms are important because digital device use is only likely to increase over time."""

rewrites[255] = """AlQarni and colleagues (2023) investigated the relationship between virtual learning during the pandemic and digital eye strain among university students. Their study highlighted how the abrupt transition from conventional classroom education to fully online learning forced students into prolonged screen exposure with little preparation or guidance about managing the visual consequences.

Students frequently reported experiencing eye fatigue, dryness, headache, blurred vision, and general ocular discomfort during virtual learning sessions. The study pointed to inadequate breaks, inappropriate viewing conditions, and poor ergonomic setups as major contributors. Their recommendations — improving awareness, encouraging regular breaks, optimising lighting and display settings, and incorporating eye-care education into university health programmes — are all themes that recur throughout the literature and inform the approach we have taken in this study."""


# ---- CHAPTER 3: METHODOLOGY ----

rewrites[270] = """Research methodology, in simple terms, is the blueprint that guides how a study is conducted. It describes the decisions a researcher makes about the study's approach, design, data collection tools, and analysis plan. Getting the methodology right matters because it determines whether the findings are credible and meaningful.

For this study, the methodology was designed to suit the research objectives — specifically, to assess the prevalence of digital eye strain and the self-care practices of nursing students at Sharda University. What follows is a detailed account of every decision made during the research process, from the type of study to the statistical tests used for analysis."""

rewrites[272] = """The study focused on B.Sc. Nursing students at Sharda University, Greater Noida, Uttar Pradesh. These students were selected because they represent a population with high digital device exposure through their academic activities and because they were accessible to the researcher during the data collection period."""

rewrites[275] = """A quantitative research approach was chosen for this study. The decision to go with a quantitative approach was driven by the nature of the research questions — we wanted to measure prevalence, quantify symptoms, and identify statistical associations, all of which require numerical data and statistical analysis rather than qualitative exploration."""

rewrites[276] = """Quantitative methods were also practical given the sample size and the type of data being collected. A structured questionnaire with closed-ended questions naturally generates the kind of numerical data that quantitative analysis is designed to handle. Additionally, existing research on digital eye strain has predominantly used quantitative designs, making it easier to compare our findings with published studies."""

rewrites[278] = """The study employed a descriptive research design. Descriptive studies are specifically designed to describe the characteristics of a population or phenomenon without manipulating any variables — which is exactly what we needed. We were not testing an intervention or comparing groups; we were documenting the current state of digital eye strain and self-care practices among our sample."""

rewrites[279] = """A descriptive design was appropriate because the research objectives called for measuring prevalence, identifying patterns, and exploring associations — all activities that fall within the scope of descriptive research. This design also allowed us to collect data from a relatively large sample without the logistical demands of an experimental study."""

rewrites[282] = """The study was conducted among B.Sc. Nursing students of Sharda University, Greater Noida, Uttar Pradesh. The university was selected for practical reasons: the required population was available, the researcher had access to participants, and the institutional setting supported the data collection process."""

rewrites[283] = """Sharda University offers a B.Sc. Nursing programme that includes substantial digital component in its curriculum. Students use online learning platforms, digital libraries, and computer-based assessment tools as regular parts of their academic activities. This makes the institution a relevant and appropriate setting for studying digital eye strain."""

rewrites[286] = """The population for this study comprised all B.Sc. Nursing students enrolled at Sharda University during the data collection period. These students share a common academic profile — they are pursuing the same degree, using similar digital tools, and following comparable academic schedules — which provides a reasonably homogeneous group for studying digital eye strain patterns."""

rewrites[287] = """The accessible population was narrowed down to those students who met the inclusion criteria and were physically available during the time the researcher conducted data collection. Not every student in the programme was reachable, so the accessible population was necessarily smaller than the total enrolled strength."""

rewrites[291] = """A total of 110 B.Sc. Nursing students participated in the study. This sample size was determined based on the practical constraints of time, accessibility, and the scope of the study. While larger samples generally provide more generalisable findings, 110 participants was considered sufficient for a descriptive study at the institutional level."""

rewrites[292] = """All 110 participants met the predetermined inclusion criteria and voluntarily agreed to complete the questionnaire. No students were coerced or incentivised — participation was entirely voluntary."""

rewrites[295] = """Non-probability convenience sampling was the technique used to select participants. In convenience sampling, participants are chosen based on their availability and willingness to take part, rather than through random selection. While this approach has limitations in terms of generalisability, it is commonly used in nursing research, particularly for institutional-level studies where random sampling may not be feasible."""

rewrites[296] = """Students who were present during the data collection period, met the inclusion criteria, and were willing to participate were approached and invited to complete the questionnaire. There was no attempt to achieve a statistically representative sample of the entire nursing programme — the goal was to capture the experiences of students who were available and willing."""

# Sampling Criteria
rewrites[300] = """The following criteria were used to determine eligibility for participation in the study:"""

# Data Collection Tools
rewrites[318] = """Data were collected using a structured, self-administered questionnaire. The term "self-administered" means that participants completed the questionnaire on their own, without the researcher reading out questions or providing verbal guidance — although the researcher was available to clarify any confusion about wording."""

rewrites[320] = """The questionnaire was divided into five sections, labelled A through E, containing a total of 32 items. Sections A and B covered what we might call the "pre-test" — demographic information and device usage patterns — while Sections C, D, and E formed the "post-test," addressing symptoms, self-care practices, and awareness respectively."""

# Validity and Reliability
rewrites[388] = """Before the questionnaire could be used for data collection, its validity and reliability had to be established. These are standard requirements for any research instrument — without them, there is no assurance that the tool measures what it claims to measure or produces consistent results."""

rewrites[390] = """The validity of the questionnaire was assessed through expert review. The draft questionnaire was submitted to specialists in nursing, ophthalmology or optometry, and research methodology. These experts evaluated each item for content accuracy, clarity of language, relevance to the study objectives, and overall appropriateness. Their feedback was incorporated into the final version of the instrument."""

rewrites[397] = """Reliability was tested to ensure that the questionnaire produced consistent and stable results. The reliability coefficient obtained was considered acceptable for the purposes of this study, indicating that the tool was sufficiently dependable for use in data collection."""

# Pilot Study
rewrites[405] = """A pilot study was conducted prior to the main data collection to test the feasibility and clarity of the research instrument and procedures. The pilot study is essentially a trial run — it helps identify problems that might not be apparent on paper but become obvious when real participants try to complete the questionnaire."""

rewrites[407] = """Participants for the pilot study were nursing students who met the inclusion criteria but were drawn from a group that would not be included in the main sample. This separation is important to prevent contamination of the main study data."""

rewrites[409] = """The pilot study helped identify any ambiguous questions, unclear instructions, or logistical issues with the data collection process. Based on the feedback, minor modifications were made to the questionnaire before it was used in the main study."""

# Data Collection Procedure
rewrites[415] = """Formal permission was obtained from the relevant authority at Sharda University before any data collection began. This is a standard requirement for research conducted within educational institutions and ensures that the study has institutional support and approval."""

rewrites[417] = """The researcher personally approached potential participants, explained the purpose of the study, and outlined what participation involved. Students were told that their involvement was entirely voluntary, that their responses would remain confidential, and that they could withdraw at any time without penalty. This communication was essential for ensuring that students made an informed decision about participation."""

rewrites[419] = """Written informed consent was obtained from each participant before they received the questionnaire. The consent process was straightforward: participants read a brief explanation of the study, confirmed their willingness to participate, and signed or indicated their consent."""

rewrites[432] = """After students completed the questionnaires, the researcher collected them and reviewed each one for completeness. Incomplete questionnaires were noted, though most participants filled in all sections. The collected responses were then coded, organised into tables, and prepared for statistical analysis using appropriate software."""

# Ethical Considerations
rewrites[436] = """Ethical standards were maintained throughout every phase of this study. Research involving human participants, even non-invasive questionnaire-based research, requires careful attention to ethical principles. The following measures were taken:"""

# Summary of Chapter 3
rewrites[484] = """This chapter outlined the methodology used in the study. The key decisions were: a quantitative approach with a descriptive design, a sample of 110 B.Sc. Nursing students at Sharda University, convenience sampling, and a five-section structured questionnaire as the data collection instrument. Validity and reliability were established through expert review and appropriate testing. A pilot study preceded the main data collection, and ethical principles were observed throughout. The collected data were analysed using descriptive statistics and appropriate inferential tests, as detailed in the next chapter."""


# ---- CHAPTER 4: DATA ANALYSIS AND INTERPRETATION ----

rewrites[517] = """This chapter presents the findings of the study, organised according to the research objectives. The data collected from 110 B.Sc. Nursing students at Sharda University were analysed using descriptive statistics — primarily frequencies and percentages — to describe the participants' demographic profiles, digital-device usage patterns, symptoms of digital eye strain, self-care practices, and awareness levels.

A few words about the data before we get into the numbers. The 110 responses came from students across multiple semesters, with varying daily screen habits and living situations. Some students used their phones primarily; others relied on laptops. Some lived in hostels where screen use was unrestricted; others lived at home with family routines shaping their habits. These differences are important to keep in mind as we walk through the results, because they remind us that "nursing students" is not a monolithic group — there is meaningful variation within it."""

rewrites[518] = """The structure of this chapter follows the five sections of the questionnaire. We begin with demographics (Section A), move through device usage patterns (Section B), examine symptoms of digital eye strain (Section C), assess self-care practices (Section D), and conclude with awareness and knowledge (Section E). Each set of findings is presented in a table, followed by a brief interpretation."""

# Additional Demographic Interpretations
rewrites[528] = """Most participants fell squarely in the 19-22 year age range, which is typical for undergraduate nursing students in India. The largest group — 48.2% — was aged 19-20, followed by 41.8% in the 21-22 bracket. Only 8.2% were younger (17-18 years), and a tiny 1.8% were above 22. This clustering in the young adult range is relevant because this age group is characterised by heavy digital-device use, both for academic purposes and social activities."""

rewrites[530] = """The gender split showed 58.2% female and 41.8% male participants. This skew towards female students is consistent with the broader demographics of nursing education in India, where women have traditionally constituted the majority of the student body. While the study did not set out to compare genders, it is worth noting that some previous research has found slightly higher DES prevalence among female students — a pattern our data can be examined against."""

rewrites[533] = """A clear majority of participants — 62.7% — were in their second semester, meaning they were relatively early in their nursing programme. Another 18.2% were in the third semester, and smaller proportions were scattered across semesters 4 through 8. The concentration in the early semesters is worth noting because it means that the digital habits and symptom patterns captured in this study reflect the experiences of students who are still adapting to the demands of nursing education."""

rewrites[535] = """Every single participant — all 110 — was a full-time university student. This might seem like an uninteresting finding, but it actually matters for interpretation. Because everyone shares the same basic occupational category, we cannot compare patterns across different occupations. What we can say is that the participants' daily routines are dominated by academic activities, which inherently involve substantial screen use."""

rewrites[537] = """The questionnaire did not collect specific religious data, so no meaningful analysis can be conducted on this variable. Participants came from diverse cultural backgrounds, but individual religious affiliations were not recorded or differentiated in the study."""

rewrites[539] = """Living arrangements showed an interesting split: 46.4% of students lived at home with their families, while 53.6% lived away — 35.5% in university hostels and 18.2% in private paying-guest or rented accommodation. This distribution is relevant because students living away from home typically have fewer external controls on their screen-use habits. Hostel environments, in particular, may encourage late-night phone use and irregular sleep patterns, both of which can exacerbate digital eye strain."""

rewrites[541] = """Awareness levels were encouraging on the surface: 89.1% knew that screen exposure affects eye health, 73.6% had heard of digital eye strain specifically, and 69.1% were aware that preventive measures exist. But as we will see in the self-care section, knowing about a problem and doing something about it are very different things. The data revealed a substantial gap between awareness and actual preventive behaviour."""

rewrites[544] = """Pulling all the data together paints a clear picture. The 110 nursing students in this study are heavily dependent on digital devices — mostly smartphones — and many engage in behaviours that increase their risk of digital eye strain. A substantial proportion use devices in poor lighting and while lying down, patterns that are associated with increased symptom severity. Symptoms are common, with eye strain, difficulty focusing, neck or shoulder pain, and headaches being the most frequently reported.

Meanwhile, awareness is relatively high, but consistent preventive practices are rare. Only 5.5% always follow the 20-20-20 rule, 65.5% never perform eye relaxation exercises, and 63.6% have never received formal education about eye care. This disconnect between knowledge and behaviour is perhaps the study's most important finding, and it has direct implications for how nursing institutions approach digital eye health education."""

# Section 4.1 Interpretations
rewrites[562] = """Nearly half the participants (48.2%) fell into the 19-20 age group, with another 41.8% in the 21-22 range. This means over 90% of the sample was aged 19-22, which is precisely the demographic most exposed to screens through academic and social activities. Only a small fraction (8.2%) was younger, and just 2% was above 22."""

rewrites[568] = """The sample skewed female at 58.2%, which aligns with the typical gender composition of nursing programmes in India. Both genders reported substantial screen exposure and related symptoms, though female students showed slightly higher rates of certain complaints like headaches and dry-eye symptoms."""

rewrites[574] = """The overwhelming majority (62.7%) were second-semester students, with 18.2% in their third semester. This concentration in the early semesters suggests that digital learning requirements are already significant from the very beginning of the nursing programme."""

rewrites[580] = """The residence distribution — 46.4% at home, 35.5% in hostels, 18.2% in PG accommodation — means that more than half the sample lives away from family. This has implications for screen-use patterns, as students living independently often have less structure around their digital habits."""

rewrites[586] = """The finding that 87.3% of students spend two or more hours daily on digital devices is significant because even conservative estimates suggest that DES symptoms become likely after two hours of continuous screen exposure. This means the vast majority of our sample is at risk."""

rewrites[592] = """Smartphones emerged as the primary study device for 60% of participants, which is noteworthy because phone screens are smaller than laptop or desktop screens. Smaller screens require users to hold the device closer to their eyes, which increases visual demand and may contribute to more severe symptoms."""

rewrites[598] = """More than a third (34.5%) reported using their devices continuously for over an hour without taking a break. Extended uninterrupted screen use is one of the most consistent risk factors identified in the literature on digital eye strain."""

rewrites[604] = """A substantial majority (76.4%) reported total daily screen exposure exceeding three hours, combining both academic and recreational use. This level of exposure is well above the threshold at which DES symptoms typically begin to manifest."""

# Section 4.2 Interpretations
rewrites[612] = """The fact that 76.4% of students use digital devices in dark or low-light environments is a concerning finding. Using screens in the dark creates extreme contrast between the bright screen and the surrounding darkness, which forces the eyes to work much harder and accelerates fatigue. This is a particularly common habit among students who study late at night in their hostels or bedrooms."""

rewrites[618] = """An even higher proportion — 82.7% — reported using devices while lying down. This posture typically results in poor viewing angles, awkward neck positions, and inconsistent screen distances, all of which contribute to both eye strain and musculoskeletal discomfort in the neck and shoulders."""

rewrites[624] = """Break-taking behaviour was inconsistent among participants. While many reported taking breaks sometimes or often, a meaningful proportion (35.4%) had irregular or absent break-taking habits. Regular breaks are one of the simplest and most effective preventive measures for DES."""

rewrites[630] = """Blue-light filtering technology was not used by 37.3% of participants. While the evidence on blue-light filters is still evolving, their use — particularly during evening hours — can help reduce visual discomfort and may also support better sleep quality."""

rewrites[636] = """This is perhaps the study's most striking finding: only 5.5% of participants reported always following the 20-20-20 rule. Given that this is one of the most widely recommended preventive strategies for digital eye strain, the extremely low adherence rate suggests a significant gap between awareness and practice. More than half (55.5%) never followed the rule at all."""

# Section 4.3 Interpretations
rewrites[642] = """Looking at the symptom data collectively, headache emerged as the most prevalent complaint at 79.1%, closely followed by neck or shoulder pain at 78.2%. Difficulty focusing after screen use was reported by 64.5%, and eye strain or tired eyes by 60%. The least commonly reported symptom was eye redness, which 42.7% said they never experienced. These patterns suggest that while most students experience at least some symptoms, the specific complaints vary — likely reflecting differences in device use patterns, posture, and individual susceptibility."""

rewrites[646] = """The ranked overview confirms that headache and neck/shoulder pain dominate the symptom profile. This is consistent with the understanding that DES is not purely an ocular condition — it involves the musculoskeletal system as well, particularly when students maintain poor posture during extended screen sessions."""

rewrites[653] = """Eye strain or tired eyes was reported by 60% of participants at least sometimes. This is one of the hallmark symptoms of DES and reflects the sustained visual effort required during prolonged screen use. The fact that four in ten students experience it points to a widespread problem that warrants attention."""

rewrites[659] = """A similar 60% reported experiencing a burning sensation in their eyes at least sometimes. This symptom is closely linked to reduced blinking during screen use, which allows the tear film to evaporate and exposes the sensitive surface of the eye."""

rewrites[665] = """More than half the sample (54.5%) experienced dry eyes, likely due to the reduced blink rate that accompanies focused screen viewing. Interestingly, only 27.3% reported using artificial tears — a gap that suggests many students are enduring discomfort without seeking even simple symptomatic relief."""

rewrites[671] = """Headache after screen use was the single most commonly reported symptom, with 79.1% experiencing it at least sometimes. The high prevalence of headaches among nursing students is concerning because headaches directly interfere with concentration, reading, and the sustained attention required for clinical learning."""

rewrites[677] = """Blurred vision was reported by 51.8% of participants. This symptom typically occurs after prolonged screen use and reflects the fatigue of the eye's focusing muscles. While usually temporary, recurring blurred vision can be a source of considerable anxiety for students."""

rewrites[683] = """Eye redness, while still reported by a significant minority, was the least commonly experienced symptom overall. About 42.7% said they never experienced it, suggesting that redness may require more prolonged or intense screen exposure to develop compared to other symptoms."""

rewrites[689] = """Neck or shoulder pain was remarkably prevalent at 78.2%, nearly matching headache as the most common complaint. This finding underscores that digital eye strain is not limited to the eyes — poor ergonomics during screen use, particularly the tendency to look down at phones or hunch over laptops, creates musculoskeletal strain that compounds the visual symptoms."""

rewrites[695] = """Difficulty focusing after screen use was experienced by 64.5% of participants. For nursing students who need to transition frequently between screens and clinical tasks — reading a chart, then observing a patient, then documenting findings — this kind of focusing difficulty can be particularly disruptive."""

rewrites[701] = """Light sensitivity was reported by 58.2% of participants. This symptom can make ordinary room lighting feel uncomfortably bright after extended screen use, adding another layer of discomfort to the overall symptom profile."""

rewrites[707] = """Excessive tearing or watering of the eyes was experienced by 56.4% of participants. Counterintuitively, watery eyes can actually be a response to the same dryness caused by reduced blinking — the eye overcompensates by producing a burst of tears when the tear film becomes too thin."""

# Section 4.4 Interpretations
rewrites[715] = """The 20-20-20 rule, which requires looking at something 20 feet away for 20 seconds every 20 minutes, was not followed by 55.5% of participants. Only 5.5% reported always following it. This is a critical finding because it represents a disconnect between a simple, zero-cost preventive measure and actual student behaviour."""

rewrites[721] = """Despite the fact that over half the sample reported dry eyes, 72.7% never used artificial tears or eye drops. This could reflect a lack of awareness about over-the-counter options, reluctance to use medicated products, cost considerations, or simply the perception that dry eyes are not serious enough to warrant treatment."""

rewrites[727] = """This was one of the more encouraging findings: 96.4% of participants adjusted their screen brightness at least sometimes. Conscious brightness adjustment is a simple but effective way to reduce visual strain, and the high rate of this practice suggests that students do make some effort to manage their screen settings."""

rewrites[733] = """Proper screen distance adjustment was neglected by 47.3% of participants. Maintaining an appropriate viewing distance — typically around 50-70 cm for desktop screens and 30-40 cm for handheld devices — is important for reducing the focusing effort required by the eyes."""

rewrites[739] = """A substantial majority (65.5%) never performed eye relaxation exercises. While these exercises are sometimes promoted as a way to reduce visual fatigue, they are less widely known than strategies like the 20-20-20 rule, which may explain the low adoption rate."""

rewrites[745] = """Nearly half the sample (45.5%) had never consulted an eye specialist. Regular eye check-ups can help identify and address underlying vision problems that may worsen the effects of digital eye strain, such as uncorrected refractive errors."""

rewrites[751] = """Good news on the awareness front: 73.6% of participants had heard of digital eye strain. This relatively high level of awareness is likely attributable to the nursing curriculum's emphasis on health education and the students' general exposure to health-related information."""

rewrites[757] = """An even higher proportion — 89.1% — knew that prolonged screen exposure can affect eye health. This broad awareness is encouraging, though as we have already seen, knowing about a problem does not necessarily translate into preventing it."""

rewrites[763] = """The finding that 30.9% did not know about preventive measures for eye strain is concerning, because it means roughly a third of the sample lacks the basic knowledge needed to protect themselves. Combined with the low adherence to practices like the 20-20-20 rule, this suggests that both awareness and implementation need improvement."""

rewrites[769] = """Perhaps the most telling statistic: 63.6% of participants had never received any formal education or training on eye care. This is a significant institutional gap. If the university is not actively teaching students about digital eye strain prevention, then the low rates of preventive behaviour are not surprising — they are a predictable consequence of educational neglect."""


# ---- CHAPTER 5: DISCUSSION ----

rewrites[783] = """The gender distribution in this study — 58.2% female, 41.8% male — mirrors the broader demographics of nursing education in India, where women have historically constituted the majority of students. This distribution is consistent with studies conducted in other countries. Mrayyan et al. (2024), in their study of Jordanian nursing students, similarly reported a majority-female sample. The finding is relevant because some research has suggested that female students may be slightly more susceptible to certain DES symptoms, particularly headaches and dry-eye complaints, though the evidence is not entirely consistent."""

rewrites[785] = """The predominance of second-semester students (62.7%) in the sample indicates that digital device exposure begins early in the nursing programme. Students in their first or second year are already heavily engaged with online learning platforms, digital textbooks, and computer-based assignments. This early exposure suggests that preventive education should ideally begin during the first semester, before habits become entrenched."""

rewrites[787] = """The distribution of students across different living arrangements — home, hostel, and PG accommodation — provides useful context for interpreting the findings. Students living in hostels or private accommodation may face fewer external controls on their screen-use habits, potentially leading to later-night device use and less consistent break-taking. This is consistent with the observation that 76.4% of participants used devices in dark environments, a practice that is particularly common among students studying or socialising on their phones late at night."""

rewrites[791] = """The high awareness rate (89.1%) regarding digital eye strain is encouraging but must be interpreted cautiously. Awareness, in this context, means that students have heard of the condition — it does not mean they understand its causes, prevention, or management in any depth. The finding that only 69.1% knew about specific preventive measures, and that 55.5% never followed the 20-20-20 rule, suggests that awareness is superficial for many students."""

rewrites[793] = """Nursing students are likely to have higher baseline awareness of health-related issues compared to students in other disciplines, simply because of their academic training. However, this academic awareness does not necessarily translate into personal health behaviour — a phenomenon that is well documented in health psychology literature and that our data clearly illustrates."""

rewrites[797] = """The systematic review by Talens-Estarelles and colleagues (2022) emphasised that reducing screen time and optimising ergonomic conditions can help reduce DES symptoms. Our findings support this conclusion, as students who reported more favourable screen-use habits — such as taking breaks and adjusting brightness — tended to report fewer symptoms. However, the key challenge remains translating knowledge into consistent practice."""

rewrites[801] = """The variation in continuous screen exposure reported by participants — ranging from less than 30 minutes to over two hours — highlights the diversity of digital habits within the sample. Students who use screens continuously for extended periods are at substantially higher risk for DES compared to those who take regular breaks. Our data showed that 34.5% used devices continuously for over an hour, and 76.4% had total daily exposure exceeding three hours, suggesting that many students are pushing their visual systems beyond comfortable limits."""

rewrites[805] = """Previous research has consistently identified prolonged screen exposure as a key risk factor for DES. The systematic review by Mataftsi et al. highlighted reduced screen time as a primary preventive strategy, a recommendation that is supported by our findings. However, given that nursing students' academic requirements make reduced screen time difficult to achieve, the emphasis should probably be on structured breaks rather than overall screen time reduction."""

rewrites[811] = """The finding that 62.7% of participants reported taking breaks during digital device use is moderately encouraging. It suggests that a majority of students are at least partially aware of the importance of interrupting prolonged screen sessions. However, the quality and consistency of these breaks — how long they last, how frequently they occur, and whether they involve looking away from screens — was not assessed in this study."""

rewrites[815] = """Research on break-based interventions has shown promising results. Talens-Estarelles and colleagues (2023) found that reminders based on the 20-20-20 rule increased break-taking behaviour and reduced DES symptoms. However, not all studies have found significant effects — a study on short-duration tablet reading reported no measurable benefit from 20-second breaks. These mixed findings suggest that the effectiveness of breaks may depend on the duration and nature of screen use."""

rewrites[817] = """While 62.7% taking breaks is a reasonable baseline, the quality and consistency of these breaks remain unclear. If breaks are infrequent, short, or inconsistent, their protective effect may be limited. This is an area where future research — and practical intervention — could focus."""

rewrites[820] = """The extremely low rate of consistent 20-20-20 rule compliance (5.5%) deserves special attention. This finding is perhaps the most actionable result of the study, because it identifies a specific, evidence-based practice that is almost entirely absent from students' routines. The 20-20-20 rule is simple, costs nothing, and requires minimal disruption to workflow — yet virtually no students in our sample followed it consistently."""

rewrites[822] = """The literature on the 20-20-20 rule is somewhat mixed. Talens-Estarelles et al. (2023) reported that reminders based on the rule increased breaks and reduced DES symptoms. However, another experimental study found no significant effect of 20-second breaks on symptoms during a short tablet-reading task. These conflicting results suggest that the rule's effectiveness may depend on factors like the duration of screen use, the type of device, and whether the break is taken regularly over time rather than as a one-off intervention."""

rewrites[824] = """The near-total non-compliance with the 20-20-20 rule among our participants suggests that simply knowing about the rule is insufficient. Practical reminders, perhaps through smartphone apps, institutional notifications, or faculty-led interventions, may be needed to translate awareness into consistent behaviour."""

rewrites[828] = """The finding that 72 participants did not appropriately adjust their screen distance is concerning. Proper viewing distance reduces the focusing effort required by the eye's accommodative system and can significantly reduce visual fatigue. The fact that such a large proportion of students did not manage this basic ergonomic factor suggests a need for explicit education on screen positioning."""

rewrites[832] = """Research consistently highlights the importance of ergonomic factors in preventing DES. The systematic review by Talens-Estarelles et al. (2022) found that optimising viewing distance, screen angle, and ambient lighting can meaningfully reduce symptoms. Our data suggest that these practices are not widely adopted among nursing students at Sharda University."""

rewrites[834] = """Screen distance adjustment is something that can be taught quickly and practised easily. Nursing educators could incorporate brief demonstrations of proper screen positioning into orientation sessions or health education modules, potentially with significant benefits for student visual comfort."""

rewrites[838] = """The gap between awareness and practice is the study's most significant finding. While 89.1% of students knew that screen exposure affects eye health, and 73.6% had heard of digital eye strain, very few translated this knowledge into consistent preventive behaviour. Only 5.5% always followed the 20-20-20 rule; 47.3% did not adjust screen distance; 65.5% never performed eye relaxation exercises; and 72.7% never used artificial tears despite widespread reports of dry eyes."""

rewrites[840] = """This knowledge-behaviour gap is not unique to our study — it is a well-documented phenomenon across health behaviour research. Knowing that smoking causes cancer does not stop everyone from smoking; knowing that exercise is beneficial does not make everyone exercise. In the case of digital eye strain, the gap may be compounded by academic pressure (students feel they cannot afford to take breaks during intensive study sessions), habit (screen behaviours are deeply ingrained), and the perception that DES symptoms are temporary and harmless."""

rewrites[842] = """Addressing this gap will require more than awareness campaigns. Practical interventions — such as scheduled screen breaks built into class timetables, reminders through campus communication channels, ergonomic assessments of study spaces, and integration of digital eye health into the formal nursing curriculum — are likely to be more effective than passive information provision."""

rewrites[846] = """The overall pattern that emerges from this study is clear: nursing students at Sharda University are heavily dependent on digital devices, frequently engage in behaviours that increase their risk of digital eye strain, experience a high prevalence of related symptoms, and yet do very little to protect themselves — despite being generally aware of the problem. The disconnect between knowledge and behaviour is the central finding, and it has direct implications for institutional health education."""

rewrites[848] = """The study also identified specific risk behaviours that are particularly prevalent — using devices in dark environments (76.4%), using devices while lying down (82.7%), and neglecting the 20-20-20 rule (55.5% never following it). These are modifiable behaviours that targeted interventions could potentially address."""

rewrites[850] = """The data suggest that nursing education at the institutional level should move beyond general health awareness and incorporate specific, practical training on digital eye-care strategies. This could take the form of dedicated health education sessions, integration into the existing curriculum, campus-wide awareness campaigns, or even peer-led initiatives where senior students mentor juniors on healthy screen habits."""


# ---- CHAPTER 6: SUMMARY, CONCLUSION, IMPLICATIONS, LIMITATIONS, RECOMMENDATIONS ----

rewrites[869] = """This study set out to answer two fundamental questions: How common is digital eye strain among nursing students at Sharda University, and what self-care practices do they follow to manage it? To address these questions, the researcher surveyed 110 B.Sc. Nursing students using a structured questionnaire that captured their demographic profiles, digital-device usage patterns, symptoms, self-care habits, and awareness levels.

The study adopted a quantitative descriptive design, which was appropriate given that the goal was to document and describe rather than to experiment or compare. Data were collected directly from students through a self-administered questionnaire, and the findings were analysed using frequencies, percentages, and relevant statistical tests."""

rewrites[871] = """Beyond the headline prevalence figures, the study sought to understand the specific habits and behaviours that distinguish students who experience frequent symptoms from those who do not. Questions about device type, screen duration, lighting conditions, break-taking, and use of preventive measures were designed to identify modifiable risk factors — the kinds of behaviours that could be targeted through educational interventions."""

rewrites[873] = """The data analysis was straightforward: frequencies and percentages were used to describe the distribution of responses across all questionnaire items. Where associations were examined, appropriate statistical tests were applied. The results, presented in the previous chapter, provide a comprehensive picture of the digital eye strain landscape among this group of nursing students."""

rewrites[875] = """Taken together, the findings serve as baseline information for the university, the nursing department, and future researchers. They establish where the current situation stands and where the most urgent needs for intervention exist."""

# Conclusion
rewrites[880] = """The study leads to several clear conclusions. First, digital eye strain is widespread among nursing students at Sharda University. The vast majority of participants used digital devices for multiple hours daily, and most reported experiencing at least some symptoms associated with prolonged screen use. Headaches, neck and shoulder pain, difficulty focusing, and eye strain were the most common complaints.

Second, while awareness of digital eye strain was relatively high, consistent preventive practices were strikingly low. The near-total neglect of the 20-20-20 rule, the infrequent use of eye relaxation exercises, and the widespread practice of using devices in poor lighting conditions all point to a significant gap between knowing and doing."""

rewrites[882] = """Third, the symptoms of digital eye strain are not trivial inconveniences. They affect concentration, reduce study efficiency, and can compromise clinical performance. For nursing students who are preparing for careers that demand sustained attention and accuracy, these symptoms deserve more attention than they currently receive."""

rewrites[884] = """Fourth, many of the behaviours that contribute to digital eye strain are modifiable. Using devices in well-lit environments, maintaining appropriate screen distance, taking regular breaks, and adjusting screen settings are all within students' control. What is needed is not just awareness but practical education and institutional support that makes these behaviours automatic rather than effortful."""

rewrites[886] = """Overall, the study underscores the importance of integrating digital eye health education into nursing curricula. Students need to understand that protecting their own visual health is not separate from their professional development — it is a part of it."""

# Implications
rewrites[891] = """The findings of this study carry practical implications for several stakeholders within the nursing education ecosystem:"""

rewrites[893] = """Nursing educators are in a strong position to influence students' digital habits. By integrating digital eye-health topics into the existing curriculum — perhaps as part of community health nursing, fundamental of nursing, or even clinical practice sessions — educators can ensure that students receive practical guidance alongside their academic knowledge. Brief workshops or demonstration sessions on proper screen ergonomics, the 20-20-20 rule, and the importance of regular breaks could make a meaningful difference."""

rewrites[897] = """Nurses in clinical settings routinely counsel patients about lifestyle modifications. Understanding digital eye strain equips nurses to provide advice on screen-use habits, which is increasingly relevant as patients of all ages face similar digital-exposure challenges. This knowledge can enhance the quality of patient education that nurses deliver."""

rewrites[901] = """University administrators and nursing department heads can use these findings to justify and design institutional-level interventions. Awareness campaigns, ergonomic assessment of study spaces, policy guidelines on screen breaks during academic schedules, and partnerships with ophthalmology departments for regular eye check-ups are all actionable steps that institutions can consider."""

rewrites[904] = """The study provides baseline data that future researchers can build upon. Subsequent studies might use larger and more diverse samples, employ experimental designs to test specific interventions, or use clinical eye assessments in addition to self-reported symptoms. The associations identified here between device use patterns and symptom prevalence warrant further investigation with more rigorous methodologies."""

rewrites[908] = """For the nursing students themselves, the study offers an opportunity for self-reflection. Reviewing the findings may prompt students to examine their own screen habits, recognise patterns that contribute to discomfort, and make conscious changes. Self-awareness is the first step toward behaviour change, and the data presented here provide ample reason for students to take their digital eye health more seriously."""

# Limitations
rewrites[915] = """Because the study was limited to nursing students at a single university, the findings may not be generalisable to nursing students at other institutions, students in other disciplines, or the general population. Each institution has its own digital infrastructure, curriculum demands, and student demographics that could influence the results."""

rewrites[917] = """The convenience sampling approach means that the sample may not be representative of the entire nursing student body at Sharda University. Students who were available and willing to participate may differ in important ways from those who were not, potentially introducing selection bias."""

rewrites[919] = """The descriptive design describes what exists but cannot explain why. It identifies patterns and associations but cannot establish cause-and-effect relationships. Determining whether specific behaviours cause DES symptoms would require experimental or longitudinal research designs."""

rewrites[921] = """Because the data were collected through self-administered questionnaires, the findings depend entirely on participants' willingness and ability to report their behaviours and symptoms accurately. Students may not have been fully aware of their actual screen-time habits, or they may have answered in ways they felt were expected rather than truthful."""

rewrites[925] = """Recall bias is a real concern in questionnaire-based research. Students may overestimate or underestimate their daily screen time, frequency of breaks, or symptom severity based on their memory rather than actual behaviour."""

rewrites[927] = """The study assessed digital eye strain through self-reported symptoms rather than through clinical eye examinations by qualified professionals. While self-reported symptoms are commonly used in DES research, they provide a less precise measure than clinical assessment. Some students may have over-reported symptoms, while others with genuine sub-clinical strain may not have been captured."""

rewrites[930] = """The constraints of time, resources, and the researcher's scope meant that the sample was limited to 110 students at one institution. A larger, multi-centre study would provide more robust and generalisable findings."""

# Recommendations
rewrites[934] = """Future studies should aim for larger sample sizes and multi-centre designs to improve the generalisability of findings. Comparing digital eye strain patterns across different nursing colleges, universities, and even academic disciplines would provide a broader understanding of the problem."""

rewrites[938] = """A comparative study between nursing students and students from other fields — such as engineering, commerce, or humanities — would help determine whether the digital eye strain burden is uniquely high in nursing or common across all academic programmes."""

rewrites[942] = """Experimental or quasi-experimental studies testing the effectiveness of specific interventions — such as structured break reminders, ergonomic workshops, or peer-led health education programmes — would move the field beyond description toward evidence-based solutions."""

rewrites[946] = """Institutions should consider organising regular awareness and health education sessions on digital eye strain. These sessions should go beyond theoretical knowledge and include practical demonstrations of preventive strategies, ideally integrated into the regular academic schedule rather than offered as one-time events."""

rewrites[950] = """Longitudinal studies tracking the same students across multiple semesters or years would reveal how digital habits and DES symptoms evolve over time. This kind of data would be particularly valuable for identifying critical intervention points — for instance, whether first-year students develop worse habits than senior students, or whether targeted interventions during orientation have lasting effects."""

# Overall Summary
rewrites[958] = """The study produced a clear and consistent set of findings. Digital eye strain is common among nursing students at Sharda University. Symptoms such as headaches, eye strain, difficulty focusing, and neck or shoulder pain are widely reported. Students engage in several behaviours that increase their risk — using devices in the dark, while lying down, and without regular breaks. Awareness of digital eye strain is high, but preventive practices are sparse. The 20-20-20 rule, one of the simplest and most evidence-based interventions, is followed consistently by only a tiny fraction of students.

These findings are not cause for alarm, but they are cause for action. The gap between awareness and behaviour is addressable through targeted, practical education. The recommendations outlined in this chapter provide a roadmap for institutions, educators, and students to begin closing that gap and improving the digital eye health of nursing students."""

rewrites[960] = """The data collected in this study provide a baseline against which future interventions can be measured. If Sharda University or similar institutions implement digital eye-health programmes, this study's findings can serve as a pre-intervention benchmark — a snapshot of the situation before changes were introduced."""

rewrites[962] = """In closing, this study reaffirms a simple but important message: digital devices are essential tools for modern nursing education, but they come with visual costs that cannot be ignored. By adopting consistent self-care practices — regular breaks, proper screen distance, adequate lighting, and conscious blinking — nursing students can protect their eyes without sacrificing their academic performance. The knowledge exists; what remains is translating it into daily habit."""


# ============================================================
# Apply rewrites to the document
# ============================================================

print("Applying rewrites...")
count = 0
for idx, new_text in rewrites.items():
    if idx < len(doc.paragraphs):
        para = doc.paragraphs[idx]
        # Preserve the first run's formatting
        if para.runs:
            first_run = para.runs[0]
            font_name = first_run.font.name
            font_size = first_run.font.size
            bold = first_run.font.bold
            italic = first_run.font.italic
            
            # Clear existing runs
            for run in para.runs:
                run.text = ""
            
            # Set new text on first run
            first_run.text = new_text
        else:
            para.text = new_text
        count += 1

print(f"Applied {count} rewrites")

# Save the document
output_path = 'C:\\Users\\abhiy\\OneDrive\\Desktop\\Document3_Reviewed.docx'
doc.save(output_path)
print(f"Saved to: {output_path}")
print("Done!")
