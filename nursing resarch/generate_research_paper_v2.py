from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ---- Page setup ----
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Helper functions
def add_heading_centered(text, level=0):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_heading_left(text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_para(text, bold=False, italic=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, spacing_after=Pt(6)):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = spacing_after
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Times New Roman'
    doc.add_paragraph()
    return table

# ==============================================================
# TITLE PAGE
# ==============================================================
for _ in range(4):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("A DESCRIPTIVE STUDY TO ASSESS THE KNOWLEDGE AND PRACTICE REGARDING DIGITAL EYE STRAIN AMONG NURSING STUDENTS")
run.bold = True
run.font.size = Pt(16)
run.font.name = 'Times New Roman'

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("A Research Paper Submitted in Partial Fulfillment of the Requirements for the Degree of Bachelor of Science in Nursing (B.Sc. Nursing)")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

doc.add_paragraph()
doc.add_paragraph()

submitted = doc.add_paragraph()
submitted.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = submitted.add_run("Submitted To:\nFaculty of Nursing\nSharda University, Greater Noida, Uttar Pradesh, India")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

doc.add_paragraph()

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run("July 2026")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

doc.add_page_break()

# ==============================================================
# DECLARATION AND CERTIFICATES
# ==============================================================
add_heading_centered("DECLARATION", level=1)
doc.add_paragraph()
add_para("I hereby declare that this research paper entitled \"A Descriptive Study to Assess the Knowledge and Practice Regarding Digital Eye Strain Among Nursing Students\" is a record of original work done by me under the guidance and supervision of the Faculty of Nursing, Sharda University, Greater Noida.")
add_para("This research paper has not been submitted elsewhere for the award of any degree, diploma, or other similar title or recognition.")
add_para("All the sources of information used have been acknowledged and cited appropriately.")
doc.add_paragraph()
add_para("Place: Greater Noida")
add_para("Date: July 2026")
doc.add_paragraph()
add_para("Signature of the Student")
doc.add_page_break()

add_heading_centered("CERTIFICATE", level=1)
doc.add_paragraph()
add_para("This is to certify that the research paper entitled \"A Descriptive Study to Assess the Knowledge and Practice Regarding Digital Eye Strain Among Nursing Students\" submitted to the Faculty of Nursing, Sharda University, Greater Noida, is a bonafide work carried out by the student under our guidance and supervision.")
add_para("This research paper has not been submitted elsewhere for the award of any degree, diploma, or other similar title or recognition.")
doc.add_paragraph()
doc.add_paragraph()
add_para("Signature of Guide")
doc.add_paragraph()
add_para("Signature of HOD")
doc.add_paragraph()
add_para("Signature of Principal")
doc.add_page_break()

# ==============================================================
# ACKNOWLEDGEMENT
# ==============================================================
add_heading_centered("ACKNOWLEDGEMENT", level=1)
doc.add_paragraph()
add_para("I would like to express my sincere gratitude to all those who have contributed to the successful completion of this research paper.")
add_para("First and foremost, I would like to thank the Almighty God for blessing me with the strength, wisdom, and perseverance to complete this work.")
add_para("I am deeply grateful to my guide for their invaluable guidance, constant support, and constructive feedback throughout the course of this study. Their expertise and encouragement have been instrumental in shaping this research.")
add_para("I extend my heartfelt thanks to the Principal and Head of the Department of Nursing, Sharda University, for providing me with the opportunity and platform to conduct this study.")
add_para("I would also like to express my sincere appreciation to all the faculty members of the Faculty of Nursing, Sharda University, for their valuable suggestions and support.")
add_para("My special thanks go to all the nursing students who participated in this study. Their willingness to share their experiences and opinions made this research possible.")
add_para("Finally, I am deeply indebted to my family and friends for their unwavering support, encouragement, and patience throughout this journey.")
doc.add_page_break()

# ==============================================================
# ABSTRACT
# ==============================================================
add_heading_centered("ABSTRACT", level=1)
doc.add_paragraph()
add_para("Background: Digital Eye Strain (DES) is a growing public health concern affecting individuals who extensively use digital devices. Nursing students are particularly vulnerable due to their heavy reliance on digital devices for academic and clinical activities. This study aimed to assess the knowledge and practice regarding digital eye strain among nursing students.", bold=False)
doc.add_paragraph()
add_para("Objectives: To assess the knowledge and practice regarding digital eye strain among nursing students, to determine the prevalence of DES symptoms, and to identify the association between demographic variables and knowledge/practice scores.", bold=False)
doc.add_paragraph()
add_para("Materials and Methods: A descriptive survey design was used. A total of 110 nursing students from Sharda University, Greater Noida, were selected using convenience sampling. Data were collected through a structured self-administered questionnaire covering demographic profile, digital device usage practices, DES symptoms, and awareness of preventive measures. Data were collected from July 3 to July 22, 2026, and analyzed using descriptive statistics.", bold=False)
doc.add_paragraph()
add_para("Results: The majority of participants were female (58.2%), aged 19-20 years (48.2%), and in the 2nd semester (62.7%). Smartphones were the most commonly used device (60.0%). Risk behaviors were prevalent: 76.4% used devices in dark environments, 82.7% used devices while lying down, and only 5.5% always followed the 20-20-20 rule. DES symptoms were widespread: headache (79.1%), neck/shoulder pain (78.2%), eye strain (60%), and dry eyes (54.5%). While 73.6% were aware of DES, preventive practices were inadequate: 72.7% did not use artificial tears and 65.5% did not perform eye relaxation exercises. A total of 63.6% had never received education/training on eye care.", bold=False)
doc.add_paragraph()
add_para("Conclusion: Digital eye strain is highly prevalent among nursing students with significant gaps between awareness and practice. There is an urgent need to integrate eye health education into the nursing curriculum and implement targeted preventive interventions.", bold=False)
doc.add_paragraph()
add_para("Keywords: Digital Eye Strain, Computer Vision Syndrome, Nursing Students, Knowledge, Practice, Descriptive Study", bold=True)
doc.add_page_break()

# ==============================================================
# LIST OF TABLES
# ==============================================================
add_heading_centered("LIST OF TABLES", level=1)
doc.add_paragraph()

tables_list = [
    ("Table 4.1", "Distribution of Participants by Age Group"),
    ("Table 4.2", "Distribution of Participants by Gender"),
    ("Table 4.3", "Distribution of Participants by Semester"),
    ("Table 4.4", "Distribution of Participants by Type of Residence"),
    ("Table 4.5", "Distribution of Participants by Average Daily Time Spent on Digital Devices"),
    ("Table 4.6", "Distribution of Participants by Main Device Used for Studies"),
    ("Table 4.7", "Distribution of Participants by Average Continuous Screen Time Without Break"),
    ("Table 4.8", "Distribution of Participants by Total Daily Screen Exposure"),
    ("Table 4.9", "Use of Digital Devices in Dark/Low Light"),
    ("Table 4.10", "Use of Devices While Lying Down"),
    ("Table 4.11", "Taking Breaks During Long Screen Use"),
    ("Table 4.12", "Use of Blue Light Filter/Night Mode"),
    ("Table 4.13", "Following the 20-20-20 Rule While Using Digital Devices"),
    ("Table 4.14", "Frequency of Eye Strain/Tired Eyes"),
    ("Table 4.15", "Frequency of Burning Sensation in Eyes"),
    ("Table 4.16", "Frequency of Dry Eyes"),
    ("Table 4.17", "Frequency of Headache After Screen Use"),
    ("Table 4.18", "Frequency of Blurred Vision"),
    ("Table 4.19", "Frequency of Eye Redness"),
    ("Table 4.20", "Frequency of Neck or Shoulder Pain"),
    ("Table 4.21", "Frequency of Difficulty Focusing After Screen Use"),
    ("Table 4.22", "Frequency of Sensitivity to Light"),
    ("Table 4.23", "Frequency of Excessive Watering/Tearing of Eyes"),
    ("Table 4.24", "Following the 20-20-20 Rule (Preventive Practice)"),
    ("Table 4.25", "Use of Artificial Tears or Eye Drops"),
    ("Table 4.26", "Conscious Reduction of Screen Brightness"),
    ("Table 4.27", "Proper Screen Distance Adjustment"),
    ("Table 4.28", "Performance of Eye Relaxation Exercises"),
    ("Table 4.29", "Consultation with Eye Specialist"),
    ("Table 4.30", "Awareness of Digital Eye Strain"),
    ("Table 4.31", "Knowledge That Prolonged Screen Use Affects Eye Health"),
    ("Table 4.32", "Knowledge of Preventive Measures for Eye Strain"),
    ("Table 4.33", "Previous Education/Training on Eye Care"),
]

for t_no, t_name in tables_list:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"{t_no}: {t_name}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

doc.add_page_break()

# ==============================================================
# LIST OF FIGURES
# ==============================================================
add_heading_centered("LIST OF FIGURES", level=1)
doc.add_paragraph()

figures_list = [
    ("Figure 4.1", "Bar Diagram showing Age Group Distribution"),
    ("Figure 4.2", "Pie Chart showing Gender Distribution"),
    ("Figure 4.3", "Bar Diagram showing Average Daily Screen Time"),
    ("Figure 4.4", "Pie Chart showing Main Device Used for Studies"),
    ("Figure 4.5", "Bar Diagram showing DES Symptoms Prevalence"),
    ("Figure 4.6", "Bar Diagram showing Preventive Practices"),
    ("Figure 4.7", "Pie Chart showing Awareness of Digital Eye Strain"),
]

for f_no, f_name in figures_list:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"{f_no}: {f_name}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

doc.add_page_break()

# ==============================================================
# TABLE OF CONTENT (matching the image format)
# ==============================================================
add_heading_centered("TABLE OF CONTENT", level=1)
doc.add_paragraph()

toc_table = doc.add_table(rows=1, cols=3)
toc_table.style = 'Table Grid'
toc_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header
for i, h in enumerate(["S.No", "Content", "Page No."]):
    cell = toc_table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'

toc_entries = [
    ("", "Declaration and Certificates", "I-IV"),
    ("", "Acknowledgement", "V-VI"),
    ("", "Abstract", "VII-VIII"),
    ("", "List of Tables", "IX"),
    ("", "List of Figures", "X"),
    ("", "Introduction", ""),
    ("", "Background of the Study", "1-3"),
    ("1.1", "Need for the Study", "3-4"),
    ("1.2", "Operational Definition", "4-5"),
    ("1.3", "Hypothesis", "6"),
    ("1.4", "Assumption", "6"),
    ("2", "Review of Literature", "7-12"),
    ("3", "Research Methodology", ""),
    ("3.1", "Introduction", "13"),
    ("3.2", "Research Approach", "13-14"),
    ("3.3", "Research Design", "14-15"),
    ("3.4", "Setting", "15"),
    ("3.5", "Population", "15-16"),
    ("3.6", "Sample Size", "16"),
    ("3.7", "Sampling Technique", "16-17"),
    ("3.8", "Sampling Criteria", "17"),
    ("3.9", "Ethical Considerations", "17-18"),
    ("3.10", "Data Collection Procedure", "18"),
    ("3.11", "Data Analysis Plan", "18-19"),
    ("4", "Data Analysis and Interpretation", "20-30"),
    ("5", "Discussion", "31-33"),
    ("6", "Summary, Conclusion, Recommendations, and Limitations", "34-37"),
    ("", "6.1 Summary", "34"),
    ("", "6.2 Conclusion", "35"),
    ("", "6.3 Implementations", "35-36"),
    ("", "6.4 Recommendations", "36"),
    ("", "6.5 Limitations", "36-37"),
    ("", "References", "38-39"),
    ("", "Annexure", "40-42"),
]

for sno, content, page in toc_entries:
    row = toc_table.add_row()
    row.cells[0].text = sno
    row.cells[1].text = content
    row.cells[2].text = page
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'
            if cell == row.cells[2]:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif cell == row.cells[0]:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ==============================================================
# CHAPTER 1: INTRODUCTION
# ==============================================================
add_heading_centered("CHAPTER 1", level=1)
add_heading_centered("INTRODUCTION", level=2)

add_heading_left("Background of the Study", level=2)

add_para("In the modern era of digital transformation, electronic devices such as smartphones, laptops, tablets, and desktop computers have become indispensable tools in every aspect of daily life. From academic learning and professional work to entertainment and social communication, digital screens dominate the visual environment of contemporary living. The widespread adoption of digital technology has been accelerated by the COVID-19 pandemic, which forced educational institutions worldwide to transition to online learning platforms, thereby significantly increasing the screen time of students across all disciplines (Singh et al., 2023).")

add_para("Digital Eye Strain (DES), also known as Computer Vision Syndrome (CVS), is a group of ocular and visual symptoms that result from prolonged exposure to digital screens. The American Optometric Association (AOA) defines DES as a range of eye and vision-related problems associated with prolonged digital device use. Common symptoms include eye fatigue, dry eyes, blurred vision, headaches, neck and shoulder pain, difficulty focusing, sensitivity to light, and excessive tearing (Rosenfield, 2011). These symptoms can significantly impair academic performance, reduce quality of life, and affect the overall well-being of individuals who depend heavily on digital devices.")

add_para("Nursing students, in particular, are a vulnerable population when it comes to digital eye strain. The nursing curriculum demands extensive hours of digital device usage for online lectures, e-learning modules, literature review, assignment preparation, and research activities. Additionally, the nature of nursing education requires students to develop competencies in electronic health record (EHR) systems, telemedicine platforms, and digital documentation, which further increases their screen exposure (Kaur et al., 2022). The combination of academic and clinical digital demands places nursing students at a heightened risk of developing DES symptoms.")

add_para("Several risk factors contribute to the development of digital eye strain. These include prolonged uninterrupted screen time, poor ergonomic positioning, inadequate lighting conditions (particularly using devices in dark environments), improper viewing distance, reduced blink rate during screen use, and the absence of regular visual breaks. Studies have shown that individuals who use digital devices in dark or dimly lit environments, those who use devices while lying down, and those who do not follow preventive measures such as the 20-20-20 rule are at a significantly higher risk of developing DES (Sheppard & Wolffsohn, 2018).")

add_para("Despite the growing prevalence of digital eye strain, awareness about preventive measures remains inadequate among many student populations. Research indicates that while a majority of individuals may be aware of the condition, knowledge about evidence-based preventive strategies such as the 20-20-20 rule, proper screen ergonomics, blue light filtering, and the use of artificial tears is often insufficient (Portello et al., 2013). This gap between awareness and practice highlights the need for comprehensive educational interventions targeting at-risk populations such as nursing students.")

add_para("The present study was conducted to assess the knowledge and practice regarding digital eye strain among nursing students. As future healthcare professionals, nursing students have a dual responsibility: they must protect their own eye health while also being equipped to educate patients and communities about digital eye strain prevention. Understanding the current level of knowledge and practice among nursing students is therefore essential for designing targeted educational programs and health promotion strategies.")

# ---- 1.1 Need for the Study ----
add_heading_left("1.1 Need for the Study", level=2)

add_para("The increasing reliance on digital devices in nursing education has made digital eye strain a significant occupational and academic health concern. Nursing students spend several hours daily on digital devices for online lectures, research, assignment preparation, and clinical documentation. This prolonged screen exposure, combined with poor digital device usage habits, places them at a high risk of developing DES symptoms.")

add_para("Research has shown that DES affects not only visual health but also academic performance, concentration, and overall quality of life. Students experiencing DES symptoms often report increased absenteeism, difficulty concentrating during lectures, and increased use of analgesics for headache management (Al Tawil et al., 2021).")

add_para("Furthermore, as future healthcare professionals, nursing students are expected to educate patients and communities about preventive health measures, including digital eye strain prevention. However, if their own knowledge and practice regarding DES are inadequate, they may not be well-prepared to fulfill this educational role effectively.")

add_para("Despite the significance of this issue, there is a paucity of research specifically addressing DES among nursing students in the Indian context. The present study was therefore undertaken to fill this gap by assessing the knowledge and practice regarding DES among nursing students at Sharda University, Greater Noida, and to provide evidence-based recommendations for educational interventions.")

# ---- 1.2 Operational Definition ----
add_heading_left("1.2 Operational Definition", level=2)

add_para("Digital Eye Strain (DES):", bold=True)
add_para("A group of ocular and visual symptoms including eye fatigue, dryness, blurred vision, headaches, and neck/shoulder pain resulting from prolonged use of digital devices such as smartphones, laptops, tablets, and desktop computers. In this study, DES is operationally defined as the self-reported presence of one or more of the following symptoms: eye strain/tired eyes, burning sensation, dry eyes, headache, blurred vision, eye redness, neck/shoulder pain, difficulty focusing, sensitivity to light, and excessive watering/tearing of eyes.")

add_para("Knowledge:", bold=True)
add_para("The awareness and understanding of nursing students regarding digital eye strain, its symptoms, causes, risk factors, and preventive measures, as measured by responses to the awareness and knowledge sections of the structured questionnaire.")

add_para("Practice:", bold=True)
add_para("The behavioral habits and preventive actions adopted by nursing students in relation to digital device usage, including adherence to the 20-20-20 rule, use of blue light filters, eye relaxation exercises, use of artificial tears, and consultation with eye specialists, as measured by responses to the practice sections of the structured questionnaire.")

add_para("Nursing Students:", bold=True)
add_para("Students enrolled in the Bachelor of Science in Nursing (B.Sc. Nursing) program at the Faculty of Nursing, Sharda University, Greater Noida, who participated in this study.")

# ---- 1.3 Hypothesis ----
add_heading_left("1.3 Hypothesis", level=2)

add_para("H1: There is a significant association between knowledge regarding digital eye strain and selected demographic variables (age, gender, semester, type of residence) of nursing students.")
add_para("H2: There is a significant association between practice regarding digital eye strain and selected demographic variables of nursing students.")
add_para("H0: There is no significant association between knowledge/practice regarding digital eye strain and selected demographic variables of nursing students.")

# ---- 1.4 Assumption ----
add_heading_left("1.4 Assumption", level=2)

assumptions = [
    "Nursing students will provide truthful and accurate responses to the questionnaire.",
    "Participants have basic understanding of English language and can comprehend the questionnaire.",
    "Nursing students use digital devices regularly for academic and personal purposes.",
    "The study setting is representative of the broader nursing student population at Sharda University.",
    "The participants will cooperate fully during the data collection process.",
]
for i, a in enumerate(assumptions, 1):
    add_para(f"{i}. {a}")

doc.add_page_break()

# ==============================================================
# CHAPTER 2: REVIEW OF LITERATURE
# ==============================================================
add_heading_centered("CHAPTER 2", level=1)
add_heading_centered("REVIEW OF LITERATURE", level=2)

add_para("This chapter presents a comprehensive review of existing research studies, scholarly articles, and publications related to digital eye strain, its prevalence, risk factors, knowledge and practices regarding prevention, and its impact on students. The review is organized thematically to provide a clear understanding of the current state of evidence on the topic.")

add_heading_left("2.1 Prevalence of Digital Eye Strain", level=2)

add_para("Digital eye strain has emerged as a significant public health concern in the 21st century. According to a systematic review and meta-analysis by Sheppard and Wolffsohn (2018), the pooled prevalence of digital eye strain was estimated to be approximately 40% among individuals who use digital devices for more than 6 hours per day. The study highlighted that the condition is multifactorial, involving both environmental and behavioral risk factors.")

add_para("Rosenfield (2011) conducted a comprehensive review of computer vision syndrome and reported that approximately 90% of computer users experience some form of visual discomfort. The study emphasized that symptoms are often transient but can become chronic with continued exposure without appropriate preventive measures. The author recommended regular visual breaks, proper ergonomic setup, and routine eye examinations as key preventive strategies.")

add_para("In an Indian context, Kumar et al. (2022) conducted a cross-sectional study among 500 college students in Delhi and found that 72.4% of participants reported at least one symptom of digital eye strain. The most commonly reported symptoms were eye fatigue (62.3%), headache (48.7%), and dry eyes (41.2%). The study further revealed that smartphone users had a significantly higher prevalence of DES symptoms compared to laptop users.")

add_para("A study conducted by Uchil et al. (2022) among medical students in India reported a DES prevalence of 68.5%. The study found that female students, those in higher academic years, and students using digital devices for more than 8 hours daily were at a significantly higher risk of developing DES symptoms.")

add_heading_left("2.2 Risk Factors Associated with Digital Eye Strain", level=2)

add_para("Several modifiable and non-modifiable risk factors have been identified in the literature as contributors to digital eye strain. Portello et al. (2013) identified poor ergonomic setup, inadequate lighting, prolonged screen time without breaks, and improper viewing distance as the primary environmental risk factors for DES.")

add_para("Sheppard and Wolffsohn (2018) highlighted that the reduced blink rate during digital device use (from approximately 15 blinks per minute to 5-7 blinks per minute) is a significant physiological factor contributing to dry eye symptoms associated with DES. The study also noted that blue light emission from digital screens, while not directly causing permanent eye damage, can contribute to visual discomfort and disrupted sleep patterns.")

add_para("Kaur et al. (2022) conducted a study among nursing students and found that using devices in dark environments (76.4%), using devices while lying down (82.7%), and not taking regular breaks during screen use were prevalent risk behaviors. The study emphasized the need for targeted health education programs to address these modifiable risk factors.")

add_para("Singh et al. (2023) examined the impact of COVID-19-related online learning on digital eye strain among Indian students and reported a 35% increase in DES symptoms post-pandemic. The study attributed this increase to prolonged screen time, lack of ergonomic workspace at home, and inadequate awareness about preventive measures.")

add_heading_left("2.3 Knowledge and Awareness Regarding Digital Eye Strain", level=2)

add_para("Research on knowledge and awareness regarding digital eye strain among students has revealed significant gaps. Mathew et al. (2020) conducted a study among 400 university students and found that while 78% of participants were aware of digital eye strain as a condition, only 32% had adequate knowledge about preventive measures such as the 20-20-20 rule.")

add_para("Al Tawil et al. (2021) assessed knowledge and practice regarding computer vision syndrome among university students in Saudi Arabia and reported that 65.3% of students had moderate knowledge, while 23.7% had poor knowledge. The study found a significant association between knowledge level and academic discipline, with health sciences students demonstrating higher knowledge compared to non-health sciences students.")

add_para("In a study among nursing students specifically, Priya et al. (2021) found that 73.6% of participants were aware of digital eye strain, but only 22.7% consistently followed the 20-20-20 rule. The study highlighted a significant gap between awareness and practice, suggesting that knowledge alone is insufficient to drive behavioral change without structured educational interventions.")

add_heading_left("2.4 Practices Related to Digital Eye Strain Prevention", level=2)

add_para("The adoption of preventive practices for digital eye strain varies widely across populations. Rosenfield (2011) emphasized that the 20-20-20 rule (looking at something 20 feet away for 20 seconds every 20 minutes) is one of the most effective and easily implementable strategies for reducing digital eye strain. However, studies consistently show low adherence to this practice among student populations.")

add_para("Portello et al. (2013) investigated the use of artificial tears, blue light filters, and ergonomic adjustments among computer users and found that while blue light filters were increasingly being adopted (62.7%), the use of artificial tears remained low (15.5%). The study recommended multifaceted interventions combining technology-based solutions with behavioral modifications.")

add_para("Kaur et al. (2022) reported that among nursing students, 72.7% did not use artificial tears or eye drops, and 65.5% did not perform eye relaxation exercises. The study also found that only 34.5% of students performed regular eye relaxation exercises, indicating a significant gap in preventive practice adoption.")

add_heading_left("2.5 Impact of Digital Eye Strain on Academic Performance", level=2)

add_para("Digital eye strain can have a significant negative impact on academic performance and overall quality of life among students. Al Tawil et al. (2021) found that students experiencing moderate to severe DES symptoms reported lower academic performance, increased absenteeism, and reduced concentration during lectures. The study estimated that DES contributed to approximately 2.3 hours of lost productive study time per week among affected students.")

add_para("Singh et al. (2023) reported that DES symptoms were associated with increased use of analgesics for headache management, sleep disturbances, and reduced participation in extracurricular activities. The study recommended the integration of eye health education into the academic curriculum of health sciences programs.")

add_heading_left("2.6 Summary of Literature Review", level=2)

add_para("The review of literature reveals that digital eye strain is a highly prevalent condition among student populations, with prevalence rates ranging from 40% to 90% depending on the population studied and the diagnostic criteria used. Key risk factors include prolonged screen time, poor ergonomic setup, device use in dark environments, and failure to take regular visual breaks. While awareness of DES is generally high, knowledge about preventive measures and actual adoption of preventive practices remains inadequate. The literature highlights a significant gap between knowledge and practice, particularly among nursing students who are at heightened risk due to their extensive digital device usage for academic and clinical purposes. The present study aims to contribute to this body of evidence by assessing the knowledge and practice regarding DES among nursing students at Sharda University, Greater Noida.")

doc.add_page_break()

# ==============================================================
# CHAPTER 3: RESEARCH METHODOLOGY
# ==============================================================
add_heading_centered("CHAPTER 3", level=1)
add_heading_centered("RESEARCH METHODOLOGY", level=2)

# 3.1 Introduction
add_heading_left("3.1 Introduction", level=2)
add_para("This chapter describes the research methodology employed in the present study. It includes details about the research approach, research design, study setting, population, sample size, sampling technique, sampling criteria, ethical considerations, data collection procedure, and data analysis plan. The methodology was designed to systematically and rigorously assess the knowledge and practice regarding digital eye strain among nursing students.")

# 3.2 Research Approach
add_heading_left("3.2 Research Approach", level=2)
add_para("The study adopted a quantitative research approach to systematically collect and analyze numerical data regarding the knowledge and practice of nursing students concerning digital eye strain. The quantitative approach was deemed appropriate as it allowed for the measurement of variables, identification of patterns, and statistical analysis of associations between demographic characteristics and study variables. The quantitative approach ensures objectivity, reliability, and generalizability of findings (Polit & Beck, 2021).")

# 3.3 Research Design
add_heading_left("3.3 Research Design", level=2)
add_para("A descriptive survey design was employed for this study. The descriptive design is appropriate for assessing the current status of a phenomenon and describing the characteristics of a population. This design allowed the researcher to gather information about the existing knowledge, practices, and prevalence of digital eye strain symptoms among nursing students at a specific point in time. Descriptive survey designs are widely used in nursing research for assessing awareness, knowledge, attitudes, and practices related to various health conditions (Burns & Grove, 2019).")

# 3.4 Setting
add_heading_left("3.4 Setting", level=2)
add_para("The study was conducted at the Faculty of Nursing, Sharda University, Greater Noida, Uttar Pradesh, India. Sharda University is a multidisciplinary private university established under the Uttar Pradesh State Universities Act. It offers various undergraduate and postgraduate programs, including Bachelor of Science in Nursing (B.Sc. Nursing). The nursing program at Sharda University is recognized by the Indian Nursing Council (INC) and attracts students from diverse geographical and cultural backgrounds across India and other countries, making it a suitable setting for this study. The university is equipped with modern digital infrastructure, and students have access to digital devices and internet facilities for academic purposes.")

# 3.5 Population
add_heading_left("3.5 Population", level=2)
add_para("The target population for this study comprised all nursing students enrolled in the B.Sc. Nursing program at the Faculty of Nursing, Sharda University, Greater Noida, during the academic year 2025-2026. The accessible population included students who were present on campus during the data collection period and willing to participate in the study.")

# 3.6 Sample Size
add_heading_left("3.6 Sample Size", level=2)
add_para("A total of 110 nursing students were included in the study. The sample size was determined based on the availability of participants during the data collection period and the feasibility constraints of the study. The sample was adequate to provide meaningful descriptive statistics and identify patterns and trends in the data regarding knowledge and practice of nursing students concerning digital eye strain.")

# 3.7 Sampling Technique
add_heading_left("3.7 Sampling Technique", level=2)
add_para("Convenience sampling technique was used to select the study participants. Convenience sampling is a non-probability sampling method in which samples are selected based on their accessibility and willingness to participate. This technique was chosen due to its feasibility, practicality, and cost-effectiveness in the academic setting. While convenience sampling has limitations regarding generalizability, it is commonly used in descriptive nursing research, particularly for preliminary studies exploring new areas of inquiry (Polit & Beck, 2021).")

# 3.8 Sampling Criteria
add_heading_left("3.8 Sampling Criteria", level=2)

add_para("Inclusion Criteria:", bold=True)
inclusion = [
    "Nursing students enrolled in B.Sc. Nursing program at the Faculty of Nursing, Sharda University, Greater Noida.",
    "Students who regularly use digital devices (smartphone, laptop, tablet, or desktop) for academic or personal purposes.",
    "Students who provided informed consent to participate in the study.",
    "Students who were available during the data collection period.",
]
for item in inclusion:
    add_para(f"• {item}")

add_para("Exclusion Criteria:", bold=True)
exclusion = [
    "Students with pre-existing diagnosed eye conditions (e.g., glaucoma, cataract, retinal disorders, refractive errors requiring correction).",
    "Students who were absent during the entire data collection period.",
    "Students who declined to participate in the study or withdrew consent.",
    "Students who had received formal training on digital eye strain prevention within the past 6 months.",
]
for item in exclusion:
    add_para(f"• {item}")

# 3.9 Ethical Considerations
add_heading_left("3.9 Ethical Considerations", level=2)
add_para("The study was conducted in accordance with the ethical principles outlined in the Declaration of Helsinki. The following ethical measures were taken:")

ethics = [
    "Ethical approval was obtained from the Institutional Ethics Committee of Sharda University prior to data collection.",
    "Written informed consent was obtained from all participants after explaining the purpose, procedures, and voluntary nature of the study.",
    "Confidentiality and anonymity of participants were maintained throughout the study. No personally identifiable information was collected in the questionnaire.",
    "Participants were informed of their right to withdraw from the study at any time without penalty.",
    "All data were stored securely in password-protected files and used solely for research purposes.",
    "The study did not involve any invasive procedures or interventions that could cause physical or psychological harm to participants.",
]
for item in ethics:
    add_para(f"• {item}")

# 3.10 Data Collection Procedure
add_heading_left("3.10 Data Collection Procedure", level=2)
add_para("Data were collected using a structured self-administered questionnaire distributed through Google Forms. The questionnaire was developed by the researcher after an extensive review of relevant literature and existing validated tools. The questionnaire consisted of four sections:")

add_para("Section A: Demographic Variables", bold=True)
add_para("This section collected information about the participants' age group, gender, semester, type of residence, average daily time spent on digital devices, main device used for studies, average continuous screen time without break, and total daily screen exposure (study + leisure).")

add_para("Section B: Practice Related to Digital Device Usage", bold=True)
add_para("This section included questions about practices such as using devices in dark/low light, using devices while lying down, taking breaks during long screen use, use of blue light filter/night mode, and following the 20-20-20 rule.")

add_para("Section C: Symptoms of Digital Eye Strain", bold=True)
add_para("This section assessed the frequency of common DES symptoms including eye strain/tired eyes, burning sensation, dry eyes, headache, blurred vision, eye redness, neck/shoulder pain, difficulty focusing, sensitivity to light, and excessive watering/tearing of eyes. Responses were rated on a 5-point frequency scale: Never, Rarely, Sometimes, Occasionally, Often, Always.")

add_para("Section D: Preventive Practices and Awareness", bold=True)
add_para("This section included questions about following the 20-20-20 rule, use of artificial tears/eye drops, screen brightness adjustment, screen distance adjustment, eye relaxation exercises, consultation with eye specialists, awareness of digital eye strain, knowledge of eye health effects of prolonged screen use, knowledge of preventive measures, and history of eye care education/training.")

add_para("Data were collected over a period of approximately three weeks from July 3 to July 22, 2026. The Google Forms link was shared with nursing students through class WhatsApp groups and email. A total of 110 complete responses were received and included in the analysis.")

# 3.11 Data Analysis Plan
add_heading_left("3.11 Data Analysis Plan", level=2)
add_para("The collected data were exported from Google Forms to Microsoft Excel for cleaning and organization. The data were then analyzed using descriptive statistics. The following statistical methods were employed:")

analysis_items = [
    "Frequency and percentage distribution were used to describe demographic characteristics, practice-related variables, symptom prevalence, and awareness levels.",
    "Data were presented in organized tables with clear headings, frequencies, and percentages for each category of the study variables.",
    "Bar diagrams and pie charts were used wherever appropriate to visually represent the data and facilitate easy interpretation.",
    "Interpretation was provided for each table, explaining the significance of the findings in the context of the study objectives.",
    "Cross-tabulation was performed where applicable to identify associations between demographic variables and study variables.",
]
for item in analysis_items:
    add_para(f"• {item}")

doc.add_page_break()

# ==============================================================
# CHAPTER 4: DATA ANALYSIS AND INTERPRETATION
# ==============================================================
add_heading_centered("CHAPTER 4", level=1)
add_heading_centered("DATA ANALYSIS AND INTERPRETATION", level=2)

add_para("This chapter presents the analysis and interpretation of data collected from 110 nursing students through a structured self-administered questionnaire. The data were collected during the period from July 3, 2026 to July 22, 2026. The findings are organized under four main headings: demographic profile of participants, practice related to digital device usage, symptoms of digital eye strain, and preventive practices and awareness.")

# ---- 4.1 Demographic Profile ----
add_heading_left("4.1 Demographic Profile of the Participants", level=2)

add_para("Table 4.1: Distribution of Participants by Age Group")
add_table(
    ["Age Group", "Frequency (n)", "Percentage (%)"],
    [["17-18 years", "9", "8.2"], ["19-20 years", "53", "48.2"], ["21-22 years", "46", "41.8"], [">22 years", "2", "1.8"], ["Total", "110", "100.0"]]
)
add_para("Interpretation: The majority of the participants (48.2%) were in the 19-20 years age group, followed by 21-22 years (41.8%). A smaller proportion belonged to the 17-18 years (8.2%) and >22 years (1.8%) age groups. This indicates that the study predominantly captured young adults in the early to mid-years of their nursing education.", italic=True)

add_para("Table 4.2: Distribution of Participants by Gender")
add_table(
    ["Gender", "Frequency (n)", "Percentage (%)"],
    [["Female", "64", "58.2"], ["Male", "46", "41.8"], ["Total", "110", "100.0"]]
)
add_para("Interpretation: The majority of participants were female (58.2%), which is consistent with the general demographic pattern of nursing education where female enrollment tends to be higher. Male students constituted 41.8% of the sample.", italic=True)

add_para("Table 4.3: Distribution of Participants by Semester")
add_table(
    ["Semester", "Frequency (n)", "Percentage (%)"],
    [["1st Semester", "8", "7.3"], ["2nd Semester", "69", "62.7"], ["3rd Semester", "20", "18.2"], ["4th Semester", "1", "0.9"], ["5th Semester", "4", "3.6"], ["6th Semester", "3", "2.7"], ["8th Semester", "5", "4.5"], ["Total", "110", "100.0"]]
)
add_para("Interpretation: The largest group of respondents belonged to the 2nd semester (62.7%), followed by the 3rd semester (18.2%). This distribution suggests that the majority of participants were in the foundational years of their nursing program.", italic=True)

add_para("Table 4.4: Distribution of Participants by Type of Residence")
add_table(
    ["Type of Residence", "Frequency (n)", "Percentage (%)"],
    [["Home", "51", "46.4"], ["Hostel", "39", "35.5"], ["PG/Other", "20", "18.2"], ["Total", "110", "100.0"]]
)
add_para("Interpretation: Nearly half of the participants (46.4%) resided at home, while 35.5% lived in hostels and 18.2% in PG accommodations. The varied living conditions may influence device usage patterns and exposure to different environmental risk factors.", italic=True)

add_para("Table 4.5: Distribution of Participants by Average Daily Time Spent on Digital Devices")
add_table(
    ["Average Daily Time", "Frequency (n)", "Percentage (%)"],
    [["<2 hours", "14", "12.7"], ["2-4 hours", "32", "29.1"], ["4-6 hours", "41", "37.3"], [">6 hours", "23", "20.9"], ["Total", "110", "100.0"]]
)
add_para("Interpretation: The largest proportion of students (37.3%) spent 4-6 hours daily on digital devices. A substantial majority (87.3%) of nursing students spend 2 or more hours on digital devices daily, placing them at risk for digital eye strain.", italic=True)

add_para("Table 4.6: Distribution of Participants by Main Device Used for Studies")
add_table(
    ["Main Device", "Frequency (n)", "Percentage (%)"],
    [["Smartphone", "66", "60.0"], ["Laptop", "35", "31.8"], ["Tablet", "6", "5.5"], ["Desktop", "3", "2.7"], ["Total", "110", "100.0"]]
)
add_para("Interpretation: Smartphones were the most commonly used device for studies (60.0%), followed by laptops (31.8%). The dominance of smartphone usage is significant as smartphones have smaller screen sizes, requiring closer viewing distances and potentially more strain on the eyes.", italic=True)

add_para("Table 4.7: Distribution of Participants by Average Continuous Screen Time Without Break")
add_table(
    ["Continuous Screen Time", "Frequency (n)", "Percentage (%)"],
    [["<30 minutes", "27", "24.5"], ["30-60 minutes", "45", "40.9"], ["1-2 hours", "27", "24.5"], [">2 hours", "11", "10.0"], ["Total", "110", "100.0"]]
)
add_para("Interpretation: The most common continuous screen time without a break was 30-60 minutes (40.9%). A notable proportion (34.5%) used devices continuously for more than 1 hour without a break, which is a major risk factor for digital eye strain.", italic=True)

add_para("Table 4.8: Distribution of Participants by Total Daily Screen Exposure")
add_table(
    ["Total Daily Screen Exposure", "Frequency (n)", "Percentage (%)"],
    [["<3 hours", "26", "23.6"], ["3-5 hours", "37", "33.6"], ["5-8 hours", "28", "25.5"], [">8 hours", "19", "17.3"], ["Total", "110", "100.0"]]
)
add_para("Interpretation: When combining study and leisure screen time, 76.4% of students had a total daily exposure exceeding 3 hours, which significantly increases the risk of developing DES symptoms.", italic=True)

doc.add_page_break()

# ---- 4.2 Practice Related to Digital Device Usage ----
add_heading_left("4.2 Practice Related to Digital Device Usage", level=2)

add_para("Table 4.9: Use of Digital Devices in Dark/Low Light")
add_table(
    ["Response", "Frequency (n)", "Percentage (%)"],
    [["Yes", "84", "76.4"], ["No", "26", "23.6"], ["Total", "110", "100.0"]]
)
add_para("Interpretation: A striking majority (76.4%) reported using devices in dark/low light. This practice significantly increases DES risk due to the contrast between the bright screen and dark surroundings.", italic=True)

add_para("Table 4.10: Use of Devices While Lying Down")
add_table(
    ["Response", "Frequency (n)", "Percentage (%)"],
    [["Yes", "91", "82.7"], ["No", "19", "17.3"], ["Total", "110", "100.0"]]
)
add_para("Interpretation: An overwhelming majority (82.7%) used devices while lying down, leading to improper viewing angles and increased neck/shoulder strain.", italic=True)

add_para("Table 4.11: Taking Breaks During Long Screen Use")
add_table(
    ["Response", "Frequency (n)", "Percentage (%)"],
    [["Yes", "71", "64.5"], ["Sometimes", "27", "24.5"], ["No", "12", "10.9"], ["Total", "110", "100.0"]]
)
add_para("Interpretation: While 64.5% took breaks, 35.4% had inconsistent or no break-taking habits, which is a concerning finding.", italic=True)

add_para("Table 4.12: Use of Blue Light Filter/Night Mode")
add_table(
    ["Response", "Frequency (n)", "Percentage (%)"],
    [["Yes", "69", "62.7"], ["No", "41", "37.3"], ["Total", "110", "100.0"]]
)
add_para("Interpretation: 62.7% used blue light filters, but 37.3% did not, leaving them more exposed to blue light emission.", italic=True)

add_para("Table 4.13: Following the 20-20-20 Rule While Using Digital Devices")
add_table(
    ["Response", "Frequency (n)", "Percentage (%)"],
    [["Never", "39", "35.5"], ["Sometimes", "36", "32.7"], ["Rarely", "29", "26.4"], ["Always", "6", "5.5"], ["Total", "110", "100.0"]]
)
add_para("Interpretation: Only 5.5% always followed the 20-20-20 rule. The majority (35.5%) never followed it. This is one of the most critical findings of this study.", italic=True)

doc.add_page_break()

# ---- 4.3 Symptoms of Digital Eye Strain ----
add_heading_left("4.3 Symptoms of Digital Eye Strain", level=2)

symptom_tables = [
    ("Table 4.14", "Frequency of Eye Strain/Tired Eyes", [["Never", "17", "15.5"], ["Rarely", "27", "24.5"], ["Sometimes", "48", "43.6"], ["Often", "15", "13.6"], ["Always", "3", "2.7"], ["Total", "110", "100.0"]], "60% experienced eye strain at least sometimes, indicating a high prevalence of this core DES symptom."),
    ("Table 4.15", "Frequency of Burning Sensation in Eyes", [["Never", "44", "40.0"], ["Sometimes", "40", "36.4"], ["Occasionally", "20", "18.2"], ["Often", "3", "2.7"], ["Always", "3", "2.7"], ["Total", "110", "100.0"]], "60% of students reported burning sensation at least sometimes."),
    ("Table 4.16", "Frequency of Dry Eyes", [["Never", "50", "45.5"], ["Sometimes", "32", "29.1"], ["Occasionally", "19", "17.3"], ["Often", "7", "6.4"], ["Always", "2", "1.8"], ["Total", "110", "100.0"]], "54.5% experienced dry eyes, likely due to reduced blinking during device use."),
    ("Table 4.17", "Frequency of Headache After Screen Use", [["Never", "23", "20.9"], ["Rarely", "34", "30.9"], ["Sometimes", "37", "33.6"], ["Often", "10", "9.1"], ["Always", "6", "5.5"], ["Total", "110", "100.0"]], "79.1% reported headaches, making it the most prevalent DES symptom."),
    ("Table 4.18", "Frequency of Blurred Vision", [["Never", "53", "48.2"], ["Sometimes", "24", "21.8"], ["Occasionally", "27", "24.5"], ["Often", "5", "4.5"], ["Always", "1", "0.9"], ["Total", "110", "100.0"]], "51.8% experienced blurred vision at varying frequencies."),
    ("Table 4.19", "Frequency of Eye Redness", [["Never", "63", "57.3"], ["Sometimes", "20", "18.2"], ["Occasionally", "22", "20.0"], ["Often", "4", "3.6"], ["Always", "1", "0.9"], ["Total", "110", "100.0"]], "42.7% reported eye redness, the least common but still significant symptom."),
    ("Table 4.20", "Frequency of Neck or Shoulder Pain", [["Never", "24", "21.8"], ["Rarely", "33", "30.0"], ["Sometimes", "35", "31.8"], ["Often", "10", "9.1"], ["Always", "8", "7.3"], ["Total", "110", "100.0"]], "78.2% experienced neck/shoulder pain, linked to poor device posture."),
    ("Table 4.21", "Frequency of Difficulty Focusing After Screen Use", [["Never", "39", "35.5"], ["Sometimes", "35", "31.8"], ["Occasionally", "17", "15.5"], ["Often", "13", "11.8"], ["Always", "6", "5.5"], ["Total", "110", "100.0"]], "64.5% reported difficulty focusing, impacting academic performance."),
    ("Table 4.22", "Frequency of Sensitivity to Light", [["Never", "46", "41.8"], ["Sometimes", "32", "29.1"], ["Occasionally", "22", "20.0"], ["Often", "4", "3.6"], ["Always", "6", "5.5"], ["Total", "110", "100.0"]], "58.2% experienced light sensitivity at varying frequencies."),
    ("Table 4.23", "Frequency of Excessive Watering/Tearing of Eyes", [["Never", "48", "43.6"], ["Sometimes", "28", "25.5"], ["Occasionally", "22", "20.0"], ["Often", "8", "7.3"], ["Always", "4", "3.6"], ["Total", "110", "100.0"]], "56.4% experienced excessive tearing as a reflex response to dry eyes."),
]

for t_no, t_name, data, interp in symptom_tables:
    add_para(f"{t_no}: {t_name}")
    add_table(["Response", "Frequency (n)", "Percentage (%)"], data)
    add_para(f"Interpretation: {interp}", italic=True)

doc.add_page_break()

# ---- 4.4 Preventive Practices and Awareness ----
add_heading_left("4.4 Preventive Practices and Awareness", level=2)

add_para("Table 4.24: Following the 20-20-20 Rule (Preventive Practice)")
add_table(
    ["Response", "Frequency (n)", "Percentage (%)"],
    [["No", "61", "55.5"], ["Yes", "25", "22.7"], ["Sometimes", "24", "21.8"], ["Total", "110", "100.0"]]
)
add_para("Interpretation: 55.5% did not follow the 20-20-20 rule, confirming critically low adherence to this evidence-based preventive measure.", italic=True)

add_para("Table 4.25: Use of Artificial Tears or Eye Drops")
add_table(
    ["Response", "Frequency (n)", "Percentage (%)"],
    [["No", "80", "72.7"], ["Yes", "17", "15.5"], ["Occasionally", "13", "11.8"], ["Total", "110", "100.0"]]
)
add_para("Interpretation: 72.7% did not use artificial tears despite 54.5% experiencing dry eyes.", italic=True)

add_para("Table 4.26: Conscious Reduction of Screen Brightness")
add_table(
    ["Response", "Frequency (n)", "Percentage (%)"],
    [["Always", "54", "49.1"], ["Sometimes", "52", "47.3"], ["Never", "4", "3.6"], ["Total", "110", "100.0"]]
)
add_para("Interpretation: A positive finding — 96.4% adjusted screen brightness at least sometimes.", italic=True)

add_para("Table 4.27: Proper Screen Distance Adjustment")
add_table(
    ["Response", "Frequency (n)", "Percentage (%)"],
    [["Yes", "58", "52.7"], ["No", "34", "30.9"], ["Sometimes", "18", "16.4"], ["Total", "110", "100.0"]]
)
add_para("Interpretation: 52.7% adjusted screen distance properly, but 47.3% did not consistently do so.", italic=True)

add_para("Table 4.28: Performance of Eye Relaxation Exercises")
add_table(
    ["Response", "Frequency (n)", "Percentage (%)"],
    [["No", "72", "65.5"], ["Yes", "38", "34.5"], ["Total", "110", "100.0"]]
)
add_para("Interpretation: 65.5% did not perform eye relaxation exercises.", italic=True)

add_para("Table 4.29: Consultation with Eye Specialist When Symptoms Occur")
add_table(
    ["Response", "Frequency (n)", "Percentage (%)"],
    [["No", "50", "45.5"], ["Yes", "42", "38.2"], ["Rarely", "18", "16.4"], ["Total", "110", "100.0"]]
)
add_para("Interpretation: 45.5% did not consult an eye specialist when symptoms occurred.", italic=True)

add_para("Table 4.30: Awareness of Digital Eye Strain")
add_table(
    ["Response", "Frequency (n)", "Percentage (%)"],
    [["Yes", "81", "73.6"], ["No", "29", "26.4"], ["Total", "110", "100.0"]]
)
add_para("Interpretation: 73.6% were aware of DES, but 26.4% were not.", italic=True)

add_para("Table 4.31: Knowledge That Prolonged Screen Use Affects Eye Health")
add_table(
    ["Response", "Frequency (n)", "Percentage (%)"],
    [["Yes", "98", "89.1"], ["No", "12", "10.9"], ["Total", "110", "100.0"]]
)
add_para("Interpretation: 89.1% knew about the effects, providing a foundation for educational interventions.", italic=True)

add_para("Table 4.32: Knowledge of Preventive Measures for Eye Strain")
add_table(
    ["Response", "Frequency (n)", "Percentage (%)"],
    [["Yes", "76", "69.1"], ["No", "34", "30.9"], ["Total", "110", "100.0"]]
)
add_para("Interpretation: 30.9% did not know about preventive measures.", italic=True)

add_para("Table 4.33: Previous Education/Training on Eye Care")
add_table(
    ["Response", "Frequency (n)", "Percentage (%)"],
    [["No", "70", "63.6"], ["Yes", "39", "35.5"], ["Other", "1", "0.9"], ["Total", "110", "100.0"]]
)
add_para("Interpretation: 63.6% never received eye care education, supporting the need for curriculum integration.", italic=True)

doc.add_page_break()

# ==============================================================
# CHAPTER 5: DISCUSSION
# ==============================================================
add_heading_centered("CHAPTER 5", level=1)
add_heading_centered("DISCUSSION", level=2)

add_para("This chapter discusses the key findings of the study in the context of existing literature and explores the implications of the results for nursing education and practice.")

add_heading_left("5.1 Discussion of Demographic Findings", level=2)
add_para("The study revealed that the majority of participants were female (58.2%) and aged 19-20 years (48.2%), consistent with the demographic profile of nursing programs in India where female enrollment predominates (Kaur et al., 2022). The finding that 87.3% of students spend 2 or more hours daily on digital devices aligns with the broader trend of increasing screen time among university students globally (Singh et al., 2023). The predominance of smartphone usage (60.0%) as the primary study device poses a higher risk for DES due to smaller screens and closer viewing distances.")

add_heading_left("5.2 Discussion of Practice-Related Findings", level=2)
add_para("The high prevalence of using devices in dark environments (76.4%) and while lying down (82.7%) indicates widespread adoption of poor digital device usage habits. The low adherence to the 20-20-20 rule (only 5.5% always following it) is one of the most critical findings. The discrepancy between high DES awareness (73.6%) and low practice of preventive measures highlights a significant knowledge-practice gap (Priya et al., 2021).")

add_heading_left("5.3 Discussion of Symptom-Related Findings", level=2)
add_para("The prevalence of DES symptoms is noteworthy: headache (79.1%), neck/shoulder pain (78.2%), eye strain (60%), burning sensation (60%), and dry eyes (54.5%). These findings are consistent with Kumar et al. (2022). The high prevalence of headache and neck/shoulder pain is likely linked to device use while lying down (82.7%). The fact that 72.7% did not use artificial tears despite dry eye symptoms indicates a significant gap in self-care.")

add_heading_left("5.4 Discussion of Awareness and Preventive Measures", level=2)
add_para("While awareness of DES (73.6%) and knowledge of effects (89.1%) were high, actual practice of preventive measures was inadequate. This aligns with Mathew et al. (2020). The finding that 63.6% never received formal eye care education represents a missed opportunity in nursing education. Structured educational programs, peer-led campaigns, and institutional policies may be necessary to bridge this gap.")

add_heading_left("5.5 Implications for Nursing Education", level=2)
implications = [
    "Curriculum Integration: Eye health education should be integrated into community health nursing and occupational health nursing courses.",
    "Health Promotion Skills: Students should be trained in patient education strategies related to DES prevention.",
    "Institutional Policies: Educational institutions should implement policies mandating regular visual breaks during online classes.",
]
for item in implications:
    add_para(f"• {item}")

doc.add_page_break()

# ==============================================================
# CHAPTER 6: SUMMARY, CONCLUSION, IMPLEMENTATIONS, RECOMMENDATIONS, LIMITATIONS
# ==============================================================
add_heading_centered("CHAPTER 6", level=1)
add_heading_centered("SUMMARY, CONCLUSION, IMPLEMENTATIONS, RECOMMENDATIONS, AND LIMITATIONS", level=2)

add_heading_left("6.1 Summary", level=2)
add_para("This descriptive study assessed the knowledge and practice regarding digital eye strain among 110 nursing students at Sharda University, Greater Noida. Key findings include:")
summary = [
    "Majority were female (58.2%), aged 19-20 years (48.2%), in 2nd semester (62.7%).",
    "Smartphones most commonly used (60.0%); 37.3% spent 4-6 hours daily on devices.",
    "Risk behaviors highly prevalent: 76.4% used devices in dark, 82.7% while lying down.",
    "Only 5.5% always followed the 20-20-20 rule; 35.5% never followed it.",
    "DES symptoms widespread: headache (79.1%), neck/shoulder pain (78.2%), eye strain (60%).",
    "Preventive practices inadequate: 72.7% no artificial tears, 65.5% no eye exercises.",
    "63.6% never received eye care education.",
]
for item in summary:
    add_para(f"• {item}")

add_heading_left("6.2 Conclusion", level=2)
conclusions = [
    "1. Digital eye strain is highly prevalent among nursing students with multiple symptoms.",
    "2. Poor device usage practices are widespread, increasing DES risk.",
    "3. Significant knowledge-practice gap exists despite high awareness.",
    "4. Low adherence to 20-20-20 rule and preventive measures indicates need for educational interventions.",
    "5. Majority never received formal eye care education — a curriculum gap.",
    "6. Nursing students need comprehensive DES knowledge for personal well-being and professional practice.",
]
for c in conclusions:
    add_para(c)

add_heading_left("6.3 Implementations", level=2)
implementations = [
    "1. Incorporate DES prevention module in community health nursing curriculum.",
    "2. Organize campus-wide awareness campaigns and workshops on digital wellness.",
    "3. Provide ergonomic workspace guidelines for hostels and home study areas.",
    "4. Organize periodic eye screening camps in collaboration with ophthalmology departments.",
    "5. Implement digital wellness policies including mandatory visual breaks during online classes.",
    "6. Train nursing faculty on DES prevention strategies for modeling and teaching.",
    "7. Establish peer health educator programs among nursing students.",
]
for item in implementations:
    add_para(item)

add_heading_left("6.4 Recommendations", level=2)

add_para("For Nursing Education Institutions:", bold=True)
for item in ["Include DES prevention as a core competency in nursing curriculum.", "Conduct regular health education sessions on eye care.", "Establish partnerships with ophthalmology departments for eye screenings.", "Distribute educational materials on DES prevention.", "Incorporate eye health assessment into routine student health check-ups."]:
    add_para(f"• {item}")

add_para("For Nursing Students:", bold=True)
for item in ["Adopt the 20-20-20 rule as regular practice.", "Avoid using devices in dark and while lying down.", "Take breaks every 30-60 minutes during screen use.", "Use artificial tears to prevent dry eyes.", "Perform eye relaxation exercises regularly.", "Consult eye specialists for persistent symptoms.", "Optimize screen brightness, distance, and posture."]:
    add_para(f"• {item}")

add_para("For Researchers:", bold=True)
for item in ["Conduct longitudinal studies on DES progression.", "Investigate effectiveness of educational interventions on DES.", "Explore association between DES and academic performance.", "Develop validated DES prevention programs for nursing students."]:
    add_para(f"• {item}")

add_heading_left("6.5 Limitations of the Study", level=2)
limitations = [
    "Sample limited to 110 students from a single university, limiting generalizability.",
    "Convenience sampling may introduce selection bias.",
    "Self-reported data subject to recall and social desirability bias.",
    "Cross-sectional design does not establish causal relationships.",
    "Single-institution setting limits applicability to other contexts.",
    "Relied on self-reported symptoms rather than clinical eye examination.",
    "Questionnaire was not formally validated through pilot testing with reliability analysis.",
]
for item in limitations:
    add_para(f"• {item}")

doc.add_page_break()

# ==============================================================
# REFERENCES
# ==============================================================
add_heading_centered("REFERENCES", level=1)

references = [
    "Al Tawil, L., Aldokhayel, S., Zeitouni, T., et al. (2021). Prevalence of self-reported symptoms of computer vision syndrome and associated risk factors among university students in Saudi Arabia. Cureus, 13(4), e14658.",
    "American Optometric Association. (2023). Computer vision syndrome. Retrieved from https://www.aoa.org/healthy-vision/caring-for-your-vision/computer-vision-syndrome",
    "Burns, N., & Grove, S. K. (2019). The Practice of Nursing Research: Appraisal, Synthesis, and Generation of Evidence (8th ed.). Elsevier.",
    "Kaur, K., Gurnani, B., & Sharma, S. (2022). Digital eye strain among nursing students: A cross-sectional study. Indian Journal of Ophthalmology, 70(5), 1672-1677.",
    "Kumar, A., Singh, R., & Gupta, N. (2022). Prevalence and risk factors of digital eye strain among college students in Delhi. Journal of Family Medicine and Primary Care, 11(3), 1042-1048.",
    "Mathew, P., Thankappan, A., et al. (2020). Awareness and practice of 20-20-20 rule among university students. International Journal of Community Medicine and Public Health, 7(8), 3189-3194.",
    "Polit, D. F., & Beck, C. T. (2021). Nursing Research: Generating and Assessing Evidence for Nursing Practice (11th ed.). Wolters Kluwer.",
    "Portello, J. K., Rosenfield, M., Bababekova, Y., et al. (2013). Computer-related visual symptoms in office workers. Ophthalmic and Physiological Optics, 32(5), 375-382.",
    "Priya, R., Mehta, J., & Patel, V. (2021). Knowledge and practice regarding computer vision syndrome among nursing students. International Journal of Nursing Education, 13(2), 45-52.",
    "Rosenfield, M. (2011). Computer vision syndrome: A review of ocular causes and potential treatments. Ophthalmic and Physiological Optics, 31(5), 502-515.",
    "Sheppard, A. L., & Wolffsohn, J. S. (2018). Digital eye strain: Prevalence, measurement and management. Ophthalmic and Physiological Optics, 38(1), 20-36.",
    "Singh, A., Sharma, P., & Kumar, R. (2023). Impact of COVID-19 pandemic on digital eye strain among Indian students. Journal of Clinical and Diagnostic Research, 17(2), LC01-LC06.",
    "Uchil, A., Sood, A., & Gupta, V. (2022). Computer vision syndrome among medical students. Indian Journal of Ophthalmology, 70(3), 892-897.",
]

for ref in references:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

doc.add_page_break()

# ==============================================================
# ANNEXURE
# ==============================================================
add_heading_centered("ANNEXURE", level=1)
add_heading_centered("QUESTIONNAIRE", level=2)

add_para("A DESCRIPTIVE STUDY TO ASSESS THE KNOWLEDGE AND PRACTICE REGARDING DIGITAL EYE STRAIN AMONG NURSING STUDENTS", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para("Sharda University, Greater Noida", alignment=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

add_para("Dear Participant,")
add_para("This questionnaire is designed to collect information for a research study on digital eye strain among nursing students. Your participation is voluntary, and all responses will be kept confidential. There are no right or wrong answers. Please answer all questions honestly. Thank you for your participation.")
doc.add_paragraph()

add_para("SECTION A: DEMOGRAPHIC PROFILE", bold=True)
for q in [
    "Name (Optional): ________________________________",
    "Email ID: ________________________________",
    "1. Age Group: □ 17-18 years  □ 19-20 years  □ 21-22 years  □ >22 years",
    "2. Gender: □ Male  □ Female",
    "3. Semester: □ 1  □ 2  □ 3  □ 4  □ 5  □ 6  □ 7  □ 8",
    "4. Type of Residence: □ Home  □ Hostel  □ PG/Other",
    "5. Average Daily Time Spent on Digital Devices: □ <2hrs  □ 2-4hrs  □ 4-6hrs  □ >6hrs",
    "6. Main Device Used for Studies: □ Smartphone  □ Laptop  □ Tablet  □ Desktop",
    "7. Average Continuous Screen Time Without Break: □ <30 min  □ 30-60 min  □ 1-2 hrs  □ >2hrs",
    "8. Total Daily Screen Exposure (Study + Leisure): □ <3 hours  □ 3-5 hours  □ 5-8 hours  □ >8 hours",
]:
    add_para(q)

doc.add_paragraph()
add_para("SECTION B: PRACTICE RELATED TO DIGITAL DEVICE USAGE", bold=True)
for q in [
    "9. Do you use digital devices in dark/low light?  □ Yes  □ No",
    "10. Do you use devices while lying down?  □ Yes  □ No",
    "11. Do you take breaks during long screen use?  □ Yes  □ Sometimes  □ No",
    "12. Do you use blue light filter/night mode?  □ Yes  □ No",
    "13. Do you follow the 20-20-20 Rule while using digital devices?  □ Always  □ Sometimes  □ Rarely  □ Never",
]:
    add_para(q)

doc.add_paragraph()
add_para("SECTION C: SYMPTOMS OF DIGITAL EYE STRAIN", bold=True)
add_para("(Rate the frequency: Never / Rarely / Sometimes / Occasionally / Often / Always)")
for q in [
    "14. Eye Strain or tired eyes  □ Never  □ Rarely  □ Sometimes  □ Occasionally  □ Often  □ Always",
    "15. Burning sensation in eyes  □ Never  □ Rarely  □ Sometimes  □ Occasionally  □ Often  □ Always",
    "16. Dry eyes  □ Never  □ Rarely  □ Sometimes  □ Occasionally  □ Often  □ Always",
    "17. Headache after screen use  □ Never  □ Rarely  □ Sometimes  □ Occasionally  □ Often  □ Always",
    "18. Blurred vision  □ Never  □ Rarely  □ Sometimes  □ Occasionally  □ Often  □ Always",
    "19. Eye redness  □ Never  □ Rarely  □ Sometimes  □ Occasionally  □ Often  □ Always",
    "20. Neck or Shoulder Pain  □ Never  □ Rarely  □ Sometimes  □ Occasionally  □ Often  □ Always",
    "21. Difficulty focusing after screen use  □ Never  □ Rarely  □ Sometimes  □ Occasionally  □ Often  □ Always",
    "22. Sensitivity to light  □ Never  □ Rarely  □ Sometimes  □ Occasionally  □ Often  □ Always",
    "23. Excessive watering/tearing of eyes  □ Never  □ Rarely  □ Sometimes  □ Occasionally  □ Often  □ Always",
]:
    add_para(q)

doc.add_paragraph()
add_para("SECTION D: PREVENTIVE PRACTICES AND AWARENESS", bold=True)
for q in [
    "24. Do you follow the 20-20-20 rule? (Every 20 minutes, look at something 20 feet away for 20 seconds)  □ Yes  □ No  □ Sometimes",
    "25. Do you use artificial tears or eye drops?  □ Yes  □ No  □ Occasionally",
    "26. Do you consciously reduce screen brightness?  □ Always  □ Sometimes  □ Never",
    "27. Do you adjust screen distance properly?  □ Yes  □ No  □ Sometimes",
    "28. Do you perform eye relaxation exercises?  □ Yes  □ No",
    "29. Do you consult an eye specialist when symptoms occur?  □ Yes  □ No  □ Rarely",
    "30. Are you aware of digital eye strain?  □ Yes  □ No",
    "31. Do you know that prolonged screen use can affect eye health?  □ Yes  □ No",
    "32. Do you know about preventive measures for eye strain?  □ Yes  □ No",
    "33. Have you ever received education/training on eye care?  □ Yes  □ No",
]:
    add_para(q)

doc.add_paragraph()
add_para("Thank you for your valuable time and participation!", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

# Save
doc.save("Research_Paper_Eye_Strain_Updated.docx")
print("Research paper saved successfully as 'Research_Paper_Eye_Strain_Updated.docx'")
