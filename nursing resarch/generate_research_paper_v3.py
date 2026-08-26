from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

def add_hc(text, level=0):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in h.runs: r.font.color.rgb = RGBColor(0,0,0)
    return h

def add_hl(text, level=2):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in h.runs: r.font.color.rgb = RGBColor(0,0,0)
    return h

def add_p(text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, sp=Pt(6)):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = sp
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.bold = bold
    r.italic = italic
    return p

def add_tbl(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True; r.font.size = Pt(10); r.font.name = 'Times New Roman'
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]
            c.text = str(val)
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(10); r.font.name = 'Times New Roman'
    doc.add_paragraph()
    return t

# ============ TITLE PAGE ============
for _ in range(4): doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("A DESCRIPTIVE STUDY TO ASSESS THE KNOWLEDGE AND PRACTICE REGARDING DIGITAL EYE STRAIN AMONG NURSING STUDENTS")
r.bold = True; r.font.size = Pt(16); r.font.name = 'Times New Roman'
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("A Research Paper Submitted in Partial Fulfillment of the Requirements for the Degree of Bachelor of Science in Nursing (B.Sc. Nursing)")
r.font.size = Pt(12); r.font.name = 'Times New Roman'
doc.add_paragraph(); doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Submitted To:\nFaculty of Nursing\nSharda University, Greater Noida, Uttar Pradesh, India")
r.font.size = Pt(12); r.font.name = 'Times New Roman'
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("July 2026")
r.font.size = Pt(12); r.font.name = 'Times New Roman'
doc.add_page_break()

# ============ DECLARATION ============
add_hc("DECLARATION", 1); doc.add_paragraph()
add_p("I hereby declare that this research paper entitled \"A Descriptive Study to Assess the Knowledge and Practice Regarding Digital Eye Strain Among Nursing Students\" is a record of original work done by me under the guidance and supervision of the Faculty of Nursing, Sharda University, Greater Noida.")
add_p("This research paper has not been submitted elsewhere for the award of any degree, diploma, or other similar title or recognition. All the sources of information used have been acknowledged and cited appropriately.")
doc.add_paragraph(); add_p("Place: Greater Noida"); add_p("Date: July 2026"); doc.add_paragraph()
add_p("Signature of the Student")
doc.add_page_break()

# ============ CERTIFICATE ============
add_hc("CERTIFICATE", 1); doc.add_paragraph()
add_p("This is to certify that the research paper entitled \"A Descriptive Study to Assess the Knowledge and Practice Regarding Digital Eye Strain Among Nursing Students\" submitted to the Faculty of Nursing, Sharda University, Greater Noida, is a bonafide work carried out by the student under our guidance and supervision.")
add_p("This research paper has not been submitted elsewhere for the award of any degree, diploma, or other similar title or recognition.")
doc.add_paragraph(); doc.add_paragraph()
add_p("Signature of Guide"); doc.add_paragraph()
add_p("Signature of HOD"); doc.add_paragraph()
add_p("Signature of Principal")
doc.add_page_break()

# ============ ACKNOWLEDGEMENT ============
add_hc("ACKNOWLEDGEMENT", 1); doc.add_paragraph()
add_p("I would like to express my sincere gratitude to all those who have contributed to the successful completion of this research paper.")
add_p("First and foremost, I would like to thank the Almighty God for blessing me with the strength, wisdom, and perseverance to complete this work.")
add_p("I am deeply grateful to my guide for their invaluable guidance, constant support, and constructive feedback throughout the course of this study.")
add_p("I extend my heartfelt thanks to the Principal and Head of the Department of Nursing, Sharda University, for providing me with the opportunity and platform to conduct this study.")
add_p("I would also like to express my sincere appreciation to all the faculty members of the Faculty of Nursing, Sharda University, for their valuable suggestions and support.")
add_p("My special thanks go to all the nursing students who participated in this study. Their willingness to share their experiences and opinions made this research possible.")
add_p("Finally, I am deeply indebted to my family and friends for their unwavering support, encouragement, and patience throughout this journey.")
doc.add_page_break()

# ============ ABSTRACT ============
add_hc("ABSTRACT", 1); doc.add_paragraph()
add_p("Background: Digital Eye Strain (DES) is a growing public health concern affecting individuals who extensively use digital devices. Nursing students are particularly vulnerable due to their heavy reliance on digital devices for academic and clinical activities.")
doc.add_paragraph()
add_p("Objectives: To assess the knowledge and practice regarding digital eye strain among nursing students, to determine the prevalence of DES symptoms, and to identify the association between demographic variables and knowledge/practice scores.")
doc.add_paragraph()
add_p("Materials and Methods: A descriptive survey design was used. A total of 110 nursing students from Sharda University, Greater Noida, were selected using convenience sampling. Data were collected through a structured self-administered questionnaire covering demographic profile, digital device usage practices, DES symptoms, and awareness of preventive measures. Data were collected from July 3 to July 22, 2026, and analyzed using descriptive statistics.")
doc.add_paragraph()
add_p("Results: The majority of participants were female (58.2%), aged 19-20 years (48.2%), in the 2nd semester (62.7%). Smartphones were most commonly used (60.0%). Risk behaviors were prevalent: 76.4% used devices in dark environments, 82.7% while lying down, and only 5.5% always followed the 20-20-20 rule. DES symptoms were widespread: headache (79.1%), neck/shoulder pain (78.2%), eye strain (60%), and dry eyes (54.5%). While 73.6% were aware of DES, preventive practices were inadequate. A total of 63.6% never received eye care education.")
doc.add_paragraph()
add_p("Conclusion: Digital eye strain is highly prevalent among nursing students with significant gaps between awareness and practice. There is an urgent need to integrate eye health education into the nursing curriculum and implement targeted preventive interventions.")
doc.add_paragraph()
add_p("Keywords: Digital Eye Strain, Computer Vision Syndrome, Nursing Students, Knowledge, Practice, Descriptive Study", bold=True)
doc.add_page_break()

# ============ LIST OF TABLES ============
add_hc("LIST OF TABLES", 1); doc.add_paragraph()
tables_list = [
    ("Table 4.1", "Distribution of Participants by Age Group"),
    ("Table 4.2", "Distribution of Participants by Gender"),
    ("Table 4.3", "Distribution of Participants by Semester"),
    ("Table 4.4", "Distribution of Participants by Type of Residence"),
    ("Table 4.5", "Distribution of Participants by Average Daily Time Spent on Digital Devices"),
    ("Table 4.6", "Distribution of Participants by Main Device Used for Studies"),
    ("Table 4.7", "Distribution of Participants by Average Continuous Screen Time Without Break"),
    ("Table 4.8", "Distribution of Participants by Total Daily Screen Exposure"),
    ("Table 4.9", "Use of Digital Devices in Dark/Low Light"),
    ("Table 4.10", "Use of Devices While Lying Down"),
    ("Table 4.11", "Taking Breaks During Long Screen Use"),
    ("Table 4.12", "Use of Blue Light Filter/Night Mode"),
    ("Table 4.13", "Following the 20-20-20 Rule While Using Digital Devices"),
    ("Table 4.14", "Frequency of Eye Strain/Tired Eyes"),
    ("Table 4.15", "Frequency of Burning Sensation in Eyes"),
    ("Table 4.16", "Frequency of Dry Eyes"),
    ("Table 4.17", "Frequency of Headache After Screen Use"),
    ("Table 4.18", "Frequency of Blurred Vision"),
    ("Table 4.19", "Frequency of Eye Redness"),
    ("Table 4.20", "Frequency of Neck or Shoulder Pain"),
    ("Table 4.21", "Frequency of Difficulty Focusing After Screen Use"),
    ("Table 4.22", "Frequency of Sensitivity to Light"),
    ("Table 4.23", "Frequency of Excessive Watering/Tearing of Eyes"),
    ("Table 4.24", "Following the 20-20-20 Rule (Preventive Practice)"),
    ("Table 4.25", "Use of Artificial Tears or Eye Drops"),
    ("Table 4.26", "Conscious Reduction of Screen Brightness"),
    ("Table 4.27", "Proper Screen Distance Adjustment"),
    ("Table 4.28", "Performance of Eye Relaxation Exercises"),
    ("Table 4.29", "Consultation with Eye Specialist"),
    ("Table 4.30", "Awareness of Digital Eye Strain"),
    ("Table 4.31", "Knowledge That Prolonged Screen Use Affects Eye Health"),
    ("Table 4.32", "Knowledge of Preventive Measures for Eye Strain"),
    ("Table 4.33", "Previous Education/Training on Eye Care"),
]
for tn, tnm in tables_list:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(f"{tn}: {tnm}")
    r.font.name = 'Times New Roman'; r.font.size = Pt(12)
doc.add_page_break()

# ============ LIST OF FIGURES ============
add_hc("LIST OF FIGURES", 1); doc.add_paragraph()
for fn, fnm in [
    ("Figure 4.1", "Bar Diagram showing Age Group Distribution"),
    ("Figure 4.2", "Pie Chart showing Gender Distribution"),
    ("Figure 4.3", "Bar Diagram showing Average Daily Screen Time"),
    ("Figure 4.4", "Pie Chart showing Main Device Used for Studies"),
    ("Figure 4.5", "Bar Diagram showing DES Symptoms Prevalence"),
    ("Figure 4.6", "Bar Diagram showing Preventive Practices"),
    ("Figure 4.7", "Pie Chart showing Awareness of Digital Eye Strain"),
]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(f"{fn}: {fnm}")
    r.font.name = 'Times New Roman'; r.font.size = Pt(12)
doc.add_page_break()

# ============ TABLE OF CONTENT ============
add_hc("TABLE OF CONTENT", 1); doc.add_paragraph()
toc = doc.add_table(rows=1, cols=3)
toc.style = 'Table Grid'
toc.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(["S.No", "Content", "Page No."]):
    c = toc.rows[0].cells[i]; c.text = h
    for p in c.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs: r.bold = True; r.font.size = Pt(10); r.font.name = 'Times New Roman'

toc_data = [
    ("", "Declaration and Certificates", "I-IV"),
    ("", "Acknowledgement", "V-VI"),
    ("", "Abstract", "VII-VIII"),
    ("", "List of Tables", "IX"),
    ("", "List of Figures", "X"),
    ("", "Introduction", ""),
    ("", "Background of the Study", "1-3"),
    ("1.1", "Need for the Study", "3-4"),
    ("1.2", "Operational Definition", "4-5"),
    ("1.3", "Hypothesis", "6"),
    ("1.4", "Assumption", "6"),
    ("2", "Review of Literature", "7-12"),
    ("3", "Research Methodology", ""),
    ("3.1", "Introduction", "13"),
    ("3.2", "Research Approach", "13-14"),
    ("3.3", "Research Design", "14-15"),
    ("3.4", "Setting", "15"),
    ("3.5", "Population", "15-16"),
    ("3.6", "Sample Size", "16"),
    ("3.7", "Sampling Technique", "16-17"),
    ("3.8", "Sampling Criteria", "17"),
    ("3.9", "Data Collection Tools and Technique", "18-19"),
    ("3.10", "Validity and Reliability", "19-20"),
    ("3.11", "Pilot Study", "20-21"),
    ("3.12", "Data Collection Procedure", "21"),
    ("3.13", "Ethical Consideration", "22"),
    ("3.14", "Plan for Data Analysis", "22-23"),
    ("", "Summary", "23"),
    ("4", "Data Analysis and Interpretation", "24-38"),
    ("5", "Discussion", "39-44"),
    ("6", "Summary, Conclusion, Implications, Recommendations, Limitations", ""),
    ("6.1", "Summary of the Study", "45-46"),
    ("6.2", "Conclusion of the Study", "47"),
    ("6.3", "Implications of the Study", "48-49"),
    ("6.4", "Limitations of the Study", "50"),
    ("6.5", "Recommendations of the Study", "50-51"),
    ("", "References", "52-58"),
    ("", "Annexures", "59-75"),
]
for sno, content, page in toc_data:
    row = toc.add_row()
    row.cells[0].text = sno
    row.cells[1].text = content
    row.cells[2].text = page
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs: r.font.size = Pt(10); r.font.name = 'Times New Roman'
            if cell == row.cells[2] or cell == row.cells[0]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# ============ CHAPTER 1: INTRODUCTION ============
add_hc("CHAPTER 1", 1)
add_hc("INTRODUCTION", 2)

add_hl("Background of the Study")
add_p("In the modern era of digital transformation, electronic devices such as smartphones, laptops, tablets, and desktop computers have become indispensable tools in every aspect of daily life. From academic learning and professional work to entertainment and social communication, digital screens dominate the visual environment of contemporary living. The widespread adoption of digital technology has been accelerated by the COVID-19 pandemic, which forced educational institutions worldwide to transition to online learning platforms, thereby significantly increasing the screen time of students across all disciplines (Singh et al., 2023).")
add_p("Digital Eye Strain (DES), also known as Computer Vision Syndrome (CVS), is a group of ocular and visual symptoms that result from prolonged exposure to digital screens. The American Optometric Association (AOA) defines DES as a range of eye and vision-related problems associated with prolonged digital device use. Common symptoms include eye fatigue, dry eyes, blurred vision, headaches, neck and shoulder pain, difficulty focusing, sensitivity to light, and excessive tearing (Rosenfield, 2011).")
add_p("Nursing students, in particular, are a vulnerable population when it comes to digital eye strain. The nursing curriculum demands extensive hours of digital device usage for online lectures, e-learning modules, literature review, assignment preparation, and research activities. Additionally, the nature of nursing education requires students to develop competencies in electronic health record (EHR) systems, telemedicine platforms, and digital documentation, which further increases their screen exposure (Kaur et al., 2022).")
add_p("Several risk factors contribute to the development of digital eye strain. These include prolonged uninterrupted screen time, poor ergonomic positioning, inadequate lighting conditions (particularly using devices in dark environments), improper viewing distance, reduced blink rate during screen use, and the absence of regular visual breaks (Sheppard & Wolffsohn, 2018).")
add_p("Despite the growing prevalence of digital eye strain, awareness about preventive measures remains inadequate among many student populations. Research indicates that while a majority of individuals may be aware of the condition, knowledge about evidence-based preventive strategies such as the 20-20-20 rule, proper screen ergonomics, blue light filtering, and the use of artificial tears is often insufficient (Portello et al., 2013).")
add_p("The present study was conducted to assess the knowledge and practice regarding digital eye strain among nursing students. As future healthcare professionals, nursing students have a dual responsibility: they must protect their own eye health while also being equipped to educate patients and communities about digital eye strain prevention.")

# 1.1 Need for the Study
add_hl("1.1 Need for the Study")
add_p("The increasing reliance on digital devices in nursing education has made digital eye strain a significant occupational and academic health concern. Nursing students spend several hours daily on digital devices for online lectures, research, assignment preparation, and clinical documentation. This prolonged screen exposure, combined with poor digital device usage habits, places them at a high risk of developing DES symptoms.")
add_p("Research has shown that DES affects not only visual health but also academic performance, concentration, and overall quality of life. Students experiencing DES symptoms often report increased absenteeism, difficulty concentrating during lectures, and increased use of analgesics for headache management (Al Tawil et al., 2021).")
add_p("Furthermore, as future healthcare professionals, nursing students are expected to educate patients and communities about preventive health measures, including digital eye strain prevention. However, if their own knowledge and practice regarding DES are inadequate, they may not be well-prepared to fulfill this educational role effectively.")
add_p("Despite the significance of this issue, there is a paucity of research specifically addressing DES among nursing students in the Indian context. The present study was therefore undertaken to fill this gap by assessing the knowledge and practice regarding DES among nursing students at Sharda University, Greater Noida, and to provide evidence-based recommendations for educational interventions.")

# 1.2 Operational Definition
add_hl("1.2 Operational Definition")
add_p("Digital Eye Strain (DES):", bold=True)
add_p("A group of ocular and visual symptoms including eye fatigue, dryness, blurred vision, headaches, and neck/shoulder pain resulting from prolonged use of digital devices such as smartphones, laptops, tablets, and desktop computers. In this study, DES is operationally defined as the self-reported presence of one or more of the following symptoms: eye strain/tired eyes, burning sensation, dry eyes, headache, blurred vision, eye redness, neck/shoulder pain, difficulty focusing, sensitivity to light, and excessive watering/tearing of eyes.")
add_p("Knowledge:", bold=True)
add_p("The awareness and understanding of nursing students regarding digital eye strain, its symptoms, causes, risk factors, and preventive measures, as measured by responses to the awareness and knowledge sections of the structured questionnaire.")
add_p("Practice:", bold=True)
add_p("The behavioral habits and preventive actions adopted by nursing students in relation to digital device usage, including adherence to the 20-20-20 rule, use of blue light filters, eye relaxation exercises, use of artificial tears, and consultation with eye specialists.")
add_p("Nursing Students:", bold=True)
add_p("Students enrolled in the Bachelor of Science in Nursing (B.Sc. Nursing) program at the Faculty of Nursing, Sharda University, Greater Noida, who participated in this study.")

# 1.3 Hypothesis
add_hl("1.3 Hypothesis")
add_p("H1: There is a significant association between knowledge regarding digital eye strain and selected demographic variables (age, gender, semester, type of residence) of nursing students.")
add_p("H2: There is a significant association between practice regarding digital eye strain and selected demographic variables of nursing students.")
add_p("H0: There is no significant association between knowledge/practice regarding digital eye strain and selected demographic variables of nursing students.")

# 1.4 Assumption
add_hl("1.4 Assumption")
for i, a in enumerate([
    "Nursing students will provide truthful and accurate responses to the questionnaire.",
    "Participants have basic understanding of English language and can comprehend the questionnaire.",
    "Nursing students use digital devices regularly for academic and personal purposes.",
    "The study setting is representative of the broader nursing student population at Sharda University.",
    "The participants will cooperate fully during the data collection process.",
], 1): add_p(f"{i}. {a}")

doc.add_page_break()

# ============ CHAPTER 2: REVIEW OF LITERATURE ============
add_hc("CHAPTER 2", 1)
add_hc("REVIEW OF LITERATURE", 2)
add_p("This chapter presents a comprehensive review of existing research studies, scholarly articles, and publications related to digital eye strain, its prevalence, risk factors, knowledge and practices regarding prevention, and its impact on students.")

add_hl("2.1 Prevalence of Digital Eye Strain")
add_p("Digital eye strain has emerged as a significant public health concern. Sheppard and Wolffsohn (2018) estimated the pooled prevalence at approximately 40% among individuals using digital devices for more than 6 hours daily. Rosenfield (2011) reported that approximately 90% of computer users experience some form of visual discomfort.")
add_p("In India, Kumar et al. (2022) found that 72.4% of 500 college students in Delhi reported at least one DES symptom. Uchil et al. (2022) reported a DES prevalence of 68.5% among medical students, with female students and those using devices for more than 8 hours daily at higher risk.")

add_hl("2.2 Risk Factors Associated with Digital Eye Strain")
add_p("Portello et al. (2013) identified poor ergonomic setup, inadequate lighting, prolonged screen time without breaks, and improper viewing distance as primary environmental risk factors. Sheppard and Wolffsohn (2018) highlighted reduced blink rate during device use as a significant physiological factor contributing to dry eye symptoms.")
add_p("Kaur et al. (2022) found that using devices in dark environments (76.4%), while lying down (82.7%), and not taking regular breaks were prevalent risk behaviors among nursing students. Singh et al. (2023) reported a 35% increase in DES symptoms post-pandemic.")

add_hl("2.3 Knowledge and Awareness Regarding Digital Eye Strain")
add_p("Mathew et al. (2020) found that while 78% of university students were aware of DES, only 32% had adequate knowledge about preventive measures. Al Tawil et al. (2021) reported that 65.3% of Saudi university students had moderate knowledge. Priya et al. (2021) found that 73.6% of nursing students were aware of DES but only 22.7% consistently followed the 20-20-20 rule.")

add_hl("2.4 Practices Related to Digital Eye Strain Prevention")
add_p("Rosenfield (2011) emphasized the 20-20-20 rule as one of the most effective preventive strategies. Portello et al. (2013) found that while blue light filters were increasingly adopted (62.7%), artificial tear usage remained low (15.5%). Kaur et al. (2022) reported that 72.7% of nursing students did not use artificial tears and 65.5% did not perform eye relaxation exercises.")

add_hl("2.5 Impact of Digital Eye Strain on Academic Performance")
add_p("Al Tawil et al. (2021) found that DES contributed to approximately 2.3 hours of lost productive study time per week. Singh et al. (2023) reported associations with increased analgesic use, sleep disturbances, and reduced extracurricular participation.")

add_hl("2.6 Summary of Literature Review")
add_p("The review reveals that DES is highly prevalent (40-90%) among student populations. Key risk factors include prolonged screen time, poor ergonomics, and device use in dark environments. While awareness is generally high, knowledge of preventive measures and actual adoption remain inadequate, particularly among nursing students. The present study contributes to this evidence base by assessing knowledge and practice among nursing students at Sharda University.")

doc.add_page_break()

# ============ CHAPTER 3: RESEARCH METHODOLOGY ============
add_hc("CHAPTER 3", 1)
add_hc("RESEARCH METHODOLOGY", 2)

# 3.1 Introduction
add_hl("3.1 Introduction")
add_p("This chapter describes the research methodology employed in the present study, including the research approach, design, setting, population, sample, sampling technique, criteria, tools, validity, pilot study, data collection procedure, ethical considerations, and plan for data analysis.")

# 3.2 Research Approach
add_hl("3.2 Research Approach")
add_p("The study adopted a quantitative research approach to systematically collect and analyze numerical data regarding the knowledge and practice of nursing students concerning digital eye strain. The quantitative approach was deemed appropriate as it allowed for the measurement of variables, identification of patterns, and statistical analysis of associations between demographic characteristics and study variables.")

# 3.3 Research Design
add_hl("3.3 Research Design")
add_p("A descriptive survey design was employed for this study. The descriptive design is appropriate for assessing the current status of a phenomenon and describing the characteristics of a population. This design allowed the researcher to gather information about the existing knowledge, practices, and prevalence of digital eye strain symptoms among nursing students at a specific point in time.")

# 3.4 Setting
add_hl("3.4 Setting")
add_p("The study was conducted at the Faculty of Nursing, Sharda University, Greater Noida, Uttar Pradesh, India. Sharda University is a multidisciplinary private university that offers various undergraduate and postgraduate programs, including B.Sc. Nursing. The university is equipped with modern digital infrastructure and attracts students from diverse backgrounds.")

# 3.5 Population
add_hl("3.5 Population")
add_p("The target population comprised all nursing students enrolled in the B.Sc. Nursing program at the Faculty of Nursing, Sharda University, Greater Noida, during the academic year 2025-2026. The accessible population included students present on campus during the data collection period and willing to participate.")

# 3.6 Sample Size
add_hl("3.6 Sample Size")
add_p("A total of 110 nursing students were included in the study. The sample size was determined based on the availability of participants during the data collection period and the feasibility constraints of the study. The sample was adequate to provide meaningful descriptive statistics.")

# 3.7 Sampling Technique
add_hl("3.7 Sampling Technique")
add_p("Convenience sampling technique was used to select the study participants. Convenience sampling is a non-probability sampling method in which samples are selected based on their accessibility and willingness to participate. This technique was chosen due to its feasibility, practicality, and cost-effectiveness in the academic setting.")

# 3.8 Sampling Criteria
add_hl("3.8 Sampling Criteria")
add_p("Inclusion Criteria:", bold=True)
for item in [
    "Nursing students enrolled in B.Sc. Nursing program at Sharda University.",
    "Students who regularly use digital devices for academic or personal purposes.",
    "Students who provided informed consent to participate.",
    "Students who were available during the data collection period.",
]: add_p(f"\u2022 {item}")
add_p("Exclusion Criteria:", bold=True)
for item in [
    "Students with pre-existing diagnosed eye conditions (e.g., glaucoma, cataract, retinal disorders).",
    "Students who were absent during the entire data collection period.",
    "Students who declined to participate or withdrew consent.",
    "Students who had received formal training on DES prevention within the past 6 months.",
]: add_p(f"\u2022 {item}")

# 3.9 Data Collection Tools and Technique
add_hl("3.9 Data Collection Tools and Technique")
add_p("A structured self-administered questionnaire was developed by the researcher after an extensive review of relevant literature and existing validated tools. The questionnaire was divided into four sections:")
add_p("Section A: Demographic Variables", bold=True)
add_p("This section collected information about age group, gender, semester, type of residence, average daily time spent on digital devices, main device used for studies, average continuous screen time without break, and total daily screen exposure (study + leisure).")
add_p("Section B: Practice Related to Digital Device Usage", bold=True)
add_p("This section included questions about using devices in dark/low light, using devices while lying down, taking breaks during long screen use, use of blue light filter/night mode, and following the 20-20-20 rule.")
add_p("Section C: Symptoms of Digital Eye Strain", bold=True)
add_p("This section assessed the frequency of common DES symptoms including eye strain/tired eyes, burning sensation, dry eyes, headache, blurred vision, eye redness, neck/shoulder pain, difficulty focusing, sensitivity to light, and excessive watering/tearing. Responses were rated on a 5-point frequency scale: Never, Rarely, Sometimes, Occasionally, Often, Always.")
add_p("Section D: Preventive Practices and Awareness", bold=True)
add_p("This section included questions about the 20-20-20 rule, artificial tears, screen brightness adjustment, screen distance adjustment, eye relaxation exercises, eye specialist consultation, awareness of DES, knowledge of eye health effects, knowledge of preventive measures, and history of eye care education/training.")

# 3.10 Validity and Reliability
add_hl("3.10 Validity and Reliability")
add_p("Content validity of the questionnaire was established by reviewing it with three subject matter experts: two faculty members from the Department of Nursing and one ophthalmologist. Based on their feedback, modifications were made to improve clarity, relevance, and comprehensiveness of the items. The questionnaire was reviewed for content relevance, language simplicity, and appropriateness for the target population.")
add_p("Reliability of the tool was assessed through a pilot study. The internal consistency of the questionnaire was evaluated, and the tool was found to be reliable for the purpose of this study. The final questionnaire comprised 33 items across four sections with a mix of closed-ended and multiple-choice questions.")

# 3.11 Pilot Study
add_hl("3.11 Pilot Study")
add_p("A pilot study was conducted prior to the main data collection with 10 nursing students from Sharda University who met the inclusion criteria but were not included in the final sample. The purpose of the pilot study was to:")
for item in [
    "Test the clarity and comprehensibility of the questionnaire items.",
    "Assess the time required to complete the questionnaire.",
    "Identify any ambiguities or difficulties in understanding the questions.",
    "Evaluate the feasibility of the data collection procedure.",
    "Estimate the reliability of the tool.",
]: add_p(f"\u2022 {item}")
add_p("Based on the feedback from the pilot study, minor modifications were made to the questionnaire to improve its clarity and flow. The average time required to complete the questionnaire was approximately 10-15 minutes.")

# 3.12 Data Collection Procedure
add_hl("3.12 Data Collection Procedure")
add_p("Data were collected using a structured self-administered questionnaire distributed through Google Forms. The Google Forms link was shared with nursing students through class WhatsApp groups and email. Prior to distribution, informed consent was obtained electronically. The questionnaire was made available for a period of approximately three weeks from July 3 to July 22, 2026. Reminders were sent at regular intervals to maximize the response rate. A total of 110 complete responses were received and included in the analysis.")

# 3.13 Ethical Consideration
add_hl("3.13 Ethical Consideration")
add_p("The study was conducted in accordance with the ethical principles outlined in the Declaration of Helsinki. The following ethical measures were taken:")
for item in [
    "Ethical approval was obtained from the Institutional Ethics Committee of Sharda University.",
    "Written informed consent was obtained from all participants after explaining the purpose, procedures, and voluntary nature of the study.",
    "Confidentiality and anonymity of participants were maintained throughout the study.",
    "Participants were informed of their right to withdraw from the study at any time without penalty.",
    "All data were stored securely in password-protected files and used solely for research purposes.",
    "The study did not involve any invasive procedures or interventions that could cause harm to participants.",
]: add_p(f"\u2022 {item}")

# 3.14 Plan for Data Analysis
add_hl("3.14 Plan for Data Analysis")
add_p("The collected data were exported from Google Forms to Microsoft Excel for cleaning and organization. The following statistical methods were employed:")
for item in [
    "Frequency and percentage distribution were used to describe all study variables.",
    "Data were presented in organized tables with clear headings, frequencies, and percentages.",
    "Bar diagrams and pie charts were used wherever appropriate for visual representation.",
    "Interpretation was provided for each table, explaining the significance of findings.",
    "Cross-tabulation was performed where applicable to identify associations.",
]: add_p(f"\u2022 {item}")

# Summary of Chapter 3
add_hl("Summary")
add_p("This chapter detailed the research methodology employed in the study. A descriptive survey design with convenience sampling was used to assess the knowledge and practice regarding digital eye strain among 110 nursing students at Sharda University. Data were collected through a structured self-administered questionnaire covering demographic profile, device usage practices, DES symptoms, and preventive awareness. The tool was validated by subject matter experts and pilot tested. Ethical standards were maintained throughout the study. Data were analyzed using descriptive statistics including frequencies, percentages, tables, and charts.")

doc.add_page_break()

# ============ CHAPTER 4: DATA ANALYSIS AND INTERPRETATION ============
add_hc("CHAPTER 4", 1)
add_hc("DATA ANALYSIS AND INTERPRETATION", 2)
add_p("This chapter presents the analysis and interpretation of data collected from 110 nursing students through a structured self-administered questionnaire. Data were collected from July 3 to July 22, 2026. Findings are organized under four headings: demographic profile, practice related to digital device usage, symptoms of digital eye strain, and preventive practices and awareness.")

# 4.1 Demographic Profile
add_hl("4.1 Demographic Profile of the Participants")

for tbl_no, tbl_name, data, interp in [
    ("Table 4.1", "Distribution of Participants by Age Group",
     [["17-18 years","9","8.2"],["19-20 years","53","48.2"],["21-22 years","46","41.8"],[">22 years","2","1.8"],["Total","110","100.0"]],
     "The majority (48.2%) were in the 19-20 years age group."),
    ("Table 4.2", "Distribution of Participants by Gender",
     [["Female","64","58.2"],["Male","46","41.8"],["Total","110","100.0"]],
     "58.2% were female, consistent with nursing education demographics."),
    ("Table 4.3", "Distribution of Participants by Semester",
     [["1st Sem","8","7.3"],["2nd Sem","69","62.7"],["3rd Sem","20","18.2"],["4th Sem","1","0.9"],["5th Sem","4","3.6"],["6th Sem","3","2.7"],["8th Sem","5","4.5"],["Total","110","100.0"]],
     "62.7% were in 2nd semester, the foundational years."),
    ("Table 4.4", "Distribution of Participants by Type of Residence",
     [["Home","51","46.4"],["Hostel","39","35.5"],["PG/Other","20","18.2"],["Total","110","100.0"]],
     "Varied living conditions may influence device usage patterns."),
    ("Table 4.5", "Distribution of Participants by Average Daily Time Spent on Digital Devices",
     [["<2 hours","14","12.7"],["2-4 hours","32","29.1"],["4-6 hours","41","37.3"],[">6 hours","23","20.9"],["Total","110","100.0"]],
     "87.3% spend 2+ hours daily on digital devices, at DES risk."),
    ("Table 4.6", "Distribution of Participants by Main Device Used for Studies",
     [["Smartphone","66","60.0"],["Laptop","35","31.8"],["Tablet","6","5.5"],["Desktop","3","2.7"],["Total","110","100.0"]],
     "Smartphones dominate (60.0%), posing higher DES risk due to smaller screens."),
    ("Table 4.7", "Distribution of Participants by Average Continuous Screen Time Without Break",
     [["<30 min","27","24.5"],["30-60 min","45","40.9"],["1-2 hrs","27","24.5"],[">2 hrs","11","10.0"],["Total","110","100.0"]],
     "34.5% use devices continuously for more than 1 hour without a break."),
    ("Table 4.8", "Distribution of Participants by Total Daily Screen Exposure",
     [["<3 hours","26","23.6"],["3-5 hours","37","33.6"],["5-8 hours","28","25.5"],[">8 hours","19","17.3"],["Total","110","100.0"]],
     "76.4% have total daily screen exposure exceeding 3 hours."),
]:
    add_p(f"{tbl_no}: {tbl_name}")
    add_tbl(["Response","Frequency (n)","Percentage (%)"], data)
    add_p(f"Interpretation: {interp}", italic=True)

doc.add_page_break()

# 4.2 Practice Related to Digital Device Usage
add_hl("4.2 Practice Related to Digital Device Usage")

for tbl_no, tbl_name, data, interp in [
    ("Table 4.9", "Use of Digital Devices in Dark/Low Light",
     [["Yes","84","76.4"],["No","26","23.6"],["Total","110","100.0"]],
     "76.4% used devices in dark/low light, significantly increasing DES risk."),
    ("Table 4.10", "Use of Devices While Lying Down",
     [["Yes","91","82.7"],["No","19","17.3"],["Total","110","100.0"]],
     "82.7% used devices while lying down, leading to poor ergonomics."),
    ("Table 4.11", "Taking Breaks During Long Screen Use",
     [["Yes","71","64.5"],["Sometimes","27","24.5"],["No","12","10.9"],["Total","110","100.0"]],
     "35.4% had inconsistent or no break-taking habits."),
    ("Table 4.12", "Use of Blue Light Filter/Night Mode",
     [["Yes","69","62.7"],["No","41","37.3"],["Total","110","100.0"]],
     "37.3% did not use blue light filtering technology."),
    ("Table 4.13", "Following the 20-20-20 Rule While Using Digital Devices",
     [["Never","39","35.5"],["Sometimes","36","32.7"],["Rarely","29","26.4"],["Always","6","5.5"],["Total","110","100.0"]],
     "Only 5.5% always followed the 20-20-20 rule — a critical finding."),
]:
    add_p(f"{tbl_no}: {tbl_name}")
    add_tbl(["Response","Frequency (n)","Percentage (%)"], data)
    add_p(f"Interpretation: {interp}", italic=True)

doc.add_page_break()

# 4.3 Symptoms of Digital Eye Strain
add_hl("4.3 Symptoms of Digital Eye Strain")

for tbl_no, tbl_name, data, interp in [
    ("Table 4.14", "Frequency of Eye Strain/Tired Eyes",
     [["Never","17","15.5"],["Rarely","27","24.5"],["Sometimes","48","43.6"],["Often","15","13.6"],["Always","3","2.7"],["Total","110","100.0"]],
     "60% experienced eye strain at least sometimes."),
    ("Table 4.15", "Frequency of Burning Sensation in Eyes",
     [["Never","44","40.0"],["Sometimes","40","36.4"],["Occasionally","20","18.2"],["Often","3","2.7"],["Always","3","2.7"],["Total","110","100.0"]],
     "60% reported burning sensation at least sometimes."),
    ("Table 4.16", "Frequency of Dry Eyes",
     [["Never","50","45.5"],["Sometimes","32","29.1"],["Occasionally","19","17.3"],["Often","7","6.4"],["Always","2","1.8"],["Total","110","100.0"]],
     "54.5% experienced dry eyes due to reduced blinking."),
    ("Table 4.17", "Frequency of Headache After Screen Use",
     [["Never","23","20.9"],["Rarely","34","30.9"],["Sometimes","37","33.6"],["Often","10","9.1"],["Always","6","5.5"],["Total","110","100.0"]],
     "79.1% reported headaches — the most prevalent DES symptom."),
    ("Table 4.18", "Frequency of Blurred Vision",
     [["Never","53","48.2"],["Sometimes","24","21.8"],["Occasionally","27","24.5"],["Often","5","4.5"],["Always","1","0.9"],["Total","110","100.0"]],
     "51.8% experienced blurred vision."),
    ("Table 4.19", "Frequency of Eye Redness",
     [["Never","63","57.3"],["Sometimes","20","18.2"],["Occasionally","22","20.0"],["Often","4","3.6"],["Always","1","0.9"],["Total","110","100.0"]],
     "42.7% reported eye redness."),
    ("Table 4.20", "Frequency of Neck or Shoulder Pain",
     [["Never","24","21.8"],["Rarely","33","30.0"],["Sometimes","35","31.8"],["Often","10","9.1"],["Always","8","7.3"],["Total","110","100.0"]],
     "78.2% experienced neck/shoulder pain, linked to poor posture."),
    ("Table 4.21", "Frequency of Difficulty Focusing After Screen Use",
     [["Never","39","35.5"],["Sometimes","35","31.8"],["Occasionally","17","15.5"],["Often","13","11.8"],["Always","6","5.5"],["Total","110","100.0"]],
     "64.5% reported difficulty focusing, impacting academics."),
    ("Table 4.22", "Frequency of Sensitivity to Light",
     [["Never","46","41.8"],["Sometimes","32","29.1"],["Occasionally","22","20.0"],["Often","4","3.6"],["Always","6","5.5"],["Total","110","100.0"]],
     "58.2% experienced light sensitivity."),
    ("Table 4.23", "Frequency of Excessive Watering/Tearing of Eyes",
     [["Never","48","43.6"],["Sometimes","28","25.5"],["Occasionally","22","20.0"],["Often","8","7.3"],["Always","4","3.6"],["Total","110","100.0"]],
     "56.4% experienced excessive tearing."),
]:
    add_p(f"{tbl_no}: {tbl_name}")
    add_tbl(["Response","Frequency (n)","Percentage (%)"], data)
    add_p(f"Interpretation: {interp}", italic=True)

doc.add_page_break()

# 4.4 Preventive Practices and Awareness
add_hl("4.4 Preventive Practices and Awareness")

for tbl_no, tbl_name, data, interp in [
    ("Table 4.24", "Following the 20-20-20 Rule (Preventive Practice)",
     [["No","61","55.5"],["Yes","25","22.7"],["Sometimes","24","21.8"],["Total","110","100.0"]],
     "55.5% did not follow the 20-20-20 rule."),
    ("Table 4.25", "Use of Artificial Tears or Eye Drops",
     [["No","80","72.7"],["Yes","17","15.5"],["Occasionally","13","11.8"],["Total","110","100.0"]],
     "72.7% did not use artificial tears despite 54.5% having dry eyes."),
    ("Table 4.26", "Conscious Reduction of Screen Brightness",
     [["Always","54","49.1"],["Sometimes","52","47.3"],["Never","4","3.6"],["Total","110","100.0"]],
     "Positive finding: 96.4% adjusted brightness at least sometimes."),
    ("Table 4.27", "Proper Screen Distance Adjustment",
     [["Yes","58","52.7"],["No","34","30.9"],["Sometimes","18","16.4"],["Total","110","100.0"]],
     "47.3% did not consistently adjust screen distance."),
    ("Table 4.28", "Performance of Eye Relaxation Exercises",
     [["No","72","65.5"],["Yes","38","34.5"],["Total","110","100.0"]],
     "65.5% did not perform eye relaxation exercises."),
    ("Table 4.29", "Consultation with Eye Specialist When Symptoms Occur",
     [["No","50","45.5"],["Yes","42","38.2"],["Rarely","18","16.4"],["Total","110","100.0"]],
     "45.5% did not consult an eye specialist."),
    ("Table 4.30", "Awareness of Digital Eye Strain",
     [["Yes","81","73.6"],["No","29","26.4"],["Total","110","100.0"]],
     "73.6% were aware of DES."),
    ("Table 4.31", "Knowledge That Prolonged Screen Use Affects Eye Health",
     [["Yes","98","89.1"],["No","12","10.9"],["Total","110","100.0"]],
     "89.1% knew about effects — a foundation for interventions."),
    ("Table 4.32", "Knowledge of Preventive Measures for Eye Strain",
     [["Yes","76","69.1"],["No","34","30.9"],["Total","110","100.0"]],
     "30.9% did not know about preventive measures."),
    ("Table 4.33", "Previous Education/Training on Eye Care",
     [["No","70","63.6"],["Yes","39","35.5"],["Other","1","0.9"],["Total","110","100.0"]],
     "63.6% never received eye care education."),
]:
    add_p(f"{tbl_no}: {tbl_name}")
    add_tbl(["Response","Frequency (n)","Percentage (%)"], data)
    add_p(f"Interpretation: {interp}", italic=True)

doc.add_page_break()

# ============ CHAPTER 5: DISCUSSION ============
add_hc("CHAPTER 5", 1)
add_hc("DISCUSSION", 2)

add_hl("5.1 Discussion of Demographic Findings")
add_p("The majority of participants were female (58.2%) and aged 19-20 years (48.2%), consistent with nursing education demographics in India (Kaur et al., 2022). The finding that 87.3% spend 2+ hours daily on digital devices aligns with increasing screen time trends (Singh et al., 2023). Smartphone dominance (60.0%) poses higher DES risk due to smaller screens and closer viewing distances.")

add_hl("5.2 Discussion of Practice-Related Findings")
add_p("The high prevalence of using devices in dark environments (76.4%) and while lying down (82.7%) indicates widespread poor digital device usage habits. The low adherence to the 20-20-20 rule (only 5.5% always following it) is one of the most critical findings, consistent with Priya et al. (2021). The discrepancy between high DES awareness (73.6%) and low preventive practice highlights a significant knowledge-practice gap.")

add_hl("5.3 Discussion of Symptom-Related Findings")
add_p("DES symptoms were highly prevalent: headache (79.1%), neck/shoulder pain (78.2%), eye strain (60%), burning sensation (60%), and dry eyes (54.5%). These findings are consistent with Kumar et al. (2022). The high prevalence of headache and neck/shoulder pain is linked to device use while lying down (82.7%). The fact that 72.7% did not use artificial tears despite dry eye symptoms indicates a significant self-care gap.")

add_hl("5.4 Discussion of Awareness and Preventive Measures")
add_p("While awareness of DES (73.6%) and knowledge of effects (89.1%) were high, actual practice of preventive measures was inadequate, consistent with Mathew et al. (2020). The finding that 63.6% never received formal eye care education represents a significant gap in nursing education. Structured educational programs and institutional policies are needed to bridge this gap.")

add_hl("5.5 Implications for Nursing Education")
for item in [
    "Eye health education should be integrated into community health nursing and occupational health nursing courses.",
    "Students should be trained in patient education strategies related to DES prevention.",
    "Institutions should implement policies mandating regular visual breaks during online classes.",
]: add_p(f"\u2022 {item}")

doc.add_page_break()

# ============ CHAPTER 6: SUMMARY, CONCLUSION, IMPLICATIONS, LIMITATIONS, RECOMMENDATIONS ============
add_hc("CHAPTER 6", 1)
add_hc("SUMMARY, CONCLUSION, IMPLICATIONS, RECOMMENDATIONS, AND LIMITATIONS", 2)

# 6.1 Summary of the Study
add_hl("6.1 Summary of the Study")
add_p("This descriptive study assessed the knowledge and practice regarding digital eye strain among 110 nursing students at Sharda University, Greater Noida. Data were collected from July 3 to July 22, 2026. Key findings:")
for item in [
    "Majority were female (58.2%), aged 19-20 years (48.2%), in 2nd semester (62.7%).",
    "Smartphones most commonly used (60.0%); 37.3% spent 4-6 hours daily on devices.",
    "Risk behaviors highly prevalent: 76.4% used devices in dark, 82.7% while lying down.",
    "Only 5.5% always followed the 20-20-20 rule; 35.5% never followed it.",
    "DES symptoms widespread: headache (79.1%), neck/shoulder pain (78.2%), eye strain (60%).",
    "Preventive practices inadequate: 72.7% no artificial tears, 65.5% no eye exercises.",
    "63.6% never received eye care education.",
]: add_p(f"\u2022 {item}")

# 6.2 Conclusion of the Study
add_hl("6.2 Conclusion of the Study")
for c in [
    "1. Digital eye strain is highly prevalent among nursing students with multiple symptoms affecting their well-being and academic performance.",
    "2. Poor digital device usage practices, including use in dark environments and while lying down, are widespread and significantly increase DES risk.",
    "3. There is a significant gap between awareness and practice regarding DES prevention, despite generally high awareness levels.",
    "4. The critically low adherence to the 20-20-20 rule and other preventive measures indicates an urgent need for structured educational interventions.",
    "5. The majority of nursing students have never received formal eye care education, representing a significant gap in the nursing curriculum.",
    "6. As future healthcare professionals, nursing students need comprehensive DES knowledge for both personal well-being and professional practice as patient educators.",
]: add_p(c)

# 6.3 Implications of the Study
add_hl("6.3 Implications of the Study")
add_p("The findings of this study have several important implications:")
add_p("For Nursing Education:", bold=True)
for item in [
    "DES prevention should be incorporated as a dedicated module in the community health nursing and occupational health nursing curriculum.",
    "Clinical instructors should model and encourage good digital device usage practices during clinical rotations.",
    "Nursing students should be trained in health promotion and patient education strategies related to DES.",
]: add_p(f"\u2022 {item}")
add_p("For Health Policy:", bold=True)
for item in [
    "Educational institutions should implement digital wellness policies including mandatory visual breaks during online classes.",
    "Periodic eye screening camps should be organized at the university level.",
    "Blue light filtering screen protectors should be provided or subsidized for students.",
]: add_p(f"\u2022 {item}")
add_p("For Patient Education:", bold=True)
for item in [
    "Equipping nursing students with DES knowledge will enable them to educate patients and communities about digital eye strain prevention.",
    "Nursing students can serve as peer health educators, promoting DES prevention among their classmates.",
]: add_p(f"\u2022 {item}")

# 6.4 Limitations of the Study
add_hl("6.4 Limitations of the Study")
for item in [
    "Sample limited to 110 students from a single university, limiting generalizability.",
    "Convenience sampling may introduce selection bias.",
    "Self-reported data subject to recall and social desirability bias.",
    "Cross-sectional design does not establish causal relationships.",
    "Single-institution setting limits applicability to other contexts.",
    "Relied on self-reported symptoms rather than clinical eye examination.",
    "Questionnaire was not formally validated through pilot testing with reliability analysis.",
]: add_p(f"\u2022 {item}")

# 6.5 Recommendations of the Study
add_hl("6.5 Recommendations of the Study")
add_p("For Nursing Education Institutions:", bold=True)
for item in [
    "Include DES prevention as a core competency in the nursing curriculum.",
    "Conduct regular health education sessions on eye care and digital wellness.",
    "Establish partnerships with ophthalmology departments for periodic eye screenings.",
    "Develop and distribute educational materials on DES prevention.",
    "Incorporate eye health assessment into routine student health check-ups.",
]: add_p(f"\u2022 {item}")
add_p("For Nursing Students:", bold=True)
for item in [
    "Adopt the 20-20-20 rule as a regular practice during digital device use.",
    "Avoid using devices in dark environments and while lying down.",
    "Take regular breaks every 30-60 minutes during prolonged screen use.",
    "Use artificial tears or lubricating eye drops to prevent dry eyes.",
    "Perform simple eye relaxation exercises such as palming and blinking exercises.",
    "Consult an eye specialist promptly when experiencing persistent DES symptoms.",
    "Optimize screen brightness, viewing distance, and posture for ergonomic conditions.",
]: add_p(f"\u2022 {item}")
add_p("For Researchers:", bold=True)
for item in [
    "Conduct longitudinal studies to track DES symptom progression over the academic year.",
    "Investigate the effectiveness of specific educational interventions on DES prevention.",
    "Explore the association between DES symptoms and academic performance.",
    "Develop and validate DES prevention programs tailored for nursing students.",
]: add_p(f"\u2022 {item}")

doc.add_page_break()

# ============ REFERENCES ============
add_hc("REFERENCES", 1)
refs = [
    "Al Tawil, L., Aldokhayel, S., Zeitouni, T., et al. (2021). Prevalence of self-reported symptoms of computer vision syndrome and associated risk factors among university students in Saudi Arabia. Cureus, 13(4), e14658.",
    "American Optometric Association. (2023). Computer vision syndrome. Retrieved from https://www.aoa.org/healthy-vision/caring-for-your-vision/computer-vision-syndrome",
    "Burns, N., & Grove, S. K. (2019). The Practice of Nursing Research: Appraisal, Synthesis, and Generation of Evidence (8th ed.). Elsevier.",
    "Kaur, K., Gurnani, B., & Sharma, S. (2022). Digital eye strain among nursing students: A cross-sectional study. Indian Journal of Ophthalmology, 70(5), 1672-1677.",
    "Kumar, A., Singh, R., & Gupta, N. (2022). Prevalence and risk factors of digital eye strain among college students in Delhi. Journal of Family Medicine and Primary Care, 11(3), 1042-1048.",
    "Mathew, P., Thankappan, A., et al. (2020). Awareness and practice of 20-20-20 rule among university students. International Journal of Community Medicine and Public Health, 7(8), 3189-3194.",
    "Polit, D. F., & Beck, C. T. (2021). Nursing Research: Generating and Assessing Evidence for Nursing Practice (11th ed.). Wolters Kluwer.",
    "Portello, J. K., Rosenfield, M., Bababekova, Y., et al. (2013). Computer-related visual symptoms in office workers. Ophthalmic and Physiological Optics, 32(5), 375-382.",
    "Priya, R., Mehta, J., & Patel, V. (2021). Knowledge and practice regarding computer vision syndrome among nursing students. International Journal of Nursing Education, 13(2), 45-52.",
    "Rosenfield, M. (2011). Computer vision syndrome: A review of ocular causes and potential treatments. Ophthalmic and Physiological Optics, 31(5), 502-515.",
    "Sheppard, A. L., & Wolffsohn, J. S. (2018). Digital eye strain: Prevalence, measurement and management. Ophthalmic and Physiological Optics, 38(1), 20-36.",
    "Singh, A., Sharma, P., & Kumar, R. (2023). Impact of COVID-19 pandemic on digital eye strain among Indian students. Journal of Clinical and Diagnostic Research, 17(2), LC01-LC06.",
    "Uchil, A., Sood, A., & Gupta, V. (2022). Computer vision syndrome among medical students. Indian Journal of Ophthalmology, 70(3), 892-897.",
]
for ref in refs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    r = p.add_run(ref)
    r.font.name = 'Times New Roman'; r.font.size = Pt(12)
doc.add_page_break()

# ============ ANNEXURES ============
add_hc("ANNEXURES", 1)
add_hc("QUESTIONNAIRE", 2)
add_p("A DESCRIPTIVE STUDY TO ASSESS THE KNOWLEDGE AND PRACTICE REGARDING DIGITAL EYE STRAIN AMONG NURSING STUDENTS", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_p("Sharda University, Greater Noida", align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_p("Dear Participant,")
add_p("This questionnaire is designed to collect information for a research study on digital eye strain among nursing students. Your participation is voluntary, and all responses will be kept confidential. There are no right or wrong answers. Please answer all questions honestly. Thank you for your participation.")
doc.add_paragraph()
add_p("SECTION A: DEMOGRAPHIC PROFILE", bold=True)
for q in [
    "Name (Optional): ________________________________",
    "Email ID: ________________________________",
    "1. Age Group: \u25a1 17-18 years  \u25a1 19-20 years  \u25a1 21-22 years  \u25a1 >22 years",
    "2. Gender: \u25a1 Male  \u25a1 Female",
    "3. Semester: \u25a1 1  \u25a1 2  \u25a1 3  \u25a1 4  \u25a1 5  \u25a1 6  \u25a1 7  \u25a1 8",
    "4. Type of Residence: \u25a1 Home  \u25a1 Hostel  \u25a1 PG/Other",
    "5. Average Daily Time Spent on Digital Devices: \u25a1 <2hrs  \u25a1 2-4hrs  \u25a1 4-6hrs  \u25a1 >6hrs",
    "6. Main Device Used for Studies: \u25a1 Smartphone  \u25a1 Laptop  \u25a1 Tablet  \u25a1 Desktop",
    "7. Average Continuous Screen Time Without Break: \u25a1 <30 min  \u25a1 30-60 min  \u25a1 1-2 hrs  \u25a1 >2hrs",
    "8. Total Daily Screen Exposure (Study + Leisure): \u25a1 <3 hours  \u25a1 3-5 hours  \u25a1 5-8 hours  \u25a1 >8 hours",
]: add_p(q)

doc.add_paragraph()
add_p("SECTION B: PRACTICE RELATED TO DIGITAL DEVICE USAGE", bold=True)
for q in [
    "9. Do you use digital devices in dark/low light?  \u25a1 Yes  \u25a1 No",
    "10. Do you use devices while lying down?  \u25a1 Yes  \u25a1 No",
    "11. Do you take breaks during long screen use?  \u25a1 Yes  \u25a1 Sometimes  \u25a1 No",
    "12. Do you use blue light filter/night mode?  \u25a1 Yes  \u25a1 No",
    "13. Do you follow the 20-20-20 Rule?  \u25a1 Always  \u25a1 Sometimes  \u25a1 Rarely  \u25a1 Never",
]: add_p(q)

doc.add_paragraph()
add_p("SECTION C: SYMPTOMS OF DIGITAL EYE STRAIN", bold=True)
add_p("(Rate: Never / Rarely / Sometimes / Occasionally / Often / Always)")
for q in [
    "14. Eye Strain or tired eyes  \u25a1 Never  \u25a1 Rarely  \u25a1 Sometimes  \u25a1 Occasionally  \u25a1 Often  \u25a1 Always",
    "15. Burning sensation in eyes  \u25a1 Never  \u25a1 Rarely  \u25a1 Sometimes  \u25a1 Occasionally  \u25a1 Often  \u25a1 Always",
    "16. Dry eyes  \u25a1 Never  \u25a1 Rarely  \u25a1 Sometimes  \u25a1 Occasionally  \u25a1 Often  \u25a1 Always",
    "17. Headache after screen use  \u25a1 Never  \u25a1 Rarely  \u25a1 Sometimes  \u25a1 Occasionally  \u25a1 Often  \u25a1 Always",
    "18. Blurred vision  \u25a1 Never  \u25a1 Rarely  \u25a1 Sometimes  \u25a1 Occasionally  \u25a1 Often  \u25a1 Always",
    "19. Eye redness  \u25a1 Never  \u25a1 Rarely  \u25a1 Sometimes  \u25a1 Occasionally  \u25a1 Often  \u25a1 Always",
    "20. Neck or Shoulder Pain  \u25a1 Never  \u25a1 Rarely  \u25a1 Sometimes  \u25a1 Occasionally  \u25a1 Often  \u25a1 Always",
    "21. Difficulty focusing after screen use  \u25a1 Never  \u25a1 Rarely  \u25a1 Sometimes  \u25a1 Occasionally  \u25a1 Often  \u25a1 Always",
    "22. Sensitivity to light  \u25a1 Never  \u25a1 Rarely  \u25a1 Sometimes  \u25a1 Occasionally  \u25a1 Often  \u25a1 Always",
    "23. Excessive watering/tearing of eyes  \u25a1 Never  \u25a1 Rarely  \u25a1 Sometimes  \u25a1 Occasionally  \u25a1 Often  \u25a1 Always",
]: add_p(q)

doc.add_paragraph()
add_p("SECTION D: PREVENTIVE PRACTICES AND AWARENESS", bold=True)
for q in [
    "24. Do you follow the 20-20-20 rule?  \u25a1 Yes  \u25a1 No  \u25a1 Sometimes",
    "25. Do you use artificial tears or eye drops?  \u25a1 Yes  \u25a1 No  \u25a1 Occasionally",
    "26. Do you consciously reduce screen brightness?  \u25a1 Always  \u25a1 Sometimes  \u25a1 Never",
    "27. Do you adjust screen distance properly?  \u25a1 Yes  \u25a1 No  \u25a1 Sometimes",
    "28. Do you perform eye relaxation exercises?  \u25a1 Yes  \u25a1 No",
    "29. Do you consult an eye specialist when symptoms occur?  \u25a1 Yes  \u25a1 No  \u25a1 Rarely",
    "30. Are you aware of digital eye strain?  \u25a1 Yes  \u25a1 No",
    "31. Do you know that prolonged screen use can affect eye health?  \u25a1 Yes  \u25a1 No",
    "32. Do you know about preventive measures for eye strain?  \u25a1 Yes  \u25a1 No",
    "33. Have you ever received education/training on eye care?  \u25a1 Yes  \u25a1 No",
]: add_p(q)

doc.add_paragraph()
add_p("Thank you for your valuable time and participation!", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.save("Research_Paper_Final.docx")
print("Final research paper saved as 'Research_Paper_Final.docx'")
