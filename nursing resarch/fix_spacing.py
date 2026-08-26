import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from copy import deepcopy

# Load the reviewed document
doc = Document('Document3_Reviewed.docx')

print("=== CLEANING UP SPACING ===")
print(f"Before: {len(doc.paragraphs)} paragraphs")

# Remove consecutive empty paragraphs (keep max 1)
body = doc.element.body
paras_to_remove = []
prev_was_empty = False

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if not text:
        if prev_was_empty:
            paras_to_remove.append(para._element)
        prev_was_empty = True
    else:
        prev_was_empty = False

# Remove the extra empty paragraphs
for elem in paras_to_remove:
    body.remove(elem)

print(f"Removed {len(paras_to_remove)} extra empty paragraphs")
print(f"After cleanup: {len(doc.paragraphs)} paragraphs")

# Count remaining empty paragraphs
empty_count = sum(1 for p in doc.paragraphs if not p.text.strip())
print(f"Remaining empty paragraphs: {empty_count}")

# Check for remaining AI patterns
print("\n=== AI DETECTION CHECK ===")
ai_patterns = [
    "plays an important role",
    "it is important to note",
    "in today's world",
    "the present study",
    "the findings revealed",
    "the authors concluded",
    "the study aimed to",
    "this highlights the importance",
    "the study was conducted",
    "the study demonstrated",
    "furthermore,",
    "moreover,",
    "additionally,",
    "in conclusion,",
    "overall, the findings",
]

ai_count = 0
for i, para in enumerate(doc.paragraphs):
    text = para.text.lower()
    for pattern in ai_patterns:
        if pattern in text:
            print(f"  [{i}] Found: '{pattern}' in: {para.text[:80]}...")
            ai_count += 1
            break

print(f"\nTotal AI-pattern matches found: {ai_count}")

# Save the cleaned document
output_path = 'C:\\Users\\abhiy\\OneDrive\\Desktop\\Document3_Final.docx'
doc.save(output_path)
print(f"\nSaved to: {output_path}")
print("Done!")
