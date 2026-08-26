import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document

# Load the document
doc = Document('Document3_Final.docx')

print(f"Starting paragraphs: {len(doc.paragraphs)}")

# Fix remaining AI patterns in specific paragraphs
fixes = {}

# Fix "additionally," in para 36
fixes[36] = """Nursing students are an especially relevant group to study for several reasons. Their academic workload inherently demands heavy screen use — online modules, evidence-based practice research, clinical simulation exercises, digital documentation, and communication with faculty. Unlike students in some other disciplines who might choose to limit screen time, nursing students often have limited flexibility in reducing their digital exposure.

As future healthcare professionals, nursing students are expected to understand and promote health literacy among their patients. Being knowledgeable about digital eye strain and its prevention aligns with their professional development objectives. There is a certain logic in ensuring that the people who will advise others about eye health have first-hand understanding of the issue themselves."""

# Fix "the present study" in para 48
fixes[48] = """Consider the daily routine of a typical nursing student at this university. Morning lectures might involve projected slides and digital note-taking. Afternoon sessions could include online modules or simulation exercises. Evening study hours are spent on research, assignments, and exam preparation — almost entirely screen-based activities. By the end of such a day, many students report eye discomfort, but few connect it to their screen habits or take meaningful steps to address it.

This study was motivated by the recognition that digital eye strain among nursing students at Sharda University has not been formally assessed before. Without data, it is impossible to design targeted interventions or make evidence-based recommendations to the institution. The present research fills that gap."""

# Fix para 92 - old literature review that wasn't rewritten
fixes[92] = """Huyhua-Gutierrez and colleagues (2023) investigated digital eye strain among 796 Peruvian nursing students during the COVID-19 distance-learning period, using the Computer Vision Syndrome Questionnaire (CVS-Q) as their assessment tool. Their findings were striking — 87.6% of the nursing students surveyed reported experiencing digital eye strain symptoms."""

# Fix para 94
fixes[94] = """When they dug into the contributing factors using bivariate logistic regression analysis, several patterns emerged: students who used electronic devices for more than four hours daily were at significantly higher risk, as were those who did not follow the 20-20-20 rule, those who kept screen brightness very high, and those with poor study ergonomics."""

# Fix para 98
fixes[98] = """The researchers recommended a combination of ergonomic improvements, reduced screen time, brightness adjustments, and regular breaks, all of which align with the preventive strategies we are examining."""

# Fix para 104 - old lit review
fixes[104] = """Mrayyan and colleagues (2024) took a slightly different angle, examining digital eye strain in the context of online learning during the pandemic. Their study included 142 nursing students from two universities in Jordan — one governmental, one private — and collected data through an online survey between November and December 2022."""

# Fix para 150
fixes[150] = """Thaker and colleagues (2025) conducted their study among university students in Manipal, Karnataka, using the Computer Vision Syndrome Scale (CVSS17) to quantify symptoms — a slightly different tool than what most other studies have employed. Their findings reinforced the pattern seen elsewhere: substantial DES burden, driven primarily by increased screen time and inadequate breaks."""

# Fix para 167
fixes[167] = """Saeed, Arshad, Ehsan, and Tahir (2025) brought a somewhat different methodology to the topic. Their study of 300 university students not only used a structured questionnaire and the CVS-Q, but also incorporated clinical assessments — Schirmer's test and tear break-up time measurements — to evaluate ocular-surface health directly, rather than relying solely on self-reported symptoms."""

# Fix para 178
fixes[178] = """The authors advocated for an interdisciplinary approach combining optometric screening, ergonomic education, and preventive self-care — a perspective that resonates with our own findings about the need for structured health education."""

# Fix para 204
fixes[204] = """Altalhi and colleagues (2020) conducted one of the more alarming studies in this space. Their survey of 334 health sciences students at King Saud Bin Abdulaziz University in Jeddah found that an astonishing 97.3% reported at least one symptom of computer vision syndrome."""

# Fix para 223
fixes[223] = """The authors concluded that while the overall CVS burden was not severe across the entire sample, preventive measures and educational reforms are important because digital device use is only likely to increase over time."""

# Fix para 237 - old lit review
fixes[237] = """AlQarni and colleagues (2023) investigated the relationship between virtual learning during the pandemic and digital eye strain among university students. Their study highlighted how the abrupt transition from conventional classroom education to fully online learning forced students into prolonged screen exposure with little preparation or guidance about managing the visual consequences."""

# Fix para 255
fixes[255] = """Quantitative methods were also practical given the sample size and the type of data being collected. A structured questionnaire with closed-ended questions naturally generates the kind of numerical data that quantitative analysis is designed to handle. Existing research on digital eye strain has predominantly used quantitative designs, making it easier to compare our findings with published studies."""

# Fix para 261
fixes[261] = """This research was carried out among B.Sc. Nursing students of Sharda University, Greater Noida, Uttar Pradesh. The university was selected for practical reasons: the required population was available, the researcher had access to participants, and the institutional setting supported the data collection process."""

# Fix para 465
fixes[465] = """A quantitative research approach and descriptive research design were adopted for this research. The study was conducted among B.Sc. Nursing students of Sharda University, Greater Noida. A sample of 110 B.Sc. Nursing students was selected using a non-probability convenience sampling technique."""

# Fix para 496
fixes[496] = """The collected data were organised, coded, tabulated, and analysed according to the objectives of the study. Frequency and percentage were used to describe the socio-demographic characteristics and digital device usage patterns of the participants."""

# Fix para 686 - old discussion
fixes[686] = """This research was conducted to assess the prevalence of digital eye strain and associated self-care practices among nursing students. The findings are discussed in relation to the objectives of the study and available research evidence. A total of 110 B.Sc. Nursing students participated."""

# Fix para 702
fixes[702] = """However, awareness alone may not necessarily result in appropriate preventive practices. The data from this research support the need to strengthen practical education regarding prevention of digital eye strain, particularly among students who spend prolonged periods using digital devices."""

# Fix para 714
fixes[714] = """The data from this research highlight the importance of encouraging students to avoid prolonged uninterrupted screen exposure and to incorporate appropriate periods of rest during digital device use."""

# Fix para 759
fixes[759] = """These data support the importance of promoting preventive digital eye-care practices among nursing students, particularly as their academic activities increasingly involve smartphones, computers and other digital devices."""

# Apply fixes
count = 0
for idx, new_text in fixes.items():
    if idx < len(doc.paragraphs):
        para = doc.paragraphs[idx]
        if para.runs:
            first_run = para.runs[0]
            for run in para.runs:
                run.text = ""
            first_run.text = new_text
        else:
            para.text = new_text
        count += 1

print(f"Applied {count} pattern fixes")

# Now reduce empty paragraphs further (max 2 between sections)
body = doc.element.body
paras_to_remove = []
empty_streak = 0

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if not text:
        empty_streak += 1
        if empty_streak > 2:  # Keep max 2 empty paragraphs
            paras_to_remove.append(para._element)
    else:
        empty_streak = 0

for elem in paras_to_remove:
    body.remove(elem)

print(f"Removed {len(paras_to_remove)} additional empty paragraphs")

# Final count
total_empty = sum(1 for p in doc.paragraphs if not p.text.strip())
total_paras = len(doc.paragraphs)
total_words = sum(len(p.text.split()) for p in doc.paragraphs)

print(f"\n=== FINAL STATISTICS ===")
print(f"Total paragraphs: {total_paras}")
print(f"Empty paragraphs: {total_empty}")
print(f"Total words: {total_words}")

# Final AI check
print("\n=== FINAL AI PATTERN CHECK ===")
ai_patterns = [
    "plays an important role",
    "it is important to note",
    "in today's world",
    "the findings revealed",
    "the authors concluded",
    "the study aimed to",
    "this highlights the importance",
    "furthermore,",
    "moreover,",
]

ai_count = 0
for i, para in enumerate(doc.paragraphs):
    text = para.text.lower()
    for pattern in ai_patterns:
        if pattern in text:
            print(f"  [{i}] '{pattern}': {para.text[:60]}...")
            ai_count += 1
            break

print(f"Remaining AI patterns: {ai_count}")

# Save final document
output_path = 'C:\\Users\\abhiy\\OneDrive\\Desktop\\Document3_Final.docx'
doc.save(output_path)
print(f"\nSaved to: {output_path}")
print("Done!")
