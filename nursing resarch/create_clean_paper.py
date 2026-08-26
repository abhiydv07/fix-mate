import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Load original to get content
old_doc = Document('Document3_Final.docx')

# Create new document
new_doc = Document()

# Set default font
style = new_doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Set paragraph spacing
paragraph_format = style.paragraph_format
paragraph_format.space_after = Pt(6)
paragraph_format.space_before = Pt(0)
paragraph_format.line_spacing = 1.5

# Get all non-empty paragraphs with their styles
print("Copying content...")
count = 0
for para in old_doc.paragraphs:
    text = para.text.strip()
    if text:  # Only copy non-empty paragraphs
        # Get original style
        original_style = para.style.name if para.style else 'Normal'
        
        # Create new paragraph
        new_para = new_doc.add_paragraph()
        
        # Set style based on original
        if 'Heading 1' in original_style:
            new_para.style = new_doc.styles['Heading 1']
        elif 'Heading 2' in original_style:
            new_para.style = new_doc.styles['Heading 2']
        elif 'List Paragraph' in original_style:
            new_para.style = new_doc.styles['List Paragraph']
        else:
            new_para.style = new_doc.styles['Normal']
        
        # Copy text
        new_para.text = text
        
        # Copy formatting from first run if exists
        if para.runs:
            first_run = para.runs[0]
            if first_run.font.name:
                new_para.runs[0].font.name = first_run.font.name
            if first_run.font.size:
                new_para.runs[0].font.size = first_run.font.size
            if first_run.font.bold is not None:
                new_para.runs[0].font.bold = first_run.font.bold
            if first_run.font.italic is not None:
                new_para.runs[0].font.italic = first_run.font.italic
        
        count += 1

print(f"Copied {count} paragraphs")

# Save new document
output_path = 'C:\\Users\\abhiy\\OneDrive\\Desktop\\Document3_Clean.docx'
new_doc.save(output_path)

# Verify
verify_doc = Document(output_path)
total_paras = len(verify_doc.paragraphs)
empty_count = sum(1 for p in verify_doc.paragraphs if not p.text.strip())
total_words = sum(len(p.text.split()) for p in verify_doc.paragraphs)

print(f"\n=== VERIFICATION ===")
print(f"Total paragraphs: {total_paras}")
print(f"Empty paragraphs: {empty_count}")
print(f"Total words: {total_words}")
print(f"\nSaved to: {output_path}")
print("Done!")
